#!/usr/bin/python3
# -*- coding: utf-8 -*-

from gnr.web.gnrbaseclasses import BaseComponent
from gnr.core.gnrdecorator import public_method
from gnr.core.gnrbag import Bag
from docx import Document
from datetime import datetime
import copy
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.shared import qn
from time import sleep
class View(BaseComponent):

    def th_struct(self,struct):
        r = struct.view().rows()
        r.fieldcell('_row_count', counter=True, name='N.',width='3em')
        r.fieldcell('sof_id')
        r.fieldcell('date')
        r.fieldcell('targa')
        r.fieldcell('tara')
        r.fieldcell('lordo')
        r.fieldcell('netto')
        r.fieldcell('total')

    def th_order(self):
        return 'sof_id'

    def th_query(self):
        return dict(column='id', op='contains', val='')

class ViewFromTrucks(BaseComponent):

    def th_struct(self,struct):
        r = struct.view().rows()
        r.fieldcell('_row_count', counter=True, name='N.',width='3em')
        r.fieldcell('sof_id')
        r.fieldcell('date', edit=dict(hasDownArrow=True),batch_assign=True)
        r.fieldcell('targa', edit=True)
        r.fieldcell('tara', edit=True)
        r.fieldcell('lordo', edit=True)
        r.fieldcell('netto', edit=True)
        r.fieldcell('total', edit=True)

    def th_order(self):
        return '_row_count'

    def th_query(self):
        return dict(column='id', op='contains', val='')
    
    def th_options(self):
        return dict(grid_selfDragRows=True)
    
    def th_view(self,view):
        bar = view.top.bar.replaceSlots('addrow','addrow,resourcePrints,10,importa_trucks,10,outturn,10,batchAssign')
        btn_importa_trucks = bar.importa_trucks.paletteImporter(paletteCode='xls_importer',
                            dockButton_iconClass=False,
                            title='!!Importa distinta trasporti',
                            importButton_label='Importa distinta trasporti',
                            previewLimit=500,
                            dropMessage='Trascina qui il tuo file o clicca per cercarlo', filetype='excel',
                            importButton_action="genro.publish('import_trucks',{filepath:imported_file_path})",
                            matchColumns='*')
        bar.dataRpc(self.importaDistinta, subscribe_import_trucks=True,record='=#FORM.record',
                      _onResult="genro.publish('xls_importer_onResult',result);this.form.reload();",
                      _onError="genro.publish('xls_importer_onResult',{error:error});")
        btn_outturn=bar.outturn.button('!![en]Outturn report')
        btn_outturn.dataRpc('nome_temp',self.outturn,record='=#FORM.record',_virtual_column='$workport',_lockScreen=dict(thermo=True))#,_lockScreen=dict(message='Please Wait')) 
      
       
    @public_method
    def importaDistinta(self, filepath=None, record=None, **kwargs):
        "Importa Excel distinta trasporti"
             
        reader = self.utils.getReader(filepath)
        result = Bag()
        sof_id = record['id']
        progressivo=0
        for row in reader():
            if row['netto']: 
                progressivo += row['netto']    
            else:
                progressivo = None    
            new_distinta = self.db.table('shipsteps.sof_trucks').newrecord(sof_id=sof_id,date=row['data'],plate=row['targa'],tare=row['tara'],gross=row['lordo'],
                        net=row['netto'],total=progressivo, **row)

            self.db.table('shipsteps.sof_trucks').insert(new_distinta)
            result.addItem('inserted_row', row)
            
        self.db.commit()
        return result
    
    @public_method
    def outturn(self,record, **kwargs): 
        tbl_arrival=self.db.table('shipsteps.arrival')
        tbl_sof_trucks = self.db.table('shipsteps.sof_trucks')
        tbl_sof = self.db.table('shipsteps.sof')
        tbl_cargo = self.db.table('shipsteps.cargo_unl_load')
        record_arrival = tbl_arrival.record(id=record['arrival_id'], ignoreMissing=True, virtual_columns='workport,vesselname,imo').output('bag') 
        record_trucks = tbl_sof_trucks.query(columns="*",
                         where='$sof_id = :sofid',sofid=record['id']).fetch()
        record_cargo = tbl_cargo.query(columns="*",
                         where='$arrival_id = :arrid',arrid=record_arrival['id']).fetch()
        
        record_sof = tbl_sof.record(id=record['id'], ignoreMissing=True, virtual_columns='carico_sofbl,shortage,tot_mov,tot_cargo_sof').output('bag') 
        sof_cargo=record_sof['@sof_cargo_sof']
        cargo=[]
        for r in sof_cargo.keys():
            bl_mis=sof_cargo[r]['@cargo_unl_load_id']['@measure_id.description']
            bl_qt=str(sof_cargo[r]['@cargo_unl_load_id']['quantity'])
            bl_qt = f"{str(bl_qt).replace('.',',')}"
            bl_descr=sof_cargo[r]['@cargo_unl_load_id']['description']
            bln=sof_cargo[r]['@cargo_unl_load_id']['bln']
            bl_date=str(sof_cargo[r]['@cargo_unl_load_id']['bl_date'].strftime("%d/%m/%Y"))
            bl_issue=sof_cargo[r]['@cargo_unl_load_id']['@place_origin_goods.citta_nazione']
            cargo.append(bl_mis + ' ' + bl_qt + ' ' + bl_descr + ' - BL no.' + bln + ' issued ' + bl_issue + ' ' + bl_date )
        #print(x)
        workport=record_arrival['workport']
        vessel=record_arrival['vesselname']
        imo=record_arrival['imo']
        banchina=record_arrival['@dock_id.dock_name']
        #cargo=record_sof['carico_sofbl']
        ops_commenced=record_sof['ops_commenced'].strftime("%d/%m/%Y %H:%M")
        ops_completed=record_sof['ops_completed'].strftime("%d/%m/%Y %H:%M")
        tot_cargo=str(round(record_sof['tot_cargo_sof'],3))
        tot_cargo = f"{str(tot_cargo).replace('.',',')}" #con replace inseriamo la virgola al posto del punto
        outturn=str(round(record_sof['tot_mov'],3))
        outturn = f"{str(outturn).replace('.',',')}"
        shortage=str(round(record_sof['shortage'],3))
        shortage = f"{str(shortage).replace('.',',')}"
        shortage_perc=str(round(record_sof['shortage']/record_sof['tot_cargo_sof']*100,3))+str('%')
        shortage_perc = f"{str(shortage_perc).replace('.',',')}"
        nome_file = 'Outturn.docx'
        nome_file_out = 'Outturn_filled.docx'
        file_sn_out = self.site.storageNode('home:form_standard', 'Outturn_filled.docx')
        data_report = record['ops_completed'].strftime("%d/%m/%Y")
        
        variables = {
          "${port}": workport,    
          "${date}": data_report,
          "${vessel}": vessel,
          "${imo}": imo,
          "${banchina}": banchina,
          "${ops_comm}": ops_commenced,
          "${ops_compl}": ops_completed,
          "${tot_cargo}": tot_cargo,
          "${tot_mov}": outturn,
          "${shortage}": shortage,
          "${shortage_perc}": shortage_perc,
        }
        file_sn = self.site.storageNode('home:form_standard', nome_file)
        template_file_path = file_sn.internal_path
              
        output_file_path = file_sn_out.internal_path

        template_document = Document(template_file_path)

        for variable_key, variable_value in variables.items():
            for paragraph in template_document.paragraphs:
                self.replace_text_in_paragraph(paragraph, variable_key, variable_value)

            for table in template_document.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            self.replace_text_in_paragraph(paragraph, variable_key, variable_value)
        
        table_cargo = template_document.tables[0]
        for r in cargo:
            row_cells = table_cargo.add_row()
            row_cells.cells[0].text = str(r)

        table = template_document.tables[2]
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        #borders = OxmlElement('w:tblBorders')
        #bottom_border = OxmlElement('w:bottom')
        #bottom_border.set(qn('w:val'), 'single')
        #bottom_border.set(qn('w:sz'), '4')
        #borders.append(bottom_border)
        #table._tbl.tblPr.append(borders)
       

        col1 = table.columns[1]
        cell = col1.cells[0]
        
        paragr = cell.paragraphs[0]
        style = template_document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(8)
        
        for r in self.utils.quickThermo(record_trucks):#, labelfield='_row_count'):
            row_cells = table.add_row()  
            row_cells.cells[0].text = str(r['_row_count'])
            row_cells.cells[0].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
            row_cells.cells[1].text = str(r['date'].strftime("%d/%m/%Y"))
            row_cells.cells[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
            row_cells.cells[2].text = str(r['targa'])
            row_cells.cells[2].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
            row_cells.cells[3].text = str(r['tara'])
            row_cells.cells[3].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
            row_cells.cells[4].text = str(r['lordo'])
            row_cells.cells[4].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
            row_cells.cells[5].text = str(r['netto'])
            row_cells.cells[5].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
            row_cells.cells[6].text = str(r['total'])
            row_cells.cells[6].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
            print(r)
        
        #for r, rowiteration in enumerate(table.rows):
        #    if r > 0: #qui saltiamo il primo rigo della tabella dove abbiamo le intestazioni
        #        for c, cells in enumerate(rowiteration.cells):
        #            if c < 6: #se il numero di cella è più basso di 6 il contenuto viene centrato altrimenti posizionato a destra
        #                table.cell(r,c).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        #            else:
        #                table.cell(r,c).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
                
        #table.cell(1,6).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        
        template_document.save(output_file_path)
        #apriamo direttamente il file salvato con il programma standard di sistema
        filename=output_file_path
              
        path_pdf = self.site.storageNode('home:form_standard', nome_file_out)
        
        #path_pdf=path + str('/') + nome_form + str('.pdf')
        result=self.site.storageNode(path_pdf)
        #print(x)
        self.setInClientData(path='gnr.clientprint',
                              value=result.url(), fired=True)
        #print(x)
        #return nome_form
    def copy_cell_properties(self,source_cell, dest_cell):
        '''Copies cell properties from source cell to destination cell.

        Copies cell background shading, borders etc. in a python-docx Document.

        Args:
            source_cell (docx.table._Cell): the source cell with desired formatting
            dest_cell (docx.table._Cell): the destination cell to which to apply formatting
        '''
        # get properties from source cell
        # (tcPr = table cell properties)
        cell_properties = source_cell._tc.get_or_add_tcPr()
        
        # remove any table cell properties from destination cell
        # (otherwise, we end up with two <w:tcPr> ... </w:tcPr> blocks)
        dest_cell._tc.remove(dest_cell._tc.get_or_add_tcPr())

        # make a shallow copy of the source cell properties
        # (otherwise, properties get *moved* from source to destination)
        cell_properties = copy.copy(cell_properties)

        # append the copy of properties from the source cell, to the destination cell
        dest_cell._tc.append(cell_properties)

    def replace_text_in_paragraph(self,paragraph, key, value):
        if key in paragraph.text:
            inline = paragraph.runs
            for item in inline:
                if key in item.text:
                    item.text = item.text.replace(key, value)

class Form(BaseComponent):

    def th_form(self, form):
        pane = form.record
        fb = pane.formbuilder(cols=2, border_spacing='4px')
        fb.field('sof_id' )
        fb.field('date' )
        fb.field('targa' )
        fb.field('tara' )
        fb.field('lordo' )
        fb.field('netto' )
        fb.field('total' )


    def th_options(self):
        return dict(dialog_height='400px', dialog_width='600px' )
