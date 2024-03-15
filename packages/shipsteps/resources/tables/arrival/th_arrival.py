#!/usr/bin/python3
# -*- coding: utf-8 -*-

from turtle import left
from gnr.web.gnrbaseclasses import BaseComponent
from gnr.core.gnrdecorator import public_method
from gnr.core.gnrdecorator import metadata
from gnr.web.gnrbaseclasses import TableTemplateToHtml
from datetime import datetime
from gnr.core.gnrlang import GnrException
import os.path
from gnr.core.gnrbag import Bag
from docx import Document
import os
import subprocess #per apertura file tramite programma di sistema
#import datetime

class View(BaseComponent):
    
    def th_struct(self,struct):
        "Vista standard"
        r = struct.view().rows()  
        arrival = r.columnset('colset_arrival', name='Arrival', color='white',background='DarkSlateBlue', font_weight='bold')
        arrival.fieldcell('agency_id', width='7em')
        arrival.fieldcell('reference_num', width='8em')
        arrival.fieldcell('protfald', width='7em')
        arrival.fieldcell('visit_id',width='8em')
        arrival.fieldcell('nsis_prot',width='8em')
        arrival.fieldcell('voy_n', width='4em')
        arrival.fieldcell('date', width='5em')
        arrival.fieldcell('vessel_details_id', width='15em', font_weight='bold')
        arrival.fieldcell('tip_mov', width='6em', font_weight='bold')
        expect = r.columnset('colset_expect', name='Expected Times', color='white',background='SeaGreen', font_weight='bold')
        expect.fieldcell('eta', width='5em', Short=True)
        expect.fieldcell('etb', width='5em')
        expect.fieldcell('et_start', width='5em')
        expect.fieldcell('etc', width='5em')
        expect.fieldcell('ets', width='5em')
        draft = r.columnset('colset_draft', name='Draft', color='white',background='SaddleBrown', font_weight='bold')
        draft.fieldcell('draft_aft_arr', width='3em')
        draft.fieldcell('draft_fw_arr', width='3em')
        draft.fieldcell('draft_aft_dep', width='3em')
        draft.fieldcell('draft_fw_dep', width='3em')
        extra = r.columnset('colset_extra', name='Signature / Mandatory / Invoicing', color='white',background='Tan', font_weight='bold')
        extra.fieldcell('firma_div',width='5em')
        extra.fieldcell('mandatory', width='16em')
        extra.fieldcell('invoice_det_id', width='30em')
        berth = r.columnset('colset_berth', name='Berth', color='white',background='Gray', font_weight='bold')
        berth.fieldcell('dock_id', width='5em')
        berth.fieldcell('info_moor', width='6em')
        member = r.columnset('colset_member', name='People on board', color='white',background='DarkCyan', font_weight='bold')
        member.fieldcell('master_name', width='6em')
        member.fieldcell('n_crew', width='3em')
        member.fieldcell('n_passengers', width='3em', name='Pax no.')
        lastport = r.columnset('colset_lastport', name='Last Port', color='white',background='SlateGray', font_weight='bold')
        lastport.fieldcell('last_port', width='7em')
        lastport.fieldcell('departure_lp', width='5em')
        nextport = r.columnset('colset_nextport', name='Next Port', color='white',background='Maroon', font_weight='bold')
        nextport.fieldcell('next_port', width='7em')
        nextport.fieldcell('eta_np', width='5em')
        gpg = r.columnset('colset_gpg', name='GPG', color='Black',background='lightgray', font_weight='bold')
        gpg.fieldcell('@gpg_arr.n_gpg', name= 'GPG n.')
        gpg.fieldcell('@gpg_arr.date_start', name= 'GPG start', width='5em')
        gpg.fieldcell('@gpg_arr.date_end', name= 'GPG end', width='5em')
       #r.fieldcell('@cargo_lu_arr.description', name= 'cargo_descr', width='5em')
        #r.fieldcell('carico', name='Descrizione Carico',width='30em')
        #r.fieldcell('cargo', width='50em')
        #r.fieldcell('@cargo_lu_arr.cargo_on_board', width='50em')
        cargo = r.columnset('colset_cargo', name='Cargo', color='white',background='RosyBrown', font_weight='bold')
        cargo.fieldcell('cargo_dest')
        cargo.fieldcell('cargo_lu_en', width='30em')
        cargo.fieldcell('@cargo_lu_arr.tot_cargo', width='8em')
        cargo.fieldcell('ship_rec', width='30em')
        
    def th_struct_clientecarico(self,struct):
        "Vista ClienteCarico"
        r = struct.view().rows()
        arrival = r.columnset('colset_arrival', name='Arrival', color='white',background='DarkSlateBlue', font_weight='bold')
        arrival.fieldcell('agency_id', width='7em')
        arrival.fieldcell('reference_num', width='8em')
        arrival.fieldcell('vessel_details_id', width='15em', font_weight='bold')
        arrival.fieldcell('tip_mov', width='6em', font_weight='bold')
        cargo = r.columnset('colset_cargo', name='Cargo', color='white',background='RosyBrown', font_weight='bold')
        cargo.fieldcell('cargo_dest')
        cargo.fieldcell('cargo_lu_en', width='30em')
        cargo.fieldcell('@cargo_lu_arr.tot_cargo', width='8em', totalize=True)
        cargo.fieldcell('ship_rec', width='30em')
        cargo.fieldcell('ship_or_rec', width='30em')
        cargo.fieldcell('cargo_op', width='5em')

    def th_hiddencolumns(self):
        return "$cargo_op"
    
    @metadata(variable_struct=True)
    def th_sections_anno(self):
        #f = self.db.table('shipsteps.arrival').query(columns='$etb',where='$etb is not null', order_by='$etb').selection().output('records')
        f = self.db.table('shipsteps.arrival').query(columns="to_char($etb,'YYYY')",where='$etb is not null',distinct="to_char($etb,'YYYY')").fetch()
        
        result=[]
        result.append(dict(code='tutti',caption='!![en]All'))
        for r in f:
            result.append(dict(code=r[0],caption=r[0],condition="to_char($etb,'YYYY')=:anno",condition_anno=r[0]))
        return result

   #@metadata(variable_struct=True)
   #def th_sections_operation(self):
   #    #prendiamo agency_id nel currentEnv
   #    ag_id=self.db.currentEnv.get('current_agency_id')
   #    #effettuaiamo la ricerca di tutti i clienti filtrando quelli relativi all'agency_id
   #    f = self.db.table('shipsteps.cargo_unl_load').query(where="$operation = 'U' or $operation = 'L'",order_by='@receiver_id.name ,@shipper_id.name').selection().output('records')
   #    #creaiamo una lista vuota dove andremo ad appendere i dizionari con il valore tutti e con i clienti
   #    shiprec = self.db.table('shipsteps.ship_rec').query().selection().output('records')
   #    
   #    #print(x)
   #    result=[]
   #    Dup={}
   #    result.append(dict(code='tutti',caption='!![en]All'))
   #    for r in range(len(f)):
   #        if f[r]['operation'] == 'U': 
   #            for rec in shiprec:
   #                if rec['id'] ==  f[r]['receiver_id']:
   #                    if rec['id'] in Dup: #verifichiamo se in result abbiamo già inserito il ricevitore per non creare un duplicato
   #                        ItemNumber = Dup[rec['id']]
   #                    else:
   #                        result.append(dict(code=rec['id'],caption=rec['name'],
   #                        condition='$shiprec=:receiver',condition_receiver=rec['name']))
   #                        Dup[rec['id']] = ItemNumber = len(result)-1
   #                    
   #        if f[r]['operation'] == 'L':
   #            for rec in shiprec:
   #                if rec['id'] ==  f[r]['shipper_id']:
   #                    if rec['id'] in Dup: #verifichiamo se in result abbiamo già inserito il caricatore per non creare un duplicato
   #                        ItemNumber = Dup[rec['id']]
   #                    else:
   #                        result.append(dict(code=rec['id'],caption=rec['name'],
   #                        condition='$shiprec=:shipper',condition_shipper=rec['name']))
   #                        Dup[rec['id']] = ItemNumber = len(result)-1
   #    
   #    return result
    
    #@metadata(variable_struct=True)
    #def th_sections_shiprec(self):
    #    return [dict(code='tutti',caption='!![en]All cargo handled'),
    #            dict(code='unload',caption='!![en]Unloading',
    #                    condition="$cargo_op='U'"),
    #            dict(code='load',caption='!![en]Loading',condition="$cargo_op='L'")] 
    #
    def th_top_toolbarsuperiore(self,top):
         top.slotToolbar('5,*,sections@tip_mov,*,5',
                         childname='superiore',_position='<bar',gradient_from='#999',gradient_to='#666')
        
    def th_order(self):
        return 'reference_num:d' 

    def th_query(self):
        return dict(column='@vessel_details_id.@imbarcazione_id.nome', op='contains', val='', runOnStart=True)

    def th_options(self):
        #se l'utente connesso ha i privilegi di admin/superadmin/sviluppatore visualizzerà il tasto - per la cancellazione del record
        
        usertags = self.db.currentEnv.get('userTags')
    
        if 'admin' in usertags or 'superadmin' in usertags or '_DEV_' in usertags:
            del_row = True
        else:
            del_row = False
        
        return dict(view_preview_tpl='dati_nave',partitioned=True, delrow=del_row)
 
class View_Filtered_Arrivals(BaseComponent):
    def th_struct(self,struct):
        
        r = struct.view().rows()
        arrival = r.columnset('colset_arrival', name='Arrival', color='white',background='DarkSlateBlue', font_weight='bold')
        arrival.fieldcell('agency_id', width='7em')
        arrival.fieldcell('reference_num', width='8em')
        arrival.fieldcell('vessel_details_id', width='15em', font_weight='bold')
        arrival.fieldcell('tip_mov', width='6em', font_weight='bold')
        cargo = r.columnset('colset_cargo', name='Cargo', color='white',background='RosyBrown', font_weight='bold')
        cargo.fieldcell('cargo_dest')
        cargo.fieldcell('cargo_lu_en', width='30em')
        cargo.fieldcell('@cargo_lu_arr.tot_cargo', width='8em', totalize=True)
        cargo.fieldcell('ship_rec', width='30em')
        cargo.fieldcell('ship_or_rec', width='30em')
        cargo.fieldcell('cargo_op', width='5em')

    def th_hiddencolumns(self):
        return "$cargo_op"
    
    @metadata(variable_struct=True)
    def th_sections_anno(self):
        #f = self.db.table('shipsteps.arrival').query(columns='$etb',where='$etb is not null', order_by='$etb').selection().output('records')
        f = self.db.table('shipsteps.arrival').query(columns="to_char($etb,'YYYY')",where='$etb is not null',distinct="to_char($etb,'YYYY')").fetch()
        
        result=[]
        result.append(dict(code='tutti',caption='!![en]All'))
        for r in f:
            result.append(dict(code=r[0],caption=r[0],condition="to_char($etb,'YYYY')=:anno",condition_anno=r[0]))
        return result

    @metadata(variable_struct=True)
    def th_sections_operation(self):
        #prendiamo agency_id nel currentEnv
        ag_id=self.db.currentEnv.get('current_agency_id')
        #effettuaiamo la ricerca di tutti i clienti filtrando quelli relativi all'agency_id
        f = self.db.table('shipsteps.cargo_unl_load').query(where="$operation = 'U' or $operation = 'L'",order_by='@receiver_id.name ,@shipper_id.name').selection().output('records')
        #creaiamo una lista vuota dove andremo ad appendere i dizionari con il valore tutti e con i clienti
        shiprec = self.db.table('shipsteps.ship_rec').query().selection().output('records')
        
        #print(x)
        result=[]
        Dup={}
        result.append(dict(code='tutti',caption='!![en]All'))
        for r in range(len(f)):
            if f[r]['operation'] == 'U': 
                for rec in shiprec:
                    if rec['id'] ==  f[r]['receiver_id']:
                        if rec['id'] in Dup: #verifichiamo se in result abbiamo già inserito il ricevitore per non creare un duplicato
                            ItemNumber = Dup[rec['id']]
                        else:
                            result.append(dict(code=rec['id'],caption=rec['name'],
                            condition='$shiprec=:receiver',condition_receiver=rec['name']))
                            Dup[rec['id']] = ItemNumber = len(result)-1
                        
            if f[r]['operation'] == 'L':
                for rec in shiprec:
                    if rec['id'] ==  f[r]['shipper_id']:
                        if rec['id'] in Dup: #verifichiamo se in result abbiamo già inserito il caricatore per non creare un duplicato
                            ItemNumber = Dup[rec['id']]
                        else:
                            result.append(dict(code=rec['id'],caption=rec['name'],
                            condition='$shiprec=:shipper',condition_shipper=rec['name']))
                            Dup[rec['id']] = ItemNumber = len(result)-1
        
        return result
    
    @metadata(variable_struct=True)
    def th_sections_shiprec(self):
        return [dict(code='tutti',caption='!![en]All cargo handled'),
                dict(code='unload',caption='!![en]Unloading',
                        condition="$cargo_op='U'"),
                dict(code='load',caption='!![en]Loading',condition="$cargo_op='L'")] 
    
    def th_top_toolbarsuperiore(self,top):
        top.slotToolbar('5,sections@tip_mov,*,sections@shiprec,*,sections@anno,*,sections@operation,5',
                        childname='superiore',_position='<bar',gradient_from='#999',gradient_to='#666',sections_operation_multiButton=False,
                        sections_operation_lbl='!![en]Shippers-Receivers',sections_operation_lbl_color='white',
                        sections_operation_width='40em',sections_anno_multiButton=False,
                        sections_anno_lbl='!![en]Year',sections_anno_lbl_color='white',
                        sections_anno_width='6em')
        
    def th_order(self):
        return 'reference_num:d' 

    def th_query(self):
        return dict(column='@vessel_details_id.@imbarcazione_id.nome', op='contains', val='', runOnStart=True)

    def th_options(self):
        #se l'utente connesso ha i privilegi di admin/superadmin/sviluppatore visualizzerà il tasto - per la cancellazione del record
       
        usertags = self.db.currentEnv.get('userTags')
    
        if 'admin' in usertags or 'superadmin' in usertags or '_DEV_' in usertags:
            del_row = True
        else:
            del_row = False
   
        return dict(view_preview_tpl='dati_nave',partitioned=True, delrow=del_row)
    
class Form(BaseComponent):
    py_requires="gnrcomponents/attachmanager/attachmanager:AttachManager"

    def th_form(self, form):
        #con lo store.handler possiamo inserire tutte le virtual_columns che vogliamo avere disponibili nello store
        form.store.handler('load',virtual_columns='$workport,$docbefore_cp,$gdfdep_timeexp,$etb_date,$refcode')
        
        ##all'apertura del form arrival calcoliamo quali sono le partenze finanza non flaggate nella tasklist in modo da notificarle tramite datacontroller 
        #tbl_arrival=self.db.table('shipsteps.arrival')
        #tbl_arr_time=self.db.table('shipsteps.arrival_time')
        #tbl_tasklist=self.db.table('shipsteps.tasklist')
        #arrival=tbl_arrival.query(columns='$id,$reference_num', 
        #                      where='$agency_id=:ag_id', ag_id=self.db.currentEnv.get('current_agency_id')).fetch()
        #text_a='PRINT THE GDF FORM VESSEL DEPARTURE FOR ARRIVALS: '
        #refnum=[]
        #for r in arrival:
        #    sailed=tbl_arr_time.readColumns(columns='$sailed', where='$arrival_id=:a_id', a_id=r['id'])
        #    gdfdep=tbl_tasklist.readColumns(columns='$form_gdfdep', where='$arrival_id=:a_id', a_id=r['id'])
        #    
        #    if sailed is not None:
        #        if sailed < datetime.now() and (gdfdep == None) or (gdfdep == False):
        #            refnum.append(r['reference_num'])
        #text_b=','.join(refnum)
        #
        #html_text=text_a+text_b
        #if len(refnum) > 0:            
        #    form.data('.refnum',html_text)
        #    form.dataController('alert(msg)', msg='=.refnum',_if='msg', _onStart=True)

        tc = form.center.tabContainer(selected='.current_tab')
        bc = tc.borderContainer(title='!![en]<strong>Arrival</strong>')
        tc_car = tc.tabContainer(title='!![en]<strong>Cargo</strong>',region='center')#,hidden='^#FORM.record.@last_port.nazione_code?=!(#v=="IT"||#v=="LM")')#,hidden="^#FORM.record.@last_port.nazione_code?=#v!='IT'")
        bc_extracp = tc.borderContainer(title='!![en]<strong>Extra dati CP</strong>')
        bc_att = tc.borderContainer(title='!![en]<strong>Attachments</strong>')
        tc_task = tc.tabContainer(title='!![en]<strong>Task List</strong>',region='center',selectedPage='^tabname')
        bc_tasklist = tc_task.borderContainer(title="<div style='color:red;'>Task list</div>", region='center')#title='!![en]Task List'
        
        tc_arrtimes = tc.borderContainer(title='!![en]<strong>Arrival Times</strong>')
        #tab arrivals details con tc_arrtimes disabilitati per uso unico tab e form
        #tc_details = tc.tabContainer(title='!![en]<strong>Arrival/Departure Details</strong>')
        #self.arrival_times(tc_details.contentPane(title='!![en]<strong>Vessel Times</strong>',height='100%', background = '#f2f0e8'))
        #self.arr_details(tc_details.contentPane(title='!![en]<strong>Draft/Remains Details</strong>',height='100%', background = '#f2f0e8'))

        tc_undertask = bc_tasklist.tabContainer(margin='2px', region='bottom', height='150px', splitter=True,selectedPage='^tabname')
        tc_sof = tc.borderContainer(title='!![en]<strong>SOF</strong>',selectedPage='^.tabname')
        tc_app = tc.tabContainer(title='!![en]<strong>Applications</strong>')
        tc_bl = tc.borderContainer(title='!![en]<strong>Loading Cargoes</strong>')
        tc_usma = tc.borderContainer(title='!![en]<strong>Sanimare certificates</strong>')
        #tc_parapon = bc_task3.tabContainer(title='pippo')
        self.extraDatiCP(bc_extracp.borderContainer(region='center', splitter=True, background = '#f2f0e8'))
        #self.usmaCert(bc_usma.borderContainer(title='!![en]Renew certificates Sanimare',region='center', splitter=True, background = 'seashell'))
        tc_bl.contentPane(title='!![en]Loading Cargoes',height='100%').remote(self.cargodocsCertLazyMode,_waitingMessage='!![en]Please wait')
        tc_usma.contentPane(title='!![en]Renew certificates Sanimare',height='100%').remote(self.usmaCertLazyMode,_waitingMessage='!![en]Please wait')
        tc_email = tc.borderContainer(title='!![en]<strong>Email in/out</strong>')
        tc_email.contentPane(title='!![en]Email in/out',height='100%').remote(self.emailInOutLazyMode,_waitingMessage='!![en]Please wait')

        self.allegatiArrivo(bc_att.contentPane(title='!![en]Attachments', height='100%'))

        self.datiArrivo(bc.borderContainer(region='top',height='400px', splitter=True, background = '#f2f0e8'))
        #visualizziamo in alto alle pagine tasklist e arrival times gli expected times
        self.times(bc_tasklist.borderContainer(region='top',height='10%', background = 'SlateGrey', splitter=True, closable=True))
        self.times_sof(tc_sof.borderContainer(region='top',height='10%', background = 'SlateGrey', splitter=True, closable=True))
        self.taskList(bc_tasklist.borderContainer(region='center',height='auto', background = '#f2f0e8', splitter=True))
        tc_task.contentPane(title='!![en]Shore pass').remote(self.shorePassLazyMode,_waitingMessage='!![en]Please wait')
        tc_task.contentPane(title='!![en]Pax List').remote(self.paxListLazyMode,_waitingMessage='!![en]Please wait')
        tc_task.contentPane(title='!![en]Vessel Services',pageName='services').remote(self.servicesLazyMode,_waitingMessage='!![en]Please wait')
        
        #self.sof(tc_sof.contentPane(title='!![en]Sof',height='100%'))
        #tc_sof.contentPane(title='!![en]Sof',pageName='sof',height='100%').remote(self.sofLazyMode,_waitingMessage='!![en]Please wait')
        tc_sof.borderContainer(region='center',height='auto', background = '#f2f0e8', splitter=True).contentPane(title='!![en]Sof',pageName='sof',height='100%').remote(self.sofLazyMode,_waitingMessage='!![en]Please wait')
        
        #self.allegatiArrivo(tc_task.contentPane(title='Attachments', region='center', height='100%', splitter=True))
        
        self.garbage(tc_undertask.contentPane(title='!![en]Garbage', pageName='garbage'))
        self.tributi(tc_undertask.contentPane(title='!![en]Tributes HM',pageName='tributi'))
        #self.rinfusa(tc_app.contentPane(title='!![en]Bulk Application'))
        tc_app.contentPane(title='!![en]Bulk Application').remote(self.rinfusaLazyMode,_waitingMessage='!![en]Please wait')
        #self.bunker(tc_app.contentPane(title='!![en]Bunker Application'))
        tc_app.contentPane(title='!![en]Bunker Application').remote(self.bunkerLazyMode,_waitingMessage='!![en]Please wait')
        tc_app.contentPane(title='!![en]Certificates Application').remote(self.certificateLazyMode,_waitingMessage='!![en]Please wait')

        #self.gpg(bc.borderContainer(region='center',datapath='.@gpg_arr',height='300px', splitter=True, background = 'lavenderblush'))

        tc = bc.tabContainer(margin='2px', region='center', height='auto', splitter=True)

        tc_under_car = tc_car.tabContainer(title='!![en]Cargo onboard')
        
        self.carbordoArr(tc_under_car.contentPane(title="!![en]Cargo onboard on arrival",datapath='.record'))
        self.carbordoDep(tc_under_car.contentPane(title='!![en]Cargo onboard on departure',datapath='.record'))

        self.datiCarico(tc_car.contentPane(title='!![en]Cargo loading / unloading'))
        self.datiCaricoTransit(tc_car.contentPane(title='!![en]Transit cargo',datapath='.record'))
        
        self.arrival_details(tc_arrtimes.borderContainer(title='!![en]Arrival/Departure details',height='100%', region='top', background = '#f2f0e8', pageName='arrtime'))

        tc.contentPane(title='!![en]Email Arrival').remote(self.emailArrivalLazyMode,_waitingMessage='!![en]Please wait')
        self.NoteArrival(tc.contentPane(title='Arrival Note',datapath='.record'))
        tc.contentPane(title='!![en]Vessel details').templateChunk(table='shipsteps.arrival', record_id='^#FORM.record.id',
                                                template='dettaglio_imb')

    #arrival times e details su tabContainer separati e con stackTableHandler disabiltati e utilizzato la form unica
    #def arrival_times(self, pane):
    #    pane.stackTableHandler(relation='@time_arr')
    #    
    #def arr_details(self, pane):
    #    pane.stackTableHandler(relation='@arr_details')
    @public_method
    def emailInOutLazyMode(self,pane):
        pane.stackTableHandler(table='email.message',relation='@email_arr',condition="$arrival_id=:cod",condition_cod='=#FORM.record.id',liveUpdate=True,view_store__onBuilt=True, extendedQuery=True)
        #pane.dataController("""{this.form.reload();}""",ref_num='^#FORM.record.reference_num',_if='ref_num')
        
    def carbordoArr(self,bc):
        center = bc.roundedGroup(title='!![en]Cargo on board on arrival', region='center', height = '100%', background='#f2f0e8').div(margin='10px',margin_left='2px')
        fb = center.formbuilder(cols=3, border_spacing='4px')
        fb.field('cargo_onboard', tag='simpleTextArea',height='50px', colspan=3)
        fb.div("""EXTRA CARGO ON BOARD DESCRIPTION:<br>
                  - Se trattasi di merci pericolose precisare la classe IMO<br>
                  - se trattasi di merci secche pericolose indicare per esteso l'esatta denominazione tecnica e la classe di pericolosità<br>
                  - se trattasi di altre merci secche, precisare se alla rinfusa e l'appendice di appartenenza (A-B-C) qualora soggette al D.M. 22.07.1991 + IMSBC CODE<br>
                  - se trattasi di merce rientrante nelle categorie inquinanti di cui alla Legge 979/1982 specificare tutti I dati relativi al proprietario""", colspan=3)
        fb.field('extra_cargo_onboard', tag='simpleTextArea',height='25px',colspan=3)
    def carbordoDep(self,bc):
        center = bc.roundedGroup(title='!![en]Cargo on board on departure', region='center', height = '100%',background='aliceblue').div(margin='10px',margin_left='2px')
        fb = center.formbuilder(cols=3, border_spacing='4px')
        fb.field('cargo_onboard_dep', tag='simpleTextArea',height='50px', colspan=3)
        fb.div("""EXTRA CARGO ON BOARD DESCRIPTION:<br>
                  - Se trattasi di merci pericolose precisare la classe IMO<br>
                  - se trattasi di merci secche pericolose indicare per esteso l'esatta denominazione tecnica e la classe di pericolosità<br>
                  - se trattasi di altre merci secche, precisare se alla rinfusa e l'appendice di appartenenza (A-B-C) qualora soggette al D.M. 22.07.1991 + IMSBC CODE<br>
                  - se trattasi di merce rientrante nelle categorie inquinanti di cui alla Legge 979/1982 specificare tutti I dati relativi al proprietario""", colspan=3)
        fb.field('extra_cargo_onboard_dep', tag='simpleTextArea',height='25px',colspan=3)
    
    def extraDatiCP(self, bc):
        rg_extra = bc.roundedGroup(title='!![en]Extra data CP on Arrival/Departure',table='shipsteps.extradaticp', region='center',datapath='.record.@extradatacp',width='auto', height = 'auto').div(margin='10px',margin_left='2px')
        div_arrival=rg_extra.div('&emsp;&emsp;&emsp;&emsp;&emsp;&emsp;<strong>ARRIVAL</strong>',width='99%',height='20%',margin='auto',
                        padding='2px',
                        border='1px solid silver',
                        margin_top='1px',margin_left='4px')
        fb = div_arrival.formbuilder(cols=5, border_spacing='4px',fld_width='10em')
        fb.radioButtonText(value='^.motivo_viaggio', values='op_com:Operazioni_Commerciali,altro:Altro', lbl='Motivo Viaggio: ',validate_notnull=True,width='20em',lbl_class='align_right')
        fb.radioButtonText(value='^.tipo_viaggio', values='linea:Di_Linea,occ:Occasionale', lbl='Tipo Viaggio: ',validate_notnull=True, width='20em')
        fb.field('mot_appr', width='20em', colspan=2)
        fb.br()
        fb.field('lavori', width='40em', colspan=2)
        fb.br()
        fb.field('notizie', width='40em', colspan=2)
        fb.br()
        fb.field('pilot_arr')
        fb.field('pilot_arr_vhf')
        fb.field('antifire_arr')
        fb.field('antipol_arr')
        fb.br()
        fb.field('moor_arr')
        fb.field('n_moor_arr')
        fb.field('tug_arr')
        fb.field('n_tug_arr')
        fb.br()
        fb.field('daywork', placeholder='eg. 3 giorni')
        fb.field('timework')
        fb.br()
        fb.field('naz_cte')
        fb.field('n_ita')
        fb.field('n_com')
        fb.field('n_excom')
        fb.br()
        fb.field('carmcycle_arr')
        fb.field('commveic_arr')

        div_departure=rg_extra.div('&emsp;&emsp;&emsp;&emsp;&emsp;&emsp;<strong>DEPARTURE</strong>',width='99%',height='20%',margin='auto',
                        padding='2px',
                        border='1px solid silver',
                        margin_top='1px',margin_left='4px')
        fb = div_departure.formbuilder(cols=5, border_spacing='4px',fld_width='10em')
        fb.field('pilot_dep')
        fb.field('pilot_dep_vhf')
        fb.field('antifire_dep')
        fb.field('antipol_dep')
        fb.br()
        fb.field('moor_dep')
        fb.field('n_moor_dep')
        fb.field('tug_dep')
        fb.field('n_tug_dep')
        fb.br()
        fb.field('pax_dep')
        fb.field('pax_trans')
        fb.field('pax_sba')
        fb.field('pax_imb')
        fb.br()
        fb.field('carmcycle_dep')
        fb.field('carmcycle_trans')
        fb.field('carmcycle_sba')
        fb.field('carmcycle_imb')
        fb.br()
        fb.field('commveic_dep')
        fb.field('commveic_trans')
        fb.field('commveic_sba')
        fb.field('commveic_imb')
        fb.br()
        fb.field('car_imb',colspan=2, tag='textArea', width='36em')
        fb.field('car_sba',colspan=2, tag='textArea', width='36em')

    @public_method
    def servicesLazyMode(self,pane):
        pane.inlineTableHandler(relation='@vess_services',viewResource='ViewFromVesselServices',
                                view_store__onBuilt=True)

    #def services(self,pane):
    #    pane.inlineTableHandler(relation='@vess_services',viewResource='ViewFromVesselServices')

    def datiArrivo(self,bc):
        #center = bc.roundedGroup(title='!![en]Vessel arrival', region='center',datapath='.record',width='210px', height = '100%').div(margin='10px',margin_left='2px')
        #center1 = bc.roundedGroup(title='!![en]Arrival details',region='center',datapath='.record',width='960px', height = '100%', margin_left='210px').div(margin='10px',margin_left='2px')
        #center2 = bc.roundedGroup(title='!![en]Special security guards',table='shipsteps.gpg',region='center',datapath='.record.@gpg_arr',width='240px', height = '150px', margin_left='1170px').div(margin='10px',margin_left='2px')
        #center3 = bc.roundedGroup(title='!![en]EXTRA',region='center',datapath='.record',width='240px',margin_left='1170px', margin_top='150px').div(margin='10px',margin_left='2px')

        center = bc.roundedGroup(title='!![en]Vessel arrival', region='left',datapath='.record',width='210px', height = '100%',splitter=True).div(margin='10px',margin_left='2px')
        center1 = bc.roundedGroup(title='!![en]Arrival details',region='center',datapath='.record', height = '100%', margin_left='0px',splitter=True).div(margin='10px',margin_left='2px')
        
        center2 = bc.roundedGroup(title='!![en]EXTRA',region='right',width='440px', height = '100%', margin_left='0px',splitter=True).div(margin='10px',margin_left='2px')
       # center3 = bc.roundedGroup(title='!![en]EXTRA',region='right',datapath='.record',width='240px',margin_left='0px', margin_top='150px').div(margin='10px',margin_left='2px')
        #center3 = bc.roundedGroup(title='!![en]Times',table='shipsteps.arrival_time',region='center',datapath='.record.@time_arr',width='245px', height = '350px', margin_left='1385px').div(margin='10px',margin_left='2px')
        fb = center.formbuilder(cols=1, border_spacing='4px',lblpos='T')
        #onDbChanges in caso di modifica dati su vessel_details il form arrival viene aggiornato
        fb.onDbChanges("""if(dbChanges.some(change=>change.dbevent=='U' && change.pkey==pkey)){this.form.reload()}""",
            table='shipsteps.vessel_details',pkey='=#FORM.record.vessel_details_id')
        fb.field('agency_id', readOnly=True )
        fb.field('reference_num', readOnly='^gnr.app_preference.shipsteps.ref_num')
        fb.field('date')
        fb.field('vessel_details_id',validate_notnull=True) 
        fb.field('pfda_id' , hasDownArrow=True,  auxColumns='$data,@imbarcazione_id.nome',order_by='$data DESC')
        fb.field('visit_id')
        fb.field('tip_mov' , hasDownArrow=True,  auxColumns='$description',order_by='$description')

        fb = center1.formbuilder(cols=5, border_spacing='4px',lblpos='T',fldalign='left')
        fb.field('eta' , width='10em')
        fb.field('etb' , width='10em')
        fb.field('et_start' , width='10em')
        fb.field('etc' , width='10em')
        fb.field('ets', width='10em' )
        fb.field('draft_aft_arr', width='5em', placeholder='eg:4 or 4,5')
        fb.field('draft_fw_arr' , width='5em', placeholder='eg:4 or 4,5')
        fb.field('draft_aft_dep' , width='5em', placeholder='eg:4 or 4,5')
        fb.field('draft_fw_dep' , width='5em', placeholder='eg:4 or 4,5')
        fb.field('dock_id' )
        fb.field('info_moor',width='30em', colspan=2 ,placeholder='e.g. Inizio ormeggio il ... ore ....')
        fb.field('master_name' )
        fb.field('n_crew' , width='5em',validate_regex=" ^[0-9]*$",validate_regex_error='Insert only numbers')
        fb.field('n_passengers' , width='5em',validate_regex=" ^[0-9]*$",validate_regex_error='Insert only numbers')
        fb.br()
        fb.field('last_port',columns='$descrizione,$unlocode',auxColumns='@nazione_code.nome,$unlocode', limit=20 )
        fb.field('departure_lp' , width='10em')
        fb.field('next_port',columns='$descrizione,$unlocode',auxColumns='@nazione_code.nome,$unlocode', limit=20 )
        fb.field('eta_np' , width='10em')
        fb.field('voy_n', width='10')
        fb.field('mandatory', colspan=3 , width='47em')
        fb.field('cargo_dest', colspan=2, width='29em' )
        fb.br()
        fb.field('invoice_det_id',colspan=5 ,width='78em', hasDownArrow=True)
        fb = center2.formbuilder(cols=1, border_spacing='4px',table='shipsteps.gpg',datapath='.record.@gpg_arr', fld_width='10em',hidden="^#FORM.record.@tip_mov.code?=#v!='pass'")
        #con attributo hidden che punta a tip_mov se diverso dal valore pass nascondiamo il formbuilder della gpg
        #fb.field('arrival_id')
        fb.field('date_start')
        fb.field('date_end')
        fb.field('n_gpg')

        fb = center2.formbuilder(cols=1, datapath='.record',border_spacing='4px', fld_width='18em',lblpos='T',hidden="^#FORM.record.@last_port.@nazione_code.ue_san?=#v==true")
        #nascondiamo il campo nsis in base al valore della pyColumn ue_san nella tabella Nazione pkg Unlocode
        fb.field('nsis_prot')
      # fb = center3.formbuilder(cols=1, border_spacing='4px', fld_width='18em',lblpos='T')
        fb = center2.formbuilder(cols=1, datapath='.record',border_spacing='4px', fld_width='18em',lblpos='T')
        fb.field('firma_div', tag='textArea')

    def datiCaricoBordo(self,bc):
        center = bc.roundedGroup(title='!![en]Cargo on board', region='center', height = '100%').div(margin='10px',margin_left='2px')
        fb = center.formbuilder(cols=3, border_spacing='4px')
        fb.field('cargo_onboard', tag='simpleTextArea',height='50px', colspan=3)
        fb.div("""EXTRA CARGO ON BOARD DESCRIPTION:<br>
                  - Se trattasi di merci pericolose precisare la classe IMO<br>
                  - se trattasi di merci secche pericolose indicare per esteso l'esatta denominazione tecnica e la classe di pericolosità<br>
                  - se trattasi di altre merci secche, precisare se alla rinfusa e l'appendice di appartenenza (A-B-C) qualora soggette al D.M. 22.07.1991 + IMSBC CODE<br>
                  - se trattasi di merce rientrante nelle categorie inquinanti di cui alla Legge 979/1982 specificare tutti I dati relativi al proprietario""", colspan=3)
        fb.field('extra_cargo_onboard', tag='simpleTextArea',height='25px',colspan=3)

    def datiCarico(self,pane):
        pane.inlineTableHandler(relation='@cargo_lu_arr',viewResource='ViewFromCargoLU')

    def datiCaricoTransit(self,bc):
        center = bc.roundedGroup(title='!![en]Transit cargo', region='center', height = '100%').div(margin='10px',margin_left='2px')
        fb = center.formbuilder(cols=3, border_spacing='4px')
        fb.field('transit_cargo', tag='simpleTextArea',height='50px', colspan=3)
        fb.div("""EXTRA TRANSIT CARGO DESCRIPTION:<br>
                  - Se trattasi di merci pericolose precisare la classe IMO<br>
                  - se trattasi di merci secche pericolose indicare per esteso l'esatta denominazione tecnica e la classe di pericolosità<br>
                  - se trattasi di altre merci secche, precisare se alla rinfusa e l'appendice di appartenenza (A-B-C) qualora soggette al D.M. 22.07.1991 + IMSBC CODE<br>
                  - se trattasi di merce rientrante nelle categorie inquinanti di cui alla Legge 979/1982 specificare tutti I dati relativi al proprietario""", colspan=3)
        fb.field('extra_transit_cargo', tag='simpleTextArea',height='25px',colspan=3)
   
    @public_method
    def car_ricLazyMode(self,pane):
         pane.stackTableHandler(table='shipsteps.ship_rec', formResource='Form',view_store_onStart=True,view_store__onBuilt=True)

    #def car_ric(self,pane):
    #     pane.stackTableHandler(table='shipsteps.ship_rec', formResource='Form',view_store_onStart=True)

        
    @public_method
    def charterersLazyMode(self,pane):
        pane.inlineTableHandler(table='shipsteps.charterers',datapath='#FORM.charterers',parentForm=False,saveButton=True,semaphore=True,viewResource='ViewFromCharterers',view_store_onStart=True,view_store__onBuilt=True)
        
    #def charterers(self,pane):
    #    pane.inlineTableHandler(table='shipsteps.charterers',viewResource='ViewFromCharterers',view_store_onStart=True)

    @public_method
    def sofLazyMode(self,pane):
        pane.stackTableHandler(relation='@sof_arr',view_store__onBuilt=True,liveUpdate=True)

    #def sof(self,pane):
    #    pane.stackTableHandler(relation='@sof_arr')#, formResource='FormSof')

    def arrival_details(self, bc):
        rg_times = bc.roundedGroup(title='!![en]Arrival/Departure times',table='shipsteps.arrival_time',region='left',datapath='.record.@time_arr',width='350px', height = 'auto').div(margin='10px',margin_left='2px')
        rg_details = bc.roundedGroup(title='!![en]Arrival details',table='shipsteps.arrival_det', region='center',datapath='.record.@arr_details',width='350px', height = '100%',margin_left='350px').div(margin='10px',margin_left='2px')
        rg_details_dep = bc.roundedGroup(title='!![en]Departure details',table='shipsteps.arrival_det', region='center',datapath='.record.@arr_details',width='350px', height = '100%',margin_left='350px').div(margin='10px',margin_left='2px')
        #rg_extra = bc.roundedGroup(title='!![en]Extra data CP on Arrival/Departure',table='shipsteps.extradaticp', region='center',datapath='.record.@extradatacp',width='auto', height = 'auto', margin_left='550px').div(margin='10px',margin_left='2px')
        fb = rg_times.formbuilder(cols=1, border_spacing='4px',fld_width='10em')
        fb.field('eosp')
        fb.field('aor')
        fb.field('anchored')
        fb.field('anchor_up')
        fb.field('pob')
        fb.field('first_rope')
        fb.field('moored')
        fb.field('poff')
        fb.field('gangway')
        fb.field('free_p')
        fb.field('pobd',border_color="^pildep") #tramite il datacontroller in th_sof viene assegnata alla variabile pildep il colore del bordo
        fb.field('last_line')
        fb.field('sailed',border_color="^sail") #tramite il datacontroller in th_sof viene assegnata alla variabile sail il colore del bordo
        fb.field('cosp', lbl='Commenced of <br>Sea Passage',fldvalign='center')

        btn_arrivo=fb.button('Email arrival',hidden="^checksof")#.controller.title?=#v!=null")
        btn_partenza=fb.button('Email departure',hidden="^checksof")#.controller.title?=#v!=null")
        fb.dataRpc('checksof', self.checkSof,  record='=#FORM.record', cur_tab='^#FORM.current_tab',titolo='^#FORM.controller.title',
                   _if='cur_tab!=null||titolo!=null')
        #con il datacontroller all'inserimento dei dati in pobd e sailed risettiamo le variabili pildep e sail in null per togliere il colore rosso del bordo
        fb.dataController("""if(pobd!=null){SET pildep=null;} if(sailed!=null){SET sail=null;}""",pobd='^.pobd',sailed='^.sailed')

        btn_arrivo.dataRpc('nome_temp', self.email_arrdep,record='=#FORM.record',servizio=['arr'], email_template_id='email_arrivo',
                            nome_template = 'shipsteps.arrival:email_arrivo',format_page='A4',
                            _ask=dict(title='!![en]Select the Attachments',fields=[dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                             table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM/parent/#FORM.record.id',
                             cols=4,popup=True,colspan=2)]))
        btn_partenza.dataRpc('nome_temp', self.email_arrdep,record='=#FORM.record',servizio=['arr'], email_template_id='email_departure',
                            nome_template = 'shipsteps.arrival:email_departure',format_page='A4',
                            _ask=dict(title='!![en]Select the Attachments',fields=[dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                             table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM/parent/#FORM.record.id',
                             cols=4,popup=True,colspan=2)]))
        
        div_draft=rg_details.div('&emsp;&emsp;&emsp;&emsp;&emsp;&emsp;<strong>DRAFT</strong>',width='99%',height='20%',margin='auto',
                        padding='2px',
                        border='1px solid silver',
                        margin_top='1px',margin_left='4px')
        fb = div_draft.formbuilder(cols=1, border_spacing='4px',fld_width='10em')
        fb.field('draft_aft_arr',placeholder='e.g. 4,5')
        fb.field('draft_fw_arr',placeholder='e.g. 4,5')
   
        div_draft=rg_details_dep.div('&emsp;&emsp;&emsp;&emsp;&emsp;&emsp;<strong>DRAFT</strong>',width='99%',height='20%',margin='auto',
                        padding='2px',
                        border='1px solid silver',
                        margin_top='1px',margin_left='4px')
        fb = div_draft.formbuilder(cols=1, border_spacing='4px',fld_width='10em')
        fb.field('draft_aft_dep',placeholder='e.g. 4,5')
        fb.field('draft_fw_dep',placeholder='e.g. 4,5')

        div_rem=rg_details.div('&emsp;&emsp;&emsp;&emsp;&emsp;&emsp;<strong>REMAINS</strong>',width='99%',height='20%',margin='auto',
                        padding='2px',
                        border='1px solid silver',
                        margin_top='1px',margin_left='4px')
        fb = div_rem.formbuilder(cols=1, border_spacing='4px',fld_width='10em')
        fb.field('ifo_arr',placeholder='e.g. mt.50')
        fb.field('do_arr',placeholder='e.g. mt.50')
        fb.field('lo_arr',placeholder='e.g. kgs.50')
       
        div_rem=rg_details_dep.div('&emsp;&emsp;&emsp;&emsp;&emsp;&emsp;<strong>REMAINS</strong>',width='99%',height='20%',margin='auto',
                        padding='2px',
                        border='1px solid silver',
                        margin_top='1px',margin_left='4px')
        fb = div_rem.formbuilder(cols=1, border_spacing='4px',fld_width='10em')
        fb.field('ifo_dep',placeholder='e.g. mt.50')
        fb.field('do_dep',placeholder='e.g. mt.50')
        fb.field('lo_dep',placeholder='e.g. kgs.50')

        div_fw=rg_details.div('&emsp;&emsp;&emsp;&emsp;&emsp;&emsp;<strong>FRESH WATER</strong>',width='99%',height='20%',margin='auto',
                        padding='2px',
                        border='1px solid silver',
                        margin_top='1px',margin_left='4px')
        fb = div_fw.formbuilder(cols=1, border_spacing='4px',fld_width='10em')
        fb.field('fw_arr',placeholder='e.g. mt.50')

        div_fw=rg_details_dep.div('&emsp;&emsp;&emsp;&emsp;&emsp;&emsp;<strong>FRESH WATER</strong>',width='99%',height='20%',margin='auto',
                        padding='2px',
                        border='1px solid silver',
                        margin_top='1px',margin_left='4px')
        fb = div_fw.formbuilder(cols=1, border_spacing='4px',fld_width='10em')
        fb.field('fw_dep',placeholder='e.g. mt.50')

        div_tug=rg_details.div('&emsp;&emsp;&emsp;&emsp;&emsp;&emsp;<strong>USED TUGS</strong>',width='99%',height='20%',margin='auto',
                        padding='2px',
                        border='1px solid silver',
                        margin_top='1px',margin_left='4px')
        fb = div_tug.formbuilder(cols=1, border_spacing='4px',fld_width='10em')
        fb.field('tug_in',placeholder='e.g. 1')
     
        div_tug=rg_details_dep.div('&emsp;&emsp;&emsp;&emsp;&emsp;&emsp;<strong>USED TUGS</strong>',width='99%',height='20%',margin='auto',
                        padding='2px',
                        border='1px solid silver',
                        margin_top='1px',margin_left='4px')
        fb = div_tug.formbuilder(cols=1, border_spacing='4px',fld_width='10em')
        fb.field('tug_out',placeholder='e.g. 1')

    @public_method
    def checkSof(self, record, **kwargs):
        tabname=kwargs['cur_tab']
        record_id = record['id']
        tbl_sof = self.db.table('shipsteps.sof')
        sof = tbl_sof.query(columns='$id',where = '$arrival_id = :arr_id', arr_id=record_id).fetch()
        if sof:
            return True

    @public_method
    def cargodocsCertLazyMode(self,pane):
        pane.stackTableHandler(relation='@cargodocs_arr',formResource='FormFromCargodocs',view_store__onBuilt=True)

    @public_method
    def usmaCertLazyMode(self,pane):
        #pane.inlineTableHandler(title='!![en]Renewal/Issue certificates',relation='@certusma_arr',viewResource='ViewFromCertusma',view_store__onBuilt=True)
        pane.stackTableHandler(relation='@certusma_arr',formResource='FormFromCertusma',view_store__onBuilt=True)
    
    def usmaCert2(self,bc_usma):
        rg_certusma = bc_usma.roundedGroup(title='!![en]Renewal/Issue certificates',table='shipsteps.certsanimare',region='center',datapath='.record.@certusma_arr',width='100%', height = '100%').div(margin='10px',margin_left='2px')
        fb = rg_certusma.formbuilder(cols=4, border_spacing='4px',fld_width='10em')
        fb.field('xconto',width='30em', placeholder='Es. per conto comando nave M/V...', colspan=4)  
        fb.field('docagent', placeholder="Es. Carta d'Identità")
        fb.field('doc_n')
        fb.field('issuedby', width='15em')
        fb.field('datedoc')
        fb.field('navigation', width='30em', colspan=4)
        div1=rg_certusma.div(width='99%',height='20%',margin='auto',
                        padding='2px',
                        border='1px solid silver',
                        margin_top='1px',margin_left='4px')
        fb1=div1.formbuilder(colspan=3,cols=9, border_spacing='1px')
        fb1.field('sanification',lbl='')
        fb1.div('SANIFICAZIONE / ESENZIONE SANIFICAZIONE', width='30em', colspan=2)
        fb1.field('san_nsis', lbl='!![en]NSIS Code')
        fb1.br()
        fb1.field('medicines',lbl='')
        fb1.div('CASSETTA MEDICINALE / FARMACIA DI BORDO (TAB)', width='30em')
        fb1.field('tab_medicine', width='6em')
        fb1.field('med_nsis', lbl='!![en]NSIS Code')
        fb1.br()
        fb1.field('waterbox',lbl='')
        fb1.div('CASSE ACQUA POTABILE', width='30em', colspan=2)
        fb1.field('water_nsis', lbl='!![en]NSIS Code')
        fb1.br()
        fb1.field('reg_narcotics',lbl='')
        fb1.div('VIDIMAZIONE REGISTRO STUPEFACENTI', width='30em', colspan=2)
        fb1.field('narcotics_nsis', lbl='!![en]NSIS Code')
        fb1.br()
        fb1.field('newreg_narcotics',lbl='')
        fb1.div('RILASCIO E/O CHIUSURA REGISTRO STUPEFACENTI', width='30em', colspan=2)
        fb1.field('newnarcotics_nsis', lbl='!![en]NSIS Code')
        fb1.br()
        fb1.field('destroy_med',lbl='')
        fb1.div('DISTRUZIONE FARMACI STUPEFACENTI SCADUTI', width='30em', colspan=2)
        fb1.field('destroymed_nsis', lbl='!![en]NSIS Code')
        fb2=rg_certusma.formbuilder(cols=9,colspan=9, border_spacing='4px',fld_width='10em')
        fb2.field('visit_date')
        fb2.field('date')
    
    @public_method
    def emailArrivalLazyMode(self,pane):
        pane.inlineTableHandler(title='!![en]Email arrival',relation='@arrival_email',viewResource='ViewFromEmailArrival',view_store__onBuilt=True)

    def NoteArrival(self,frame):
        frame.simpleTextArea(title='Arrival note',value='^.note',editor=True)

    #def emailArrival(self,pane):
    #    pane.inlineTableHandler(title='!![en]Email arrival',relation='@arrival_email',viewResource='ViewFromEmailArrival')

   #def sof_cargo(self,pane):
   #    pane.inlineTableHandler(table='shipsteps.sof_cargo', viewResource='ViewFromSof_Cargo')
    def times(self,frame):
        center = frame.roundedGroup(title='!![en]Expected Times',font_weight='bold', region='center',datapath='.record').div(margin='10px',margin_left='2px')
        #rg_arr = frame.roundedGroup(title='!![en]Arrival',datapath='.record',width='100%', height = '10%').div(margin='10px',margin_left='2px')
        fb = center.formbuilder(cols=6, border_spacing='4px',fld_width='10em',lbl_color='white',lbl_font_weight='bold', font_weight='bold')
        fb.field('eta' , width='10em')
        fb.field('etb' , width='10em')
        fb.field('et_start' , width='10em')
        fb.field('etc' , width='10em')
        fb.field('ets', width='10em' )
        fb.field('dock_id', width='15em' )

    def times_sof(self,frame):
        self.times(frame) #per non riscrivere lo stesso codice di times passiamo direttamente self.times(frame)
    
    def taskList(self, bc_tasklist):
        rg_prearrival = bc_tasklist.roundedGroup(title='!![en]<strong>Pre arrival</strong>',table='shipsteps.tasklist',region='left',datapath='.record.@arr_tasklist',width='220px', height = '100%').div(margin='10px',margin_left='2px')
        rg_prearrival2 = bc_tasklist.roundedGroup(title='!![en]<strong>Pre arrival - Email</strong>',table='shipsteps.tasklist',region='left',datapath='.record.@arr_tasklist',width='220px', height = '100%', margin_left='220px').div(margin='10px',margin_left='2px')
        rg_arrival = bc_tasklist.roundedGroup(title='!![en]<strong>Arrival/Departure</strong>',table='shipsteps.tasklist',region='center',datapath='.record.@arr_tasklist',width='240px', height = '100%', margin_left='440px').div(margin='10px',margin_left='2px')
        rg_arrival_nsw = bc_tasklist.roundedGroup(title='!![en]<strong>Arrival/Departure NSW</strong>',table='shipsteps.tasklist',region='center',datapath='.record.@arr_tasklist',width='240px', height = '100%', margin_left='680px').div(margin='10px',margin_left='2px')
        rg_extra = bc_tasklist.roundedGroup(title='!![en]<strong>Extra</strong>',table='shipsteps.tasklist',region='center',datapath='.record.@arr_tasklist',width='220px', height = '100%', margin_left='480px').div(margin='10px',margin_left='2px')
        
        tbl_email_services = self.db.table('shipsteps.email_services')

        #definizione primo rettangolo di stampa all'interno del roundedGroup Pre Arrival
        div1=rg_prearrival.div(width='99%',height='20%',margin='auto',
                        padding='2px',
                        border='1px solid silver',
                        margin_top='1px',margin_left='4px')
        fb1=div1.formbuilder(colspan=1,cols=3, border_spacing='1px',fld_width='150px')

        btn_cl = fb1.Button('!![en]Print Check list')
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.checklist',button=btn_cl.js_widget)
        btn_cl.dataRpc('nome_temp', self.print_template,record='=#FORM.record',nome_template='shipsteps.arrival:check_list',
                            nome_vs='=#FORM.record.@vessel_details_id.@imbarcazione_id.nome',
                            format_page='A4',_onResult="this.form.save();")
      #  fb1.dataController("if(msg=='check_list') SET .checklist=true", msg='^nome_temp')
        fb1.field('checklist', lbl='', margin_top='5px')
        fb1.semaphore('^.checklist?=#v==true?true:false', margin_top='5px')
        btn_fs = fb1.Button('!![en]Print Frontespicie')
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.frontespizio',button=btn_fs.js_widget)
        btn_fs.dataRpc('nome_temp', self.print_template,record='=#FORM.record', nome_vs='=#FORM.record.@vessel_details_id.@imbarcazione_id.nome',
                            nome_template = 'shipsteps.arrival:front_nave',format_page='A4',_onResult="this.form.save();")
       # fb1.dataController("if(msg=='front_nave') SET .frontespizio=true", msg='^nome_temp')
        fb1.field('frontespizio', lbl='', margin_top='5px')
        fb1.semaphore('^.frontespizio?=#v==true?true:false', margin_top='5px')

        btn_mn = fb1.Button('!![en]Print Vessel module')
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.modulo_nave',button=btn_mn.js_widget)
        btn_mn.dataRpc('nome_temp', self.print_template,record='=#FORM.record', nome_vs='=#FORM.record.@vessel_details_id.@imbarcazione_id.nome',
                            nome_template = 'shipsteps.arrival:mod_nave',format_page='A4',_onResult="this.form.save();")
       # fb1.dataController("if(msg=='mod_nave') SET .checklist=true", msg='^nome_temp')
        fb1.field('modulo_nave', lbl='', margin_top='5px')
        fb1.semaphore('^.modulo_nave?=#v==true?true:false', margin_top='5px')

        btn_cn = fb1.Button('!![en]Print Vessel folder')
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.cartella_nave',button=btn_cn.js_widget)
       #fb1.dataController("""var id = button.id; console.log(id);
       #                if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
       #                else if (ca==false){document.getElementById(id).style.backgroundColor = 'salmon';}
       #                else {document.getElementById(id).style.backgroundColor = 'lightgrey';}
       #                """, ca='^.cartella_nave',button=btn_cn.js_widget)                
       #btn_cn.dataRpc(None, self.print_template,record='=#FORM.record.id',_ask=dict(title='!![en]Choose lettehead to use with this form',
       #                                        fields=[dict(name='letterhead_id', lbl='!![en]Letterhead', tag='dbSelect',columns='$id',
       #                                        hasDownArrow=True, auxColumns='$name', table='adm.htmltemplate')]), nome_vs='=#FORM.record.@vessel_details_id.@imbarcazione_id.nome',
       #                    nome_template = 'shipsteps.arrival:cartella_doc')
        btn_cn.dataRpc('nome_temp', self.print_template, record='=#FORM.record', nome_vs='=#FORM.record.@vessel_details_id.@imbarcazione_id.nome',
                            nome_template = 'shipsteps.arrival:cartella_doc',format_page='A3',_onResult="this.form.save();")
       # fb1.dataController("if(msg=='cartella_doc') SET .cartella_nave=true", msg='^nome_temp')
        fb1.field('cartella_nave', lbl='', margin_top='5px')
        fb1.semaphore('^.cartella_nave?=#v==true?true:false', margin_top='5px')

        btn_ts = fb1.Button('!![en]Print Servicies table')
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.tab_servizi',button=btn_ts.js_widget)
        btn_ts.dataRpc('nome_temp', self.print_template,record='=#FORM.record', nome_vs='=#FORM.record.@vessel_details_id.@imbarcazione_id.nome',
                            nome_template = 'shipsteps.arrival:tab_servizi',format_page='A3',_onResult="this.form.save();")
       # fb1.dataController("if(msg=='tab_servizi') SET .tab_servizi=true", msg='^nome_temp')
        fb1.field('tab_servizi', lbl='', margin_top='5px')
        fb1.semaphore('^.tab_servizi?=#v==true?true:false', margin_top='5px')

        btn_fc = fb1.Button('!![en]Print Cargo frontespiece',hidden="^#FORM.record.@tip_mov.code?=#v=='pass'")#attributo hidden nascondiamo il widget se il valore tip_mov = pass
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.front_carico',button=btn_fc.js_widget)
        btn_fc.dataRpc('nome_temp', self.print_template,record='=#FORM.record', nome_vs='=#FORM.record.@vessel_details_id.@imbarcazione_id.nome',
                            nome_template = 'shipsteps.arrival:front_carico',format_page='A4',_onResult="this.form.save();")
       # fb1.dataController("if(msg=='front_carico') SET .front_carico=true", msg='^nome_temp')
        fb1.field('front_carico', lbl='', margin_top='5px',hidden="^#FORM.record.@tip_mov.code?=#v=='pass'")#attributo hidden nascondiamo il widget se il valore tip_mov = pass
        fb1.semaphore('^.front_carico?=#v==true?true:false', margin_top='5px',hidden="^#FORM.record.@tip_mov.code?=#v=='pass'")#attributo hidden nascondiamo il widget se il valore tip_mov = pass

        btn_trib = fb1.Button('!![en]Tributes', action="""{SET tabname='tributi';}""")
       #fb1.dataController("""var id = button.id; console.log(id);
       #                if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
       #                else {document.getElementById(id).style.backgroundColor = '';}
       #                """, ca='^.tributi_cp',button=btn_trib.js_widget)                     
       #btn_trib.dataRpc('nome_temp', self.print_template_tributi,record='=#FORM.record',servizio=['tributi'], 
       #                        format_page='A4',selId='=#FORM.shipsteps_tributi_cp.view.grid.selectedId',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',                            
       #                        _ask=dict(title='!![en]Select the type of payment',fields=[dict(name='tip_versamento', lbl='!![en]Type of payment', tag='filteringSelect',hasDownArrow=True,
       #                        values='bonifico:Bonifico,bollettino:Bollettino postale',
       #                        validate_notnull=True,cols=4,popup=True,colspan=2)]),_onResult="this.form.save()")
       #fb1.field('tributi_cp', lbl='', margin_top='5px')
       #fb1.semaphore('^.tributi_cp?=#v==true?true:false', margin_top='5px')
        #definizione secondo rettangolo invio email all'interno del roundedGroup Pre Arrival
        div2=rg_prearrival2.div('<center><strong></strong>',width='99%',height='20%',margin='auto',
                        padding='2px',
                        border='1px solid silver',
                        margin_top='1px',margin_left='4px')
        fb = div2.formbuilder(colspan=1,cols=3, border_spacing='1px', colswidth='4px')
        #fb = rg_prearrival.formbuilder(colspan=3,cols=6, border_spacing='4px',colswidth='10px')
       
        
        
        btn_sr = fb.Button('!![en]Shipper/Receivers', width='10em')
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.email_ship_rec',button=btn_sr.js_widget)
        btn_sr.dataRpc('nome_temp', self.email_arrival_sof,
                   record='=#FORM.record', servizio=['arr','sof'], email_template_id='email_arr_shiprec',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                   _ask=dict(title='!![en]Select the SOF and Attachments',fields=[dict(name='sof_id', lbl='!![en]sof', tag='dbSelect',columns='$id',
                             hasDownArrow=True, auxColumns='$sof_n,$ship_rec', table='shipsteps.sof',condition="$arrival_id =:cod",
                                                condition_cod='=#FORM.record.id',width='25em',validate_notnull=True),dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                             table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                             cols=4,popup=True,colspan=2)]),_onResult="this.form.save();")
        fb.field('email_ship_rec',lbl='', margin_top='5px')
        fb.semaphore('^.email_ship_rec?=#v==true?true:false', margin_top='5px')
        #datacontroller verifica il valore della variabile nome_temp di ritorno dalla funzione per invio email
        #e setta il valore della campo checkbox a true e lancia il messaggio 'Messaggio Creato'
      #  fb.dataController("if(msgspec=='ship_rec') {SET .email_ship_rec=true ; alert('Message created')} if(msgspec=='no_email') alert('You must insert destination email as TO or BCC'); if(msgspec=='no_sof') alert('You must select the SOF or you must create new one');", msgspec='^msg_special')
        
        #verifichiamo quanti servizi CP ci sono, nel caso più di uno apparirà la dbSelect per la scelta
        service_for_email = tbl_email_services.query(columns="$service_for_email_id", where='$service_for_email_id=:serv', serv='dog').fetch()
        serv_len=len(service_for_email)
        btn_dog = fb.Button('!![en]Customs/GdF', width='10em')
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.email_dogana',button=btn_dog.js_widget)
        if serv_len > 1:
            btn_dog.dataRpc('nome_temp', self.email_services,
                       record='=#FORM.record', servizio=['dogana','gdf','gdf roan'], email_template_id='email_dogana',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                        _ask=dict(title='!![en]Select the services',fields=[dict(name='services', lbl='!![en]Services', tag='dbSelect',hasDownArrow=True,
                                table='shipsteps.email_services', columns='$consignee',condition="$service_for_email_id=:cod",condition_cod='dog',alternatePkey='consignee',
                                validate_notnull=True,cols=4,popup=True,colspan=2, hasArrowDown=True),dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                                table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                                cols=4,popup=True,colspan=2)]),_onResult="this.form.save();") 
        else:
            btn_dog.dataRpc('nome_temp', self.email_services,
                       record='=#FORM.record', servizio=['dogana','gdf','gdf roan'], email_template_id='email_dogana',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                        _ask=dict(title='!![en]Select the Attachments',fields=[dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                                 table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                                 cols=4,popup=True,colspan=2)]),_onResult="this.form.save();")
        #datacontroller verifica il valore della variabile nome_temp di ritorno dalla funzione per invio email
        #e setta il valore della campo checkbox a true e lancia il messaggio 'Messaggio Creato'
       # fb.dataController("if(msgspec=='val_dog') {SET .email_dogana=true ; alert('Message created')}", msgspec='^nome_temp')
        fb.field('email_dogana',lbl='', margin_top='5px')
        fb.semaphore('^.email_dogana?=#v==true?true:false', margin_top='5px')

       #uploader = fb.button('Upload more files')
       #uploader.dataController("""
       #                        genro.dlg.multiUploaderDialog('!![en]Upload many files and assign them to users',{
       #                                    uploadPath:uploadPath,
       #                                    onResult:function(){
       #                                        genro.publish("floating_message",{message:"Upload completato", messageType:"message"});
       #                                        genro.publish('trigger_action',{user_id:user_id}); }
       #                                    });""",
       #                                    uploadPath=':import_queue',
       #                                    _ask=dict(title='Choose users to whom to assign files',
       #                                        fields=[dict(name='user_id', lbl='User', tag='dbselect', table='adm.user')]))
        #verifichiamo quanti servizi Immigration ci sono, nel caso più di uno apparirà la dbSelect per la scelta
        service_for_email = tbl_email_services.query(columns="$service_for_email_id", where='$service_for_email_id=:serv', serv='imm').fetch()
        serv_len=len(service_for_email)
        btn_fr = fb.Button('!![en]Immigration', width='10em')
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.email_frontiera',button=btn_fr.js_widget)
        if serv_len > 1:
            btn_fr.dataRpc('nome_temp', self.print_template,
                       record='=#FORM.record', servizio=['immigration'], email_template_id='email_frontiera',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                       nome_template = 'shipsteps.arrival:form_immigration', nome_vs='=#FORM.record.@vessel_details_id.@imbarcazione_id.nome',format_page='A4',
                       _ask=dict(title='!![en]Select the services',fields=[dict(name='services', lbl='!![en]Services', tag='dbSelect',hasDownArrow=True,
                                table='shipsteps.email_services', columns='$consignee',condition="$service_for_email_id=:cod",condition_cod='imm',alternatePkey='consignee',
                                validate_notnull=True,cols=4,popup=True,colspan=2, hasArrowDown=True),dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                                table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                                cols=4,popup=True,colspan=2),dict(name='std_att', lbl='!![en]Service attachments', tag='checkboxtext',hasDownArrow=True,
                                table='shipsteps.email_services_atc', columns='$description', auxColumns='$maintable_id',condition="@maintable_id.service_for_email_id=:cod",condition_cod='immigration',
                                cols=4,popup=True,colspan=2, hasArrowDown=True)]),_onResult="this.form.save();")
        else:
             btn_fr.dataRpc('nome_temp', self.print_template,
                       record='=#FORM.record', servizio=['immigration'], email_template_id='email_frontiera',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                       nome_template = 'shipsteps.arrival:form_immigration', nome_vs='=#FORM.record.@vessel_details_id.@imbarcazione_id.nome',format_page='A4',
                       _ask=dict(title='!![en]Select the Attachments',fields=[dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                                 table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                                 cols=4,popup=True,colspan=2),dict(name='std_att', lbl='!![en]Service attachments', tag='checkboxtext',hasDownArrow=True,
                                table='shipsteps.email_services_atc', columns='$description', auxColumns='$maintable_id',condition="@maintable_id.service_for_email_id=:cod",condition_cod='immigration',
                                cols=4,popup=True,colspan=2, hasArrowDown=True)]),_onResult="this.form.save();")
       # fb.dataController("if(msgspec=='val_imm') {SET .email_frontiera=true; alert('Message created')}", msgspec='^nome_temp')
        fb.field('email_frontiera',lbl='', margin_top='5px')
        fb.semaphore('^.email_frontiera?=#v==true?true:false', margin_top='5px')
        #verifichiamo quanti servizi Sanimare ci sono, nel caso più di uno apparirà la dbSelect per la scelta
        service_for_email = tbl_email_services.query(columns="$service_for_email_id", where='$service_for_email_id=:serv', serv='usma').fetch()
        serv_len=len(service_for_email)
        btn_usma = fb.Button('!![en]Sanimare', width='10em',hidden="^#FORM.record.@last_port.@nazione_code.ue_san?=#v==true")#nascondiamo il widget in base al valore della pyColumn ue_san nella tabella Nazione pkg Unlocode
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.email_usma',button=btn_usma.js_widget)
        if serv_len > 1:
            btn_usma.dataRpc('nome_temp', self.email_services,
                       record='=#FORM.record', servizio=['sanimare'], email_template_id='email_sanimare',nsisprot='=#FORM.record.nsis_prot',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                       _ask=dict(title='!![en]Select the services',fields=[dict(name='services', lbl='!![en]Services', tag='dbSelect',hasDownArrow=True,
                                table='shipsteps.email_services', columns='$consignee',condition="$service_for_email_id=:cod",condition_cod='usma',alternatePkey='consignee',
                                validate_notnull=True,cols=4,popup=True,colspan=2, hasArrowDown=True),dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                                table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                                cols=4,popup=True,colspan=2),dict(name='std_att', lbl='!![en]Service attachments', tag='checkboxtext',hasDownArrow=True,
                                table='shipsteps.email_services_atc', columns='$description', auxColumns='$maintable_id',condition="@maintable_id.service_for_email_id=:cod",condition_cod='sanimare',
                                cols=4,popup=True,colspan=2, hasArrowDown=True)]),_onResult="this.form.save();")
        else:
            btn_usma.dataRpc('nome_temp', self.email_services,
                       record='=#FORM.record', servizio=['sanimare'], email_template_id='email_sanimare',nsisprot='=#FORM.record.nsis_prot',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                       _ask=dict(title='!![en]Select the Attachments',fields=[dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                                 table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                                 cols=4,popup=True,colspan=2),dict(name='std_att', lbl='!![en]Service attachments', tag='checkboxtext',hasDownArrow=True,
                                table='shipsteps.email_services_atc', columns='$description', auxColumns='$maintable_id',condition="@maintable_id.service_for_email_id=:cod",condition_cod='sanimare',
                                cols=4,popup=True,colspan=2, hasArrowDown=True)]),_onResult="this.form.save();")
       # fb.dataController("if(msgspec=='val_usma') {SET .email_usma=true ; alert('Message created')}", msgspec='^nome_temp')
        fb.field('email_usma',lbl='', margin_top='5px',hidden="^#FORM.record.@last_port.@nazione_code.ue_san?=#v==true")#nascondiamo il widget in base al valore della pyColumn ue_san nella tabella Nazione pkg Unlocode
        fb.semaphore('^.email_usma?=#v==true?true:false', margin_top='5px',hidden="^#FORM.record.@last_port.@nazione_code.ue_san?=#v==true")#nascondiamo il widget in base al valore della pyColumn ue_san nella tabella Nazione pkg Unlocode

        btn_pilot = fb.Button('!![en]Pilot/Moor', width='10em')
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.email_pilot_moor',button=btn_pilot.js_widget)
        btn_pilot.dataRpc('nome_temp', self.email_services,
                   record='=#FORM.record', servizio=['pilot','mooringmen'], email_template_id='email_pilot_moor',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                   _ask=dict(title='!![en]Select the Attachments',fields=[dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                             table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                             cols=4,popup=True,colspan=2)]),_onResult="this.form.save();")
       # fb.dataController("if(msgspec=='val_pil_moor') {SET .email_pilot_moor=true ; alert('Message created')}", msgspec='^msg_special')
        fb.field('email_pilot_moor',lbl='', margin_top='5px')
        fb.semaphore('^.email_pilot_moor?=#v==true?true:false', margin_top='5px')

        btn_tug = fb.Button('!![en]Tug', width='10em')
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.email_tug',button=btn_tug.js_widget)
        btn_tug.dataRpc('nome_temp', self.email_services,
                   record='=#FORM.record', servizio=['tug'], email_template_id='email_tug',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                   _ask=dict(title='!![en]Select the Attachments',fields=[dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                             table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                             cols=4,popup=True,colspan=2)]),_onResult="this.form.save();")
       # fb.dataController("if(msgspec=='val_tug') {SET .email_tug=true ; alert('Message created')}", msgspec='^msg_special')
       
        fb.field('email_tug',lbl='', margin_top='5px')
        fb.semaphore('^.email_tug?=#v==true?true:false', margin_top='5px')
        
        #verifichiamo quanti servizi garbage ci sono, nel caso più di uno apparirà la dbSelect per la scelta
        service_for_email = tbl_email_services.query(columns="$service_for_email_id", where='$service_for_email_id=:serv', serv='garb').fetch()
        serv_len=len(service_for_email)
        btn_garb = fb.Button('!![en]Garbage pick-up', width='10em')
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.email_garbage',button=btn_garb.js_widget)
        if serv_len > 1:
            agency=self.db.currentEnv.get('current_agency_id')
                        
            btn_garb.dataRpc('nome_temp', self.print_template_garbage,record='=#FORM.record',servizio=['garbage'], email_template_id='garbage_email',
                                nome_template = 'shipsteps.garbage:garbage_request',format_page='A4',selId='=#FORM.shipsteps_garbage.view.grid.selectedId',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',                            
                                _ask=dict(title='!![en]Select the services',fields=[dict(name='services', lbl='!![en]Services', tag='dbSelect',hasDownArrow=True,
                                table='shipsteps.email_services', columns='$consignee', auxColumns='$email,$email_cc,$email_bcc,$email_pec,$email_cc_pec',condition="$service_for_email_id=:cod",condition_cod='garb',alternatePkey='consignee',
                                validate_notnull=True,cols=4,popup=True,colspan=2, hasArrowDown=True),dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                                table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                                cols=4,popup=True,colspan=2),dict(name='std_att', lbl='!![en]Service attachments', tag='checkboxtext',hasDownArrow=True,
                                table='shipsteps.email_services_atc', columns='$description', auxColumns='$maintable_id',condition="@maintable_id.service_for_email_id=:cod",condition_cod='garbage',
                                cols=4,popup=True,colspan=2, hasArrowDown=True)]),_onResult="this.form.save()") 
        else:
            btn_garb.dataRpc('nome_temp', self.print_template_garbage,record='=#FORM.record.id',servizio=['garbage'], email_template_id='garbage_email',
                                nome_template = 'shipsteps.garbage:garbage_request',format_page='A4',selId='=#FORM.shipsteps_garbage.view.grid.selectedId',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',                            
                                _ask=dict(title='!![en]Select the Attachments',fields=[dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                                table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                                cols=4,popup=True,colspan=2),dict(name='std_att', lbl='!![en]Service attachments', tag='checkboxtext',hasDownArrow=True,
                                table='shipsteps.email_services_atc', columns='$description', auxColumns='$maintable_id',condition="@maintable_id.service_for_email_id=:cod",condition_cod='garbage',
                                cols=4,popup=True,colspan=2, hasArrowDown=True)]),_onResult="this.form.save();")
       #btn_garb.dataRpc('nome_temp', self.print_template_garbage,record='=#FORM.record',servizio=['garbage'], email_template_id='garbage_email',
       #                    nome_template = 'shipsteps.garbage:garbage_request',format_page='A4',selId='=#FORM.shipsteps_garbage.view.grid.selectedId',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',                            
       #                    _ask=dict(title='!![en]Select the services',fields=[dict(name='services', lbl='!![en]Services', tag='dbSelect',hasDownArrow=True,
       #                    table='shipsteps.email_services', columns='$consignee',condition="$service_for_email_id=:cod",condition_cod='garb',alternatePkey='consignee',
       #                    validate_notnull=True,cols=4,popup=True,colspan=2, hasArrowDown=True),dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
       #                    table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
       #                    cols=4,popup=True,colspan=2)])) 
                           #_ask=dict(title='!![en]Select the services',fields=[dict(name='services', lbl='!![en]Services', tag='dbSelect',hasDownArrow=True,
                           # table='shipsteps.email_services', columns='$consignee',condition="$service_for_email_id=:cod",condition_cod='garb',alternatePkey='consignee',
                           # validate_notnull=True,cols=4,popup=True,colspan=2, hasArrowDown=True),dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                           # table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                           # cols=4,popup=True,colspan=2)]))

        #fb.dataController("if(msgspec=='val_garbage') {SET .email_garbage=true ; alert('Message created');} if(msgspec=='yes') alert('You must select the record as row in the garbage form'); "
        #                    , msgspec='^msg_special')
        fb.field('email_garbage',lbl='', margin_top='5px')
        fb.semaphore('^.email_garbage?=#v==true?true:false', margin_top='5px')
        #verifichiamo quanti servizi PFSO ci sono, nel caso più di uno apparirà la dbSelect per la scelta
        service_for_email = tbl_email_services.query(columns="$service_for_email_id", where='$service_for_email_id=:serv', serv='pfso').fetch()
        serv_len=len(service_for_email)
        btn_pfso = fb.Button('!![en]PFSO', width='10em')
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.email_pfso',button=btn_pfso.js_widget)
        if serv_len > 1:
            btn_pfso.dataRpc('nome_temp', self.email_services,
                       record='=#FORM.record', servizio=['pfso'], email_template_id='email_pfso',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                       _ask=dict(title='!![en]Select the services',fields=[dict(name='services', lbl='!![en]Services', tag='dbSelect',hasDownArrow=True,
                                table='shipsteps.email_services', columns='$consignee', auxColumns='$email,$email_cc,$email_bcc,$email_pec,$email_cc_pec',condition="$service_for_email_id=:cod",condition_cod='pfso',alternatePkey='consignee',
                                validate_notnull=True,cols=4,popup=True,colspan=2, hasArrowDown=True),dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                                table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                                cols=4,popup=True,colspan=2),dict(name='std_att', lbl='!![en]Service attachments', tag='checkboxtext',hasDownArrow=True,
                                table='shipsteps.email_services_atc', columns='$description', auxColumns='$maintable_id',condition="@maintable_id.service_for_email_id=:cod",condition_cod='pfso',
                                cols=4,popup=True,colspan=2, hasArrowDown=True)]),_onResult="this.form.save();") 
        else:
            btn_pfso.dataRpc('nome_temp', self.email_services,
                       record='=#FORM.record', servizio=['pfso'], email_template_id='email_pfso',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                       _ask=dict(title='!![en]Select the Attachments',fields=[dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                                 table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                                 cols=4,popup=True,colspan=2),dict(name='std_att', lbl='!![en]Service attachments', tag='checkboxtext',hasDownArrow=True,
                                table='shipsteps.email_services_atc', columns='$description', auxColumns='$maintable_id',condition="@maintable_id.service_for_email_id=:cod",condition_cod='pfso',
                                cols=4,popup=True,colspan=2, hasArrowDown=True)]),_onResult="this.form.save();")                                 
       # fb.dataController("if(msgspec=='val_pfso') {SET .email_pfso=true ; alert('Message created')}", msgspec='^msg_special')
        fb.field('email_pfso',lbl='', margin_top='5px')
        fb.semaphore('^.email_pfso?=#v==true?true:false', margin_top='5px')
        #verifichiamo quanti servizi Chimico di porto ci sono, nel caso più di uno apparirà la dbSelect per la scelta
        service_for_email = tbl_email_services.query(columns="$service_for_email_id", where='$service_for_email_id=:serv', serv='chem').fetch()
        serv_len=len(service_for_email)
        btn_chem = fb.Button('!![en]Chemist', width='10em',hidden="^#FORM.record.@tip_mov.code?=#v=='pass'")#attributo hidden nascondiamo il widget se il valore tip_mov = pass
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.email_chemist',button=btn_chem.js_widget)
        if serv_len >1:
            btn_chem.dataRpc('nome_temp', self.email_services,
                       record='=#FORM.record', servizio=['chemist'], email_template_id='email_chemist',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                        _ask=dict(title='!![en]Select the services',fields=[dict(name='services', lbl='!![en]Services', tag='dbSelect',hasDownArrow=True,
                                table='shipsteps.email_services', columns='$consignee', auxColumns='$email,$email_cc,$email_bcc,$email_pec,$email_cc_pec',condition="$service_for_email_id=:cod",condition_cod='chem',alternatePkey='consignee',
                                validate_notnull=True,cols=4,popup=True,colspan=2, hasArrowDown=True),dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                                table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                                cols=4,popup=True,colspan=2),dict(name='std_att', lbl='!![en]Service attachments', tag='checkboxtext',hasDownArrow=True,
                                table='shipsteps.email_services_atc', columns='$description', auxColumns='$maintable_id',condition="@maintable_id.service_for_email_id=:cod",condition_cod='chemist',
                                cols=4,popup=True,colspan=2, hasArrowDown=True)]),_onResult="this.form.save();") 
        else:
            btn_chem.dataRpc('nome_temp', self.email_services,
                       record='=#FORM.record', servizio=['chemist'], email_template_id='email_chemist',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                        _ask=dict(title='!![en]Select the Attachments',fields=[dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                                 table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                                 cols=4,popup=True,colspan=2),dict(name='std_att', lbl='!![en]Service attachments', tag='checkboxtext',hasDownArrow=True,
                                table='shipsteps.email_services_atc', columns='$description', auxColumns='$maintable_id',condition="@maintable_id.service_for_email_id=:cod",condition_cod='chemist',
                                cols=4,popup=True,colspan=2, hasArrowDown=True)]),_onResult="this.form.save();")                                 
        #fb.dataController("if(msgspec=='val_chemist') {SET .email_chemist=true ; alert('Message created')}", msgspec='^msg_special')
        fb.field('email_chemist',lbl='', margin_top='5px',hidden="^#FORM.record.@tip_mov.code?=#v=='pass'")#attributo hidden nascondiamo il widget se il valore tip_mov = pass
        fb.semaphore('^.email_chemist?=#v==true?true:false', margin_top='5px',hidden="^#FORM.record.@tip_mov.code?=#v=='pass'")#attributo hidden nascondiamo il widget se il valore tip_mov = pass
        #verifichiamo quanti servizi GPG ci sono, nel caso più di uno apparirà la dbSelect per la scelta
        service_for_email = tbl_email_services.query(columns="$service_for_email_id", where='$service_for_email_id=:serv', serv='gpg').fetch()
        serv_len=len(service_for_email)
        btn_gpg = fb.Button('!![en]GPG', width='10em')
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.email_gpg',button=btn_gpg.js_widget)
        if serv_len > 1:
            btn_gpg.dataRpc('nome_temp', self.email_services,
                      record='=#FORM.record', servizio=['gpg'], email_template_id='email_gpg',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                      _ask=dict(title='!![en]Select the services',fields=[dict(name='services', lbl='!![en]Services', tag='dbSelect',hasDownArrow=True,
                                table='shipsteps.email_services', columns='$consignee', auxColumns='$email,$email_cc,$email_bcc,$email_pec,$email_cc_pec',condition="$service_for_email_id=:cod",condition_cod='gpg',alternatePkey='consignee',
                                validate_notnull=True,cols=4,popup=True,colspan=2, hasArrowDown=True),dict(name='allegati', lbl='!![en]Arrival attachments', tag='checkboxtext',
                                table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                                cols=4,popup=True,colspan=2),dict(name='std_att', lbl='!![en]Service attachments', tag='checkboxtext',hasDownArrow=True,
                                table='shipsteps.email_services_atc', columns='$description', auxColumns='$maintable_id',condition="@maintable_id.service_for_email_id=:cod",condition_cod='gpg',
                                cols=4,popup=True,colspan=2, hasArrowDown=True)]),_onResult="this.form.save();")
        else:
            btn_gpg.dataRpc('nome_temp', self.email_services,
                      record='=#FORM.record', servizio=['gpg'], email_template_id='email_gpg',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                      _ask=dict(title='!![en]Select the Attachments',fields=[dict(name='allegati', lbl='!![en]Arrival attachments', tag='checkboxtext',
                                 table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                                 cols=4,popup=True,colspan=2),dict(name='std_att', lbl='!![en]Service attachments', tag='checkboxtext',hasDownArrow=True,
                                table='shipsteps.email_services_atc', columns='$description', auxColumns='$maintable_id',condition="@maintable_id.service_for_email_id=:cod",condition_cod='gpg',
                                cols=4,popup=True,colspan=2, hasArrowDown=True)]),_onResult="this.form.save();")                                 
       # fb.dataController("if(msgspec=='val_gpg') {SET .email_gpg=true ; alert('Message created')}", msgspec='^msg_special')
        fb.field('email_gpg',lbl='', margin_top='5px')
        fb.semaphore('^.email_gpg?=#v==true?true:false', margin_top='5px')

        btn_ens = fb.Button('!![en]ENS', width='10em',hidden="^#FORM.record.@tip_mov.code?=#v=='pass'")#attributo hidden nascondiamo il widget se il valore tip_mov = pass
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.email_ens',button=btn_ens.js_widget)
        btn_ens.dataRpc('nome_temp', self.email_services,
                  record='=#FORM.record', servizio=['ens'], email_template_id='email_ens',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                  _ask=dict(title='!![en]Select the Attachments',fields=[dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                             table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                             cols=4,popup=True,colspan=2)]),_onResult="this.form.save();")
       # fb.dataController("if(msgspec=='val_ens') {SET .email_ens=true ; alert('Message created')}", msgspec='^msg_special')
       
        fb.field('email_ens',lbl='', margin_top='5px',hidden="^#FORM.record.@tip_mov.code?=#v=='pass'")#attributo hidden nascondiamo il widget se il valore tip_mov = pass
        fb.semaphore('^.email_ens?=#v==true?true:false', margin_top='5px',hidden="^#FORM.record.@tip_mov.code?=#v=='pass'")#attributo hidden nascondiamo il widget se il valore tip_mov = pass

        #verifichiamo quanti servizi ADSP ci sono, nel caso più di uno apparirà la dbSelect per la scelta
        service_for_email = tbl_email_services.query(columns="$service_for_email_id", where='$service_for_email_id=:serv', serv='adsp').fetch()
        serv_len=len(service_for_email)
        btn_gbadsp = fb.Button('!![en]Garbage ADSP', width='10em')
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.email_garbage_adsp',button=btn_gbadsp.js_widget)
        if serv_len > 1:
            btn_gbadsp.dataRpc('nome_temp', self.email_services,
                  record='=#FORM.record', servizio=['adsp'], email_template_id='not_rifiuti',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                  _ask=dict(title='!![en]Select the Attachments',fields=[dict(name='services', lbl='!![en]Services', tag='dbSelect',hasDownArrow=True,
                                table='shipsteps.email_services', columns='$consignee',condition="$service_for_email_id=:cod",condition_cod='adsp',alternatePkey='consignee',
                                validate_notnull=True,cols=4,popup=True,colspan=2, hasArrowDown=True),dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                             table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',validate_notnull=True,
                             cols=4,popup=True,colspan=2),dict(name='std_att', lbl='!![en]Service attachments', tag='checkboxtext',hasDownArrow=True,
                                table='shipsteps.email_services_atc', columns='$description', auxColumns='$maintable_id',condition="@maintable_id.service_for_email_id=:cod",condition_cod='adsp',
                                cols=4,popup=True,colspan=2, hasArrowDown=True)]),_onResult="this.form.save();")
        else:
            btn_gbadsp.dataRpc('nome_temp', self.email_services,
                  record='=#FORM.record', servizio=['adsp'], email_template_id='not_rifiuti',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                  _ask=dict(title='!![en]Select the Attachments',fields=[dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                             table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',validate_notnull=True,
                             cols=4,popup=True,colspan=2),dict(name='std_att', lbl='!![en]Service attachments', tag='checkboxtext',hasDownArrow=True,
                                table='shipsteps.email_services_atc', columns='$description', auxColumns='$maintable_id',condition="@maintable_id.service_for_email_id=:cod",condition_cod='adsp',
                                cols=4,popup=True,colspan=2, hasArrowDown=True)]),_onResult="this.form.save();")
       
       
       # fb.dataController("if(msgspec=='val_adsp') {SET .email_garbage_adsp=true ; alert('Message created')}", msgspec='^msg_special')
      
        fb.field('email_garbage_adsp',lbl='', margin_top='5px')
        fb.semaphore('^.email_garbage_adsp?=#v==true?true:false', margin_top='5px')


       #btn_af = fb.Button('!![en]Email Antifire', width='101px')
       #fb.field('email_antifire',lbl='', margin_top='5px')
       #fb.semaphore('^.email_antifire', margin_top='5px')
        #verifichiamo quanti servizi Sanimare ci sono, nel caso più di uno apparirà la dbSelect per la scelta
        service_for_email = tbl_email_services.query(columns="$service_for_email_id", where='$service_for_email_id=:serv', serv='usma').fetch()
        serv_len=len(service_for_email)
        btn_riclps = fb.Button('!![en]LPS request', width='10em',hidden="^#FORM.record.@last_port.@nazione_code.ue_san?=#v==true")#nascondiamo il widget in base al valore della pyColumn ue_san nella tabella Nazione pkg Unlocode
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.email_ric_lps',button=btn_riclps.js_widget)
        if serv_len > 1:
            btn_riclps.dataRpc('nome_temp', self.email_services,
                       record='=#FORM.record', servizio=['sanimare'], email_template_id='email_ric_lps',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                       _ask=dict(title='!![en]Select the services and the Attachments to send:<br>Free-pratique-sanitary<br>Sanimare declaration<br>Crew list<br>Sanitation certificate<br>Last 3 CSR',
                                fields=[dict(name='services', lbl='!![en]Services', tag='dbSelect',hasDownArrow=True,
                                table='shipsteps.email_services', columns='$consignee',condition="$service_for_email_id=:cod",condition_cod='usma',alternatePkey='consignee',
                                validate_notnull=True,cols=4,popup=True,colspan=2, hasArrowDown=True),dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                                table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                                cols=4,popup=True,colspan=2)]),_onResult="this.form.save();")
        else:
            btn_riclps.dataRpc('nome_temp', self.email_services,
                       record='=#FORM.record', servizio=['sanimare'], email_template_id='email_ric_lps',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                       _ask=dict(title='!![en]Select the Attachments to send:<br>Free-pratique-sanitary<br>Sanimare declaration<br>Crew list<br>Sanitation certificate<br>Last 3 CSR',fields=[dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                                 table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                                 cols=4,popup=True,colspan=2)]),_onResult="this.form.save();")
       # fb.dataController("if(msgspec=='val_usma') {SET .email_usma=true ; alert('Message created')}", msgspec='^nome_temp')
        fb.field('email_ric_lps',lbl='', margin_top='5px',hidden="^#FORM.record.@last_port.@nazione_code.ue_san?=#v==true")#nascondiamo il widget in base al valore della pyColumn ue_san nella tabella Nazione pkg Unlocode
        fb.semaphore('^.email_ric_lps?=#v==true?true:false', margin_top='5px',hidden="^#FORM.record.@last_port.@nazione_code.ue_san?=#v==true")#nascondiamo il widget in base al valore della pyColumn ue_san nella tabella Nazione pkg Unlocode

        btn_update = fb.Button('!![en]Services updating', width='10em')
        btn_update.dataRpc('nome_temp', self.email_serv_upd,
                   record='=#FORM.record',email_template_id='email_arr_shiprec',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                   _ask=dict(title='!![en]Select the services to update and Attachments',fields=[dict(name='services', lbl='!![en]Services', tag='checkboxtext',
                             table='shipsteps.services_for_email', columns='$description_serv',#values='dogana,gdf,gdf roan,pilot,mooringmen,tug,immigration,sanimare,pfso,garbage,chemist,gpg',
                             validate_notnull=True,cols=4,popup=True,colspan=2),dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                             table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                             cols=4,popup=True,colspan=2)]))
        #fb.dataController("if(msgspec=='val_upd') {alert('Message created')}", msgspec='^msg_special')
        
        fb.div()
        fb.div()
     
        btn_upd_shiprec = fb.Button('!![en]Ship/Rec. updating', width='10em')
        btn_upd_shiprec.dataRpc('nome_temp', self.email_arrival_sof,
                   record='=#FORM.record', servizio=['arr','sof'], email_template_id='email_updating_shiprec',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                   _ask=dict(title='!![en]Select the SOF and Attachments',fields=[dict(name='sof_id', lbl='!![en]sof', tag='dbSelect',columns='$id',
                             hasDownArrow=True, auxColumns='$sof_n,$ship_rec', table='shipsteps.sof',condition="$arrival_id =:cod",
                                                condition_cod='=#FORM.record.id',width='25em',validate_notnull=True),dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                             table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                             cols=4,popup=True,colspan=2)]))
        #datacontroller verifica il valore della variabile msg_special di ritorno dalla funzione per invio email
        #e setta il valore della campo checkbox a true e lancia il messaggio 'Messaggio Creato'
       
        #fb.dataController("if(msgspec=='ship_rec_upd') {alert('Message created')} if(msgspec=='no_email') alert('You must insert destination email as TO or BCC'); if(msgspec=='no_sof') alert('You must select the SOF or you must create new one');", msgspec='^msg_special')
        
        div3=rg_prearrival.div('<center><strong><br>Email to Harbour Master <br> Docs before vessel arrival<br></strong>',width='99%',height='20%',margin='auto',
                        padding='2px',hidden="^#FORM.record.docbefore_cp",
                        border='1px solid silver',
                        margin_top='1px',margin_left='4px')
        fb2 = div3.formbuilder(colspan=1,cols=3, border_spacing='2px',fld_width='150px')
        #verifichiamo quanti servizi CP ci sono, nel caso più di uno apparirà la dbSelect per la scelta
        service_for_email = tbl_email_services.query(columns="$service_for_email_id", where='$service_for_email_id=:serv', serv='cp').fetch()
        serv_len=len(service_for_email)
        btn_integr = fb2.Button('!![en]Alimentary integration',hidden="^#FORM.record.@tip_mov.code?=#v!='alim'")#attributo hidden per nascondere il widget se il valore tip_mov è diverso da alim
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.email_integr',button=btn_integr.js_widget)
        if serv_len >1:
            btn_integr.dataRpc('nome_temp', self.email_services,
                      record='=#FORM.record', servizio=['capitaneria'], email_template_id='email_integrazione_alim',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                       _ask=dict(title='!![en]Select the services',fields=[dict(name='services', lbl='!![en]Services', tag='dbSelect',hasDownArrow=True,
                                table='shipsteps.email_services', columns='$consignee', auxColumns='$email,$email_cc,$email_bcc,$email_pec,$email_cc_pec',condition="$service_for_email_id=:cod",condition_cod='cp',alternatePkey='consignee',
                                validate_notnull=True,cols=4,popup=True,colspan=2, hasArrowDown=True),dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                                table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                                cols=4,popup=True,colspan=2)]),_onResult="this.form.save();")
        else:    
            btn_integr.dataRpc('nome_temp', self.email_services,
                      record='=#FORM.record', servizio=['capitaneria'], email_template_id='email_integrazione_alim',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                      _ask=dict(title='!![en]Select the Attachments',fields=[dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                                 table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                                 cols=4,popup=True,colspan=2)]),_onResult="this.form.save();")
        #fb2.dataController("if(msgspec=='val_integr') {SET .email_integr=true ; alert('Message created')}", msgspec='^msg_special')
       
        fb2.field('email_integr', lbl='', margin_top='6px',hidden="^#FORM.record.@tip_mov.code?=#v!='alim'")#attributo hidden per nascondere il widget se il valore tip_mov è diverso da alim
        fb2.semaphore('^.email_integr?=#v==true?true:false', margin_top='6px',hidden="^#FORM.record.@tip_mov.code?=#v!='alim'")#attributo hidden per nascondere il widget se il valore tip_mov è diverso da alim

        #verifichiamo quanti servizi CP ci sono, nel caso più di uno apparirà la dbSelect per la scelta
        service_for_email = tbl_email_services.query(columns="$service_for_email_id", where='$service_for_email_id=:serv', serv='cp').fetch()
        serv_len=len(service_for_email)
        btn_pmou = fb2.Button('!![en]PMOU notification',hidden='^gnr.app_preference.shipsteps.pmou')#attributo hidden per nascondere il widget se il valore nelle preferenze pmou è True
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.email_pmou',button=btn_pmou.js_widget)
        if serv_len >1:
            btn_pmou.dataRpc('nome_temp', self.email_services,
                      record='=#FORM.record', servizio=['capitaneria'], email_template_id='email_pmou',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                       _ask=dict(title='!![en]Select the services',fields=[dict(name='services', lbl='!![en]Services', tag='dbSelect',hasDownArrow=True,
                                table='shipsteps.email_services', columns='$consignee', auxColumns='$email,$email_cc,$email_bcc,$email_pec,$email_cc_pec',condition="$service_for_email_id=:cod",condition_cod='cp',alternatePkey='consignee',
                                validate_notnull=True,cols=4,popup=True,colspan=2, hasArrowDown=True),dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                                table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                                cols=4,popup=True,colspan=2)]),_onResult="this.form.save();")
        else:    
            btn_pmou.dataRpc('nome_temp', self.email_services,
                      record='=#FORM.record', servizio=['capitaneria'], email_template_id='email_pmou',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                      _ask=dict(title='!![en]Select the Attachments',fields=[dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                                 table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                                 cols=4,popup=True,colspan=2)]),_onResult="this.form.save();")
        #fb2.dataController("if(msgspec=='val_integr') {SET .email_integr=true ; alert('Message created')}", msgspec='^msg_special')
       
        fb2.field('email_pmou', lbl='', margin_top='6px',hidden='^gnr.app_preference.shipsteps.pmou')#attributo hidden per nascondere il widget se il valore nelle preferenze pmou è True
        fb2.semaphore('^.email_pmou?=#v==true?true:false', margin_top='6px',hidden='^gnr.app_preference.shipsteps.pmou')#attributo hidden per nascondere il widget se il valore nelle preferenze pmou è True

        #btn_garb_cp = fb2.Button('!![en]Email Garbage form')
        #btn_garb_cp.dataRpc('nome_temp', self.email_services,
        #          record='=#FORM.record.id', servizio=['capitaneria'], email_template_id='email_garbage_cp',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
        #          _ask=dict(title='!![en]Select the Attachments',fields=[dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
        #                     table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
        #                     cols=4,popup=True,colspan=2)]),_onResult="this.form.save();")
        #fb2.dataController("if(msgspec=='val_garb_cp') {SET .email_garbage_cp=true ; alert(msg_txt)}", msgspec='^msg_special',msg_txt = 'Email ready to be sent')
       
        #fb2.field('email_garbage_cp', lbl='', margin_top='6px')
        #fb2.semaphore('^.email_garbage_cp', margin_top='6px')
        
        fb.dataController("""if(gdfdep==true) {alert(gdfdep_txt);}""",gdfdep='^#FORM.record.gdfdep_timeexp',gdfdep_txt='Print the GDF Form vessel departure')
        #fb.dataController("if(msgspec=='val_bulk')alert(msg_txt);",msgspec='^msg_special',msg_txt = 'Email ready to be sent')
        fb.dataController("""if(msg=='val_bulk'){alert(msg_txt);} if(msg=='val_garb_cp'){SET .email_garbage_cp=true ; alert(msg_txt);}
                             if(msg=='val_integr') {SET .email_integr=true ;genro.publish("floating_message",{message:msg_txt, messageType:"message"});}
                             if(msg=='val_pmou') {SET .email_pmou=true ;genro.publish("floating_message",{message:msg_txt, messageType:"message"});}
                             if(msg=='ship_rec_upd') genro.publish("floating_message",{message:msg_txt, messageType:"message"}); if(msg=='no_email') genro.publish("floating_message",{message:'You must insert destination email as TO or BCC', messageType:"error"}); if(msg=='no_sof') genro.publish("floating_message",{message:'You must select the SOF or you must create new one', messageType:"error"});
                             if(msg=='val_upd') genro.publish("floating_message",{message:msg_txt, messageType:"message"});
                             if(msg=='val_adsp') {SET .email_garbage_adsp=true ; genro.publish("floating_message",{message:msg_txt, messageType:"message"});}
                             if(msg=='val_ens') {SET .email_ens=true ; genro.publish("floating_message",{message:msg_txt, messageType:"message"});}
                             if(msg=='val_gpg') {SET .email_gpg=true ; genro.publish("floating_message",{message:msg_txt, messageType:"message"});}
                             if(msg=='val_chemist') {SET .email_chemist=true ; genro.publish("floating_message",{message:msg_txt, messageType:"message"});}
                             if(msg=='val_pfso') {SET .email_pfso=true ; genro.publish("floating_message",{message:msg_txt, messageType:"message"});}
                             if(msg=='val_garbage') {SET .email_garbage=true ;alert(msg_txt);} if(msg=='yes') genro.publish("floating_message",{message:'You must select the record as row in the garbage form', messageType:"error"});
                             if(msg=='val_tug') {SET .email_tug=true; genro.publish("floating_message",{message:msg_txt, messageType:"message"});}
                             if(msg=='val_tug_dep') {SET .email_tug_dep=true; genro.publish("floating_message",{message:msg_txt, messageType:"message"});}
                             if(msg=='val_pil_moor') {SET .email_pilot_moor=true; genro.publish("floating_message",{message:msg_txt, messageType:"message"});}
                             if(msg=='val_usma') {SET .email_usma=true; genro.publish("floating_message",{message:msg_txt, messageType:"message"});}
                             if(msg=='val_lps') {SET .email_ric_lps=true; genro.publish("floating_message",{message:msg_txt, messageType:"message"});}
                             if(msg=='val_lps_cp') {SET .email_lps_cp=true; genro.publish("floating_message",{message:msg_txt, messageType:"message"});}
                             if(msg=='form_immigration') {SET .email_frontiera=true; alert(msg_txt);}
                             if(msg=='val_dog') {SET .email_dogana=true; genro.publish("floating_message",{message:msg_txt, messageType:"message"});}
                             if(msg=='ship_rec') {SET .email_ship_rec=true; genro.publish("floating_message",{message:msg_txt, messageType:"message"});} if(msg=='no_email') genro.publish("floating_message",{message:'You must insert destination email as TO or BCC', messageType:"error"}); if(msg=='no_sof') genro.publish("floating_message",{message:'You must select the SOF or you must create new one', messageType:"error"});
                             if(msg=='val_chemist_cp') {SET .email_certchim_cp=true; genro.publish("floating_message",{message:msg_txt, messageType:"message"});}
                             if(msg=='val_ricrifiuti') {SET .email_ric_rifiuti_cp=true; genro.publish("floating_message",{message:msg_txt, messageType:"message"});}
                             if(msg=='val_tributi') {SET .email_tributi_cp=true; genro.publish("floating_message",{message:msg_txt, messageType:"message"});}
                             if(msg=='vess_serv') {SET .form_services=true;} this.form.save();
                             if(msg=='val_deroga_gb') if(msg=='val_deroga_gb'){SET .email_garbage_cp=true ; alert(msg_txt);}
                             if(msg=='no_moored') {genro.publish("floating_message",{message:'You must insert in arrival times date and time of vessel moored', messageType:"error"});}
                             if(msg=='no_next_port') {genro.publish("floating_message",{message:'You must insert next port', messageType:"error"});}
                             if(msg=='mod61_arr') {alert(msg_txt);} if(msg=='nota_arr_no') genro.publish("floating_message",{message:'You must first print Nota Arrivo', messageType:"error"}); if(msg=='fal1_arr_no') genro.publish("floating_message",{message:'You must first print Fal1 arrival', messageType:"error"}); if(msg=='fal1arr_notarr') genro.publish("floating_message",{message:'You must first print Fal1 arrival and Nota Arrivo', messageType:"error"});
                             if(msg=='mod61_dep') {alert(msg_txt);} if(msg=='nota_part_no') genro.publish("floating_message",{message:'You must first print Dich. integrativa di partenza', messageType:"error"}); if(msg=='fal1_dep_no') genro.publish("floating_message",{message:'You must first print Fal1 departure', messageType:"error"}); if(msg=='fal1dep_notapart') genro.publish("floating_message",{message:'You must first print Fal1 departure and Dich. integrativa di partenza', messageType:"error"}); if(msg=='no_sailed') genro.publish("floating_message",{message:'You must first insert ets date and time', messageType:"error"});
                             if(msg=='fal1dep_notapart_dispval') genro.publish("floating_message",{message:'You must first print Fal1 departure and Dich. integrativa di partenza and Currency availability', messageType:"error"});
                             if(msg=='dispval_no') genro.publish("floating_message",{message:'You must first print Currency availability', messageType:"error"});
                             if(msg=='intfat') genro.publish("floating_message",{message:msg_txt, messageType:"message"});
                             if(msg=='no_sanitation') genro.publish("floating_message",{message:'you must insert sanitation certificate in the ships docs', messageType:"error"});
                             if(msg=='no_nsisprot') genro.publish("floating_message",{message:'you must insert nsis prot. number in the arrival form', messageType:"error"});
                             if(msg=='mod_nave') {SET .modulo_nave=true;}
                             if(msg=='front_carico') {SET .front_carico=true;}
                             if(msg=='tab_servizi') {SET .tab_servizi=true;}
                             if(msg=='cartella_doc') {SET .cartella_nave=true;}
                             if(msg=='front_nave') {SET .frontespizio=true;}
                             if(msg=='check_list') {SET .checklist=true;}
                             if(msg=='check_list_alim') {SET .checklist=true;}
                             if(msg=='check_list_pass') {SET .checklist=true;}
                             if(msg=='tributi_cp') {SET .tributi_cp=true;} if(msg=='no_tributi') genro.publish("floating_message",{message:'You must select the record as row in the tribute form', messageType:"error"});
                             if(msg=='form_gdf') {SET .form_gdf=true;}
                             if(msg=='form_immigration_print') {SET .form_immigration=true;}
                             if(msg=='form_provisions') {SET .form_provision=true;}
                             if(msg=='master_info') {SET .master_info=true;}
                             if(msg=='DichSanimare') {SET .form_sanimare=true;}
                             if(msg=='InterferenzeFiore') {SET .form_checklist_f=true;}"""
                             ,msg='^nome_temp',msg_txt = 'Email ready to be sent')

       #fb.dataController("""if(msgspec=='val_bulk'){alert(msg_txt);} if(msgspec=='val_garb_cp'){SET .email_garbage_cp=true ; alert(msg_txt);}
       #                     if(msgspec=='val_integr') {SET .email_integr=true ; alert(msg_txt);}
       #                     if(msgspec=='ship_rec_upd') {alert(msg_txt);} if(msgspec=='no_email') {alert('You must insert destination email as TO or BCC');} if(msgspec=='no_sof') {alert('You must select the SOF or you must create new one');}
       #                     if(msgspec=='val_upd') {alert(msg_txt);}
       #                     if(msgspec=='val_adsp') {SET .email_garbage_adsp=true ; alert(msg_txt);}
       #                     if(msgspec=='val_ens') {SET .email_ens=true ; alert(msg_txt);}
       #                     if(msgspec=='val_gpg') {SET .email_gpg=true ; alert(msg_txt);}
       #                     if(msgspec=='val_chemist') {SET .email_chemist=true ; alert(msg_txt);}
       #                     if(msgspec=='val_pfso') {SET .email_pfso=true ; alert(msg_txt);}
       #                     if(msgspec=='val_garbage') {SET .email_garbage=true ; alert(msg_txt);} if(msgspec=='yes') {alert('You must select the record as row in the garbage form');}
       #                     if(msgspec=='val_tug') {SET .email_tug=true ; alert(msg_txt);}
       #                     if(msgspec=='val_pil_moor') {SET .email_pilot_moor=true ; alert(msg_txt);}
       #                     if(msgspec=='val_usma') {SET .email_usma=true ; alert(msg_txt);}
       #                     if(msgspec=='val_imm') {SET .email_frontiera=true; alert(msg_txt);}
       #                     if(msgspec=='val_dog') {SET .email_dogana=true ; alert(msg_txt);}
       #                     if(msgspec=='ship_rec') {SET .email_ship_rec=true ; alert(msg_txt);} if(msgspec=='no_email') {alert('You must insert destination email as TO or BCC');} if(msgspec=='no_sof') {alert('You must select the SOF or you must create new one');}
       #                     if(msg=='mod61_arr') {alert(msg_txt);} if(msg=='nota_arr_no') {alert('You must first print Nota Arrivo');} if(msg=='fal1_arr_no') {alert('You must first print Fal1 arrival');} if(msg=='fal1arr_notarr') {alert('You must first print Fal1 arrival and Nota Arrivo');}
       #                     if(msg=='mod61_dep') {alert(msg_txt);} if(msg=='nota_part_no') {alert('You must first print Dich. integrativa di partenza');} if(msg=='fal1_dep_no') {alert('You must first print Fal1 departure');} if(msg=='fal1dep_notapart') {alert('You must first print Fal1 departure and Dich. integrativa di partenza');} if(msg=='no_sailed') {alert('You must first insert ets date and time');}
       #                     if(msg=='intfat') genro.publish("floating_message",{message:"email ready to be sent", messageType:"message"});
       #                     if(msg=='mod_nave') {SET .checklist=true;}
       #                     if(msg=='front_carico') {SET .front_carico=true;}
       #                     if(msg=='tab_servizi') {SET .tab_servizi=true;}
       #                     if(msg=='cartella_doc') {SET .cartella_nave=true;}
       #                     if(msg=='front_nave') {SET .frontespizio=true;}
       #                     if(msg=='check_list') {SET .checklist=true;}""",msgspec='^msg_special', msg='^nome_temp',msg_txt = 'Email ready to be sent')

        
        div_arr=rg_arrival.div('<center><strong>ARRIVAL</strong>',width='99%',height='20%',margin='auto',
                        padding='2px',
                        border='1px solid silver',
                        margin_top='1px',margin_left='4px')
        fb_arr=div_arr.formbuilder(colspan=1,cols=3, border_spacing='1px', fld_width='14em')
       
        #btn_fgdf=fb_arr.Button('!![en]Test2',action="""
        #                       var tp = {template:template_id};
        #                       var kw = objectExtract(this.getInheritedAttributes(),"batch_*",true);
        #                       kw.table = 'shipsteps.arrival'; 
        #                       kw.resource = "print_template";
        #                       kw.res_type = "print";
        #                       kw.pkey = this.form.getCurrentPkey();
        #                       kw.extra_parameters = new gnr.GnrBag({template_id:tp.template,table:kw.table});
        #                       console.log(kw);
        #                       genro.publish("table_script_run",kw)""",template_id='tjZPgxoKNZiJgvTEYCEYKg')
        #                      
        #fb_arr.br()
        #fb_arr.br()
        btn_fgdf_cp = fb_arr.Button('!![en]Form GdF')
        
       
        #btn_fgdf_cp.dataController("SET .p_date=pratique_date;",_ask=dict(title='!![en]Insert pratique date',
        #                                            fields=[dict(name='pratique_date', lbl='!![en]Date', tag='dateTextBox',
        #                                            cols=4,popup=True,colspan=2)])) 
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.form_gdf',button=btn_fgdf_cp.js_widget)
        
        btn_fgdf_cp.dataRpc('nome_temp', self.print_template,record='=#FORM.record',_ask=dict(title='!![en]Insert pratique date',
                                                    fields=[dict(name='pratique_date', lbl='!![en]Date', tag='dateTextBox',
                                                    cols=4,popup=True,colspan=2)]),#_onCalling="{SET .date_pr=pratique_date;}this.form.save();",
                                                    pratique_date='^#FORM.record.etb_date',
                            nome_template = 'shipsteps.arrival:form_gdf', nome_vs='=#FORM.record.@vessel_details_id.@imbarcazione_id.nome',
                            format_page='A4',_onResult="this.form.save();")
        #btn_fgdf_cp.dataRpc('nome_temp', self.print_template,record='=#FORM.record',
        #                    nome_template = 'shipsteps.arrival:form_gdf', nome_vs='=#FORM.record.@vessel_details_id.@imbarcazione_id.nome',
        #                    format_page='A4',_onResult="this.form.save();")
        fb_arr.field('form_gdf', lbl='', margin_top='6px')
        fb_arr.semaphore('^.form_gdf?=#v==true?true:false', margin_top='6px')

        btn_fimm_cp = fb_arr.Button('!![en]Form Immigration')
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.form_immigration',button=btn_fimm_cp.js_widget)
        btn_fimm_cp.dataRpc('nome_temp', self.print_template,record='=#FORM.record', nome_template = 'shipsteps.arrival:form_immigration', 
                                         nome_vs='=#FORM.record.@vessel_details_id.@imbarcazione_id.nome',format_page='A4', 
                                         only_print='yes',_onResult="this.form.save();")
        fb_arr.field('form_immigration', lbl='', margin_top='6px')
        fb_arr.semaphore('^.form_immigration?=#v==true?true:false', margin_top='6px')

        btn_fprov_cp = fb_arr.Button('!![en]Form Provisions')
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.form_provision',button=btn_fprov_cp.js_widget)
        btn_fprov_cp.dataRpc('nome_temp', self.print_template,record='=#FORM.record',_ask=dict(title='!![en]Insert pratique date',
                                                    fields=[dict(name='pratique_date', lbl='!![en]Date', tag='dateTextBox',
                                                    cols=4,popup=True,colspan=2)]),#_onCalling="{SET .date_pr=pratique_date;}this.form.save();",
                                                    pratique_date='^#FORM.record.etb_date',
                            nome_template = 'shipsteps.arrival:form_provisions', nome_vs='=#FORM.record.@vessel_details_id.@imbarcazione_id.nome',
                            format_page='A4',_onResult="this.form.save();")
        fb_arr.field('form_provision', lbl='', margin_top='6px')
        fb_arr.semaphore('^.form_provision?=#v==true?true:false', margin_top='6px')

        #fb_arr.br()
        
        btn_fsan = fb_arr.Button('!![en]Sanimare declaration',hidden="^#FORM.record.@last_port.@nazione_code.ue_san?=#v==true")#nascondiamo il widget in base al valore della pyColumn ue_san nella tabella Nazione pkg Unlocode
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.form_sanimare',button=btn_fsan.js_widget)
        btn_fsan.dataRpc('nome_temp', self.apridoc,record='=#FORM.record',nome_form='DichSanimare', 
                                        _virtual_column='lastport,nextport,vesselname,flag,imo,tsl',_onResult="this.form.save();")
        fb_arr.field('form_sanimare', lbl='', margin_top='6px',hidden="^#FORM.record.@last_port.@nazione_code.ue_san?=#v==true")#nascondiamo il widget in base al valore della pyColumn ue_san nella tabella Nazione pkg Unlocode
        fb_arr.semaphore('^.form_sanimare?=#v==true?true:false', margin_top='6px',hidden="^#FORM.record.@last_port.@nazione_code.ue_san?=#v==true")#nascondiamo il widget in base al valore della pyColumn ue_san nella tabella Nazione pkg Unlocode

        btn_intfiore = fb_arr.Button('!![en]CheckList Fiore',hidden="^#FORM.record.@tip_mov.code?=#v=='pass'")#attributo hidden nascondiamo il widget se il valore tip_mov = pass
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.form_checklist_f',button=btn_intfiore.js_widget)
        btn_intfiore.dataRpc('nome_temp', self.apridoc,record='=#FORM.record',nome_form='InterferenzeFiore', 
                                          _virtual_column='lastport,nextport,vesselname,flag,imo,tsl',_onResult="this.form.save();")
        fb_arr.field('form_checklist_f', lbl='', margin_top='6px',hidden="^#FORM.record.@tip_mov.code?=#v=='pass'")#attributo hidden nascondiamo il widget se il valore tip_mov = pass
        fb_arr.semaphore('^.form_checklist_f?=#v==true?true:false', margin_top='6px',hidden="^#FORM.record.@tip_mov.code?=#v=='pass'")#attributo hidden nascondiamo il widget se il valore tip_mov = pass
        fb_arr.br()

        btn_timegate = fb_arr.Button('!![en]Times gate info')
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.master_info',button=btn_timegate.js_widget)
        btn_timegate.dataRpc('nome_temp', self.print_template_gate,record='=#FORM.record',
                            nome_template = 'shipsteps.opening_gate:time_gate', nome_vs='=#FORM.record.@vessel_details_id.@imbarcazione_id.nome',
                            format_page='A4',_onResult="this.form.save();")
        fb_arr.field('master_info', lbl='', margin_top='6px')
        fb_arr.semaphore('^.master_info?=#v==true?true:false', margin_top='6px')

        #verifichiamo quanti servizi CP ci sono, nel caso più di uno apparirà la dbSelect per la scelta
        service_for_email = tbl_email_services.query(columns="$service_for_email_id", where='$service_for_email_id=:serv', serv='cp').fetch()
        serv_len=len(service_for_email)
        btn_lps_cp = fb_arr.Button('!![en]Email LPS CP',hidden="^#FORM.record.@last_port.@nazione_code.ue_san?=#v==true")#nascondiamo il widget in base al valore della pyColumn ue_san nella tabella Nazione pkg Unlocode
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.email_lps_cp',button=btn_lps_cp.js_widget)
        if serv_len > 1:
            btn_lps_cp.dataRpc('nome_temp', self.email_services,
                      record='=#FORM.record', servizio=['capitaneria'], email_template_id='email_lps_cp',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                      _ask=dict(title='!![en]Select the services',fields=[dict(name='services', lbl='!![en]Services', tag='dbSelect',hasDownArrow=True,
                                table='shipsteps.email_services', columns='$consignee', auxColumns='$email,$email_cc,$email_bcc,$email_pec,$email_cc_pec',condition="$service_for_email_id=:cod",condition_cod='cp',alternatePkey='consignee',
                                validate_notnull=True,cols=4,popup=True,colspan=2, hasArrowDown=True),dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                                table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                                cols=4,popup=True,colspan=2)]),_onResult="this.form.save();")
        else:
            btn_lps_cp.dataRpc('nome_temp', self.email_services,
                  record='=#FORM.record', servizio=['capitaneria'], email_template_id='email_lps_cp',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                  _ask=dict(title='!![en]Select the Attachments',fields=[dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                             table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',validate_notnull=True,
                             cols=4,popup=True,colspan=2)]),_onResult="this.form.save();")

        fb_arr.field('email_lps_cp', lbl='', margin_top='6px',hidden="^#FORM.record.@last_port.@nazione_code.ue_san?=#v==true")#nascondiamo il widget in base al valore della pyColumn ue_san nella tabella Nazione pkg Unlocode
        fb_arr.semaphore('^.email_lps_cp?=#v==true?true:false', margin_top='6px',hidden="^#FORM.record.@last_port.@nazione_code.ue_san?=#v==true")#nascondiamo il widget in base al valore della pyColumn ue_san nella tabella Nazione pkg Unlocode

        #verifichiamo quanti servizi CP ci sono, nel caso più di uno apparirà la dbSelect per la scelta
        service_for_email = tbl_email_services.query(columns="$service_for_email_id", where='$service_for_email_id=:serv', serv='cp').fetch()
        serv_len=len(service_for_email)
        btn_chim_cp = fb_arr.Button('!![en]Email Chemist Cert. CP',hidden="^#FORM.record.@tip_mov.code?=#v=='pass'")#attributo hidden nascondiamo il widget se il valore tip_mov = pass
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.email_certchim_cp',button=btn_chim_cp.js_widget)
        if serv_len > 1:
            btn_chim_cp.dataRpc('nome_temp', self.email_services,
                      record='=#FORM.record', servizio=['capitaneria'], email_template_id='email_chimico_cp',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                      _ask=dict(title='!![en]Select the services',fields=[dict(name='services', lbl='!![en]Services', tag='dbSelect',hasDownArrow=True,
                                table='shipsteps.email_services', columns='$consignee', auxColumns='$email,$email_cc,$email_bcc,$email_pec,$email_cc_pec',condition="$service_for_email_id=:cod",condition_cod='cp',alternatePkey='consignee',
                                validate_notnull=True,cols=4,popup=True,colspan=2, hasArrowDown=True),dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                                table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                                cols=4,popup=True,colspan=2)]),_onResult="this.form.save();")
        else:
            btn_chim_cp.dataRpc('nome_temp', self.email_services,
                  record='=#FORM.record', servizio=['capitaneria'], email_template_id='email_chimico_cp',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                  _ask=dict(title='!![en]Select the Attachments',fields=[dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                             table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',validate_notnull=True,
                             cols=4,popup=True,colspan=2)]),_onResult="this.form.save();")

        fb_arr.field('email_certchim_cp', lbl='', margin_top='6px',hidden="^#FORM.record.@tip_mov.code?=#v=='pass'")#attributo hidden nascondiamo il widget se il valore tip_mov = pass
        fb_arr.semaphore('^.email_certchim_cp?=#v==true?true:false', margin_top='6px',hidden="^#FORM.record.@tip_mov.code?=#v=='pass'")#attributo hidden nascondiamo il widget se il valore tip_mov = pass

        #verifichiamo quanti servizi CP ci sono, nel caso più di uno apparirà la dbSelect per la scelta
        service_for_email = tbl_email_services.query(columns="$service_for_email_id", where='$service_for_email_id=:serv', serv='cp').fetch()
        serv_len=len(service_for_email)
        fb_arr.br()
        btn_der_cp = fb_arr.Button('!![en]Email Waste derogation CP')
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.email_garbage_cp',button=btn_der_cp.js_widget)
        if serv_len > 1:
            btn_der_cp.dataRpc('nome_temp', self.print_template_derogagb,
                      record='=#FORM.record', servizio=['capitaneria'], email_template_id='email_deroga_garbage',
                                imbarcazione_id='=#FORM.record.@vessel_details_id.imbarcazione_id',nome_template = 'shipsteps.arrival:deroga_rifiuti',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                                moored='=#FORM.record.@time_arr.moored',nextport='=#FORM.record.@next_port.descrizione',
                      _ask=dict(title='!![en]Select the services',fields=[dict(name='services', lbl='!![en]Services', tag='dbSelect',hasDownArrow=True,
                                table='shipsteps.email_services', columns='$consignee', auxColumns='$email,$email_cc,$email_bcc,$email_pec,$email_cc_pec',condition="$service_for_email_id=:cod",condition_cod='cp',alternatePkey='consignee',
                                validate_notnull=True,cols=4,popup=True,colspan=2, hasArrowDown=True),dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                                table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                                cols=4,popup=True,colspan=2)]),_onResult="this.form.save();")
        else:                         
            btn_der_cp.dataRpc('nome_temp', self.print_template_derogagb,
                      record='=#FORM.record.id', servizio=['capitaneria'], email_template_id='email_deroga_garbage',
                                imbarcazione_id='=#FORM.record.@vessel_details_id.imbarcazione_id',nome_template = 'shipsteps.arrival:deroga_rifiuti',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                                moored='=#FORM.record.@time_arr.moored',nextport='=#FORM.record.@next_port.descrizione',
                      _ask=dict(title='!![en]Select the Attachments',fields=[dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                                 table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',validate_notnull=True,
                                 cols=4,popup=True,colspan=2)]),_onResult="this.form.save();")

        fb_arr.field('email_garbage_cp', lbl='', margin_top='6px')
        fb_arr.semaphore('^.email_garbage_cp?=#v==true?true:false', margin_top='6px')

        div_dep=rg_arrival.div('<center><strong>DEPARTURE</strong>',width='99%',height='20%',margin='auto',
                        padding='2px',
                        border='1px solid silver',
                        margin_top='1px',margin_left='4px')
        fb_dep=div_dep.formbuilder(colspan=1,cols=3, border_spacing='1px', fld_width='14em')
        #verifichiamo quanti servizi CP ci sono, nel caso più di uno apparirà la dbSelect per la scelta
        service_for_email = tbl_email_services.query(columns="$service_for_email_id", where='$service_for_email_id=:serv', serv='cp').fetch()
        serv_len=len(service_for_email)
        btn_rif_cp = fb_dep.Button('!![en]Email Waste Receipt CP')
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.email_ric_rifiuti_cp',button=btn_rif_cp.js_widget)
        if serv_len > 1:
            btn_rif_cp.dataRpc('nome_temp', self.email_services,
                      record='=#FORM.record', servizio=['capitaneria'], email_template_id='email_ricevutarifiuti_cp',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                      _ask=dict(title='!![en]Select the services',fields=[dict(name='services', lbl='!![en]Services', tag='dbSelect',hasDownArrow=True,
                                table='shipsteps.email_services', columns='$consignee', auxColumns='$email,$email_cc,$email_bcc,$email_pec,$email_cc_pec',condition="$service_for_email_id=:cod",condition_cod='cp',alternatePkey='consignee',
                                validate_notnull=True,cols=4,popup=True,colspan=2, hasArrowDown=True),dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                                table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                                cols=4,popup=True,colspan=2)]),_onResult="this.form.save();")
        else:
            btn_rif_cp.dataRpc('nome_temp', self.email_services,
                      record='=#FORM.record', servizio=['capitaneria'], email_template_id='email_ricevutarifiuti_cp',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                      _ask=dict(title='!![en]Select the Attachments',fields=[dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                                 table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',validate_notnull=True,
                                 cols=4,popup=True,colspan=2)]),_onResult="this.form.save();")
        fb_dep.field('email_ric_rifiuti_cp', lbl='', margin_top='6px')
        fb_dep.semaphore('^.email_ric_rifiuti_cp?=#v==true?true:false', margin_top='6px')
        btn_vs=fb_dep.Button('!![en]Vessel services', action="""{SET tabname='services';}""")
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.form_services',button=btn_vs.js_widget)
        fb_dep.field('form_services', lbl='', margin_top='6px')
        fb_dep.semaphore('^.form_services?=#v==true?true:false', margin_top='6px')

        btn_tugdep = fb_dep.Button('!![en]Tug departure')
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.email_tug_dep',button=btn_tugdep.js_widget)
        btn_tugdep.dataRpc('nome_temp', self.email_services,
                   record='=#FORM.record', servizio=['tug'], email_template_id='email_tug_dep',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                   _ask=dict(title='!![en]Select the Attachments',fields=[dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                             table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                             cols=4,popup=True,colspan=2)]),_onResult="this.form.save();")
       # fb.dataController("if(msgspec=='val_tug') {SET .email_tug=true ; alert('Message created')}", msgspec='^msg_special')
       
        fb_dep.field('email_tug_dep',lbl='', margin_top='5px')
        fb_dep.semaphore('^.email_tug_dep?=#v==true?true:false', margin_top='5px')

        btn_fgdf=fb_dep.Button('!![en]Form GdF Departure',action="""genro.publish("table_script_run",{table:"shipsteps.arrival",
                                                                               res_type:'print',
                                                                               resource:'Partenza_Finanza',
                                                                               pkey: pkey})
                                                                               {SET .form_gdfdep=true;}
                                                                               this.form.save();""",
                                                                               pkey='=#FORM.pkey')
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.form_gdfdep',button=btn_fgdf.js_widget)                                                                               
        fb_dep.field('form_gdfdep', lbl='', margin_top='6px')
        fb_dep.semaphore('^.form_gdfdep?=#v==true?true:false', margin_top='6px')    
        #verifichiamo quanti servizi CP ci sono, nel caso più di uno apparirà la dbSelect per la scelta
        service_for_email = tbl_email_services.query(columns="$service_for_email_id", where='$service_for_email_id=:serv', serv='cp').fetch()
        serv_len=len(service_for_email)
        btn_trib_cp = fb_dep.Button('!![en]Email Tributes CP')
        fb1.dataController("""var id = button.id; console.log(id);
                        if (ca==true){document.getElementById(id).style.backgroundColor = 'lightgreen';}
                        else {document.getElementById(id).style.backgroundColor = '';}
                        """, ca='^.email_tributi_cp',button=btn_trib_cp.js_widget)
        if serv_len > 1:
            btn_trib_cp.dataRpc('nome_temp', self.email_services,
                      record='=#FORM.record', servizio=['capitaneria'], email_template_id='email_tributi_cp',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                      _ask=dict(title='!![en]Select the services',fields=[dict(name='services', lbl='!![en]Services', tag='dbSelect',hasDownArrow=True,
                                table='shipsteps.email_services',columns='$consignee', auxColumns='$email,$email_cc,$email_bcc,$email_pec,$email_cc_pec',condition="$service_for_email_id=:cod",condition_cod='cp',alternatePkey='consignee',
                                validate_notnull=True,cols=4,popup=True,colspan=2, hasArrowDown=True),dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                                table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                                cols=4,popup=True,colspan=2)]),_onResult="this.form.save();") 
        else:
            btn_trib_cp.dataRpc('nome_temp', self.email_services,
                      record='=#FORM.record', servizio=['capitaneria'], email_template_id='email_tributi_cp',selPkeys_att='=#FORM.attachments.view.grid.currentSelectedPkeys',
                      _ask=dict(title='!![en]Select the Attachments',fields=[dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                                 table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',validate_notnull=True,
                                 cols=4,popup=True,colspan=2)]),_onResult="this.form.save();")                                  
        fb_dep.field('email_tributi_cp', lbl='', margin_top='6px')
        fb_dep.semaphore('^.email_tributi_cp?=#v==true?true:false', margin_top='6px')    
        
        #rg_arrival.div('&nbsp').field('nsw', table='shipsteps.tasklist', label='NSW', lbl='Sistema NSW')                                                           
        rg_arrival_nsw.div('&nbsp').checkbox(value='^.nsw', label='NSW', lbl='Sistema NSW')                                                           
        div_nsw=rg_arrival_nsw.div('<center><strong>NSW</strong>',width='99%',height='20%',margin='auto',
                        padding='2px',
                        border='1px solid silver',
                        margin_top='1px',margin_left='4px', hidden='^.nsw?=!#v')
        fb_nsw=div_nsw.formbuilder(colspan=1,cols=1, border_spacing='1px', fld_width='15em')
        
        btn_accosto=fb_nsw.Button('!![en]Email Domanda Accosto',
                                        action="""genro.publish("table_script_run",{table:"shipsteps.arrival",
                                                                               res_type:'print',
                                                                               resource:'Accosto',
                                                                               pkey: pkey});""",
                                                                               pkey='=#FORM.pkey')

        btn_cambio_accosto=fb_nsw.Button('!![en]Email Cambio Accosto',
                                        action="""genro.publish("table_script_run",{table:"shipsteps.arrival",
                                                                               res_type:'print',
                                                                               resource:'Cambio_accosto',
                                                                               pkey: pkey});""",
                                                                               pkey='=#FORM.pkey')
        btn_com_partenza=fb_nsw.Button('!![en]Email Comunicazione Partenza',
                                        action="""genro.publish("table_script_run",{table:"shipsteps.arrival",
                                                                               res_type:'print',
                                                                               resource:'Partenza',
                                                                               pkey: pkey});""",
                                                                               pkey='=#FORM.pkey')
       
        div_nsw2=rg_arrival_nsw.div('<center><strong>NSW ARRIVAL</strong>',width='99%',height='20%',margin='auto',
                        padding='2px',
                        border='1px solid silver',
                        margin_top='1px',margin_left='4px', hidden='^.nsw?=!#v')
        fb_nsw2=div_nsw2.formbuilder(colspan=1,cols=1, border_spacing='1px', fld_width='15em')

        btn_fal1_arr=fb_nsw2.Button('!![en]Fal1 Arrival',
                                        action="""genro.publish("table_script_run",{table:"shipsteps.arrival",
                                                                               res_type:'print',
                                                                               resource:'general_decl',
                                                                               pkey: pkey});""",
                                                                               pkey='=#FORM.pkey')
        btn_nota_arr=fb_nsw2.Button('Nota Arrivo',
                                        action="""genro.publish("table_script_run",{table:"shipsteps.arrival",
                                                                               res_type:'print',
                                                                               resource:'nota_arrivo',
                                                                               pkey: pkey});""",
                                                                               pkey='=#FORM.pkey')
        btn_arrivo = fb_nsw2.Button('!![en]Email arrival')
        #btn_arrivo.dataRpc('nome_temp', self.print_template,record='=#FORM.record',servizio=['capitaneria_nsw'], email_template_id='email_arrivo_cp',
        #                    nome_template = 'shipsteps.arrival:mod61_arr',format_page='A4',nome_vs='=#FORM.record.@vessel_details_id.@imbarcazione_id.nome',
        #                    _ask=dict(title='!![en]Select the Attachments',fields=[dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
        #                     table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
        #                     cols=4,popup=True,colspan=2)]))
        #verifichiamo quanti servizi CP ci sono, nel caso più di uno apparirà la dbSelect per la scelta
        service_for_email = tbl_email_services.query(columns="$service_for_email_id", where='$service_for_email_id=:serv', serv='cp').fetch()
        serv_len=len(service_for_email)
        if serv_len >1:
            btn_arrivo.dataRpc('nome_temp', self.print_template,
                      record='=#FORM.record', servizio=['capitaneria'], email_template_id='email_arrivo_cp',nome_template = 'shipsteps.arrival:mod61_arr',format_page='A4',nome_vs='=#FORM.record.@vessel_details_id.@imbarcazione_id.nome',
                       _ask=dict(title='!![en]Select the services and attachments',fields=[dict(name='services', lbl='!![en]Services', tag='dbSelect',hasDownArrow=True,
                                table='shipsteps.email_services', columns='$consignee', auxColumns='$email,$email_cc,$email_bcc,$email_pec,$email_cc_pec',condition="$service_for_email_id=:cod",condition_cod='cp',alternatePkey='consignee',
                                validate_notnull=True,cols=4,popup=True,colspan=2, hasArrowDown=True),dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                                table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                                cols=4,popup=True,colspan=2)]))
        else:    
            btn_arrivo.dataRpc('nome_temp', self.print_template,
                      record='=#FORM.record', servizio=['capitaneria'], email_template_id='email_arrivo_cp',nome_template = 'shipsteps.arrival:mod61_arr',format_page='A4',nome_vs='=#FORM.record.@vessel_details_id.@imbarcazione_id.nome',
                      _ask=dict(title='!![en]Select the Attachments',fields=[dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                                 table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                                 cols=4,popup=True,colspan=2)]))


        div_nsw3=rg_arrival_nsw.div('<center><strong>NSW DEPARTURE</strong>',width='99%',height='20%',margin='auto',
                        padding='2px',
                        border='1px solid silver',
                        margin_top='1px',margin_left='4px', hidden='^.nsw?=!#v')
        fb_nsw3=div_nsw3.formbuilder(colspan=1,cols=1, border_spacing='1px', fld_width='15em')
        btn_fal1_dep=fb_nsw3.Button('!![en]Fal1 Departure',
                                        action="""genro.publish("table_script_run",{table:"shipsteps.arrival",
                                                                               res_type:'print',
                                                                               resource:'general_decl_dep',
                                                                               pkey: pkey});""",
                                                                               pkey='=#FORM.pkey')
        
        btn_disp_val=fb_nsw3.Button('!![en]Currency availability',hidden="^#FORM.record.@vessel_details_id.@imbarcazione_id.flag?=#v=='IT'",
                                        action="""genro.publish("table_script_run",{table:"shipsteps.arrival",
                                                                               res_type:'print',
                                                                               resource:'dispo_valuta',
                                                                               pkey: pkey});""",
                                                                               pkey='=#FORM.pkey')
        btn_nota_dep=fb_nsw3.Button('Dich.Intergr.Partenza',
                                        action="""genro.publish("table_script_run",{table:"shipsteps.arrival",
                                                                               res_type:'print',
                                                                               resource:'dichiarazione_partenza',
                                                                               pkey: pkey});""",
                                                                               pkey='=#FORM.pkey')
        btn_departure = fb_nsw3.Button('!![en]Email Departure')
        #btn_departure.dataRpc('nome_temp', self.print_template,record='=#FORM.record',servizio=['capitaneria_nsw'], email_template_id='email_partenza_cp',
        #                    nome_template = 'shipsteps.arrival:mod61_dep',format_page='A4',nome_vs='=#FORM.record.@vessel_details_id.@imbarcazione_id.nome',
        #                    _ask=dict(title='!![en]Select the Attachments',fields=[dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
        #                     table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
        #                     cols=4,popup=True,colspan=2)]))
        
        #verifichiamo quanti servizi CP ci sono, nel caso più di uno apparirà la dbSelect per la scelta
        service_for_email = tbl_email_services.query(columns="$service_for_email_id", where='$service_for_email_id=:serv', serv='cp').fetch()
        serv_len=len(service_for_email)
        if serv_len >1:
            btn_departure.dataRpc('nome_temp', self.print_template,
                      record='=#FORM.record', servizio=['capitaneria'], email_template_id='email_partenza_cp',nome_template = 'shipsteps.arrival:mod61_dep',format_page='A4',nome_vs='=#FORM.record.@vessel_details_id.@imbarcazione_id.nome',
                       _ask=dict(title='!![en]Select the services and attachments',fields=[dict(name='services', lbl='!![en]Services', tag='dbSelect',hasDownArrow=True,
                                table='shipsteps.email_services', columns='$consignee', auxColumns='$email,$email_cc,$email_bcc,$email_pec,$email_cc_pec',condition="$service_for_email_id=:cod",condition_cod='cp',alternatePkey='consignee',
                                validate_notnull=True,cols=4,popup=True,colspan=2, hasArrowDown=True),dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                                table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                                cols=4,popup=True,colspan=2)]))
        else:    
            btn_departure.dataRpc('nome_temp', self.print_template,
                      record='=#FORM.record', servizio=['capitaneria'], email_template_id='email_partenza_cp',nome_template = 'shipsteps.arrival:mod61_dep',format_page='A4',nome_vs='=#FORM.record.@vessel_details_id.@imbarcazione_id.nome',
                      _ask=dict(title='!![en]Select the Attachments',fields=[dict(name='allegati', lbl='!![en]Attachments', tag='checkboxtext',
                                 table='shipsteps.arrival_atc', columns='$description',condition="$maintable_id =:cod",condition_cod='=#FORM.record.id',
                                 cols=4,popup=True,colspan=2)]))

        #fb_extra.dataController("""genro.publish("floating_message",{message:'prova'), messageType:"message"}""")
       #genro.publish("floating_message",{message:"Email ready to be sent", messageType:"message"});

        div_extra=rg_extra.div('<center><strong>EXTRA</strong>',width='99%',height='20%',margin='auto',
                        padding='2px',
                        border='1px solid silver',
                        margin_top='1px',margin_left='4px')
        fb_extra=div_extra.formbuilder(colspan=1,cols=1, border_spacing='1px', fld_width='15em')

        #dlg = rg_extra.dialog(nodeId='mydialog',style='width:600px;height:150px;',title='Intestazione fattura',closable=True)
        #dlg.span('Intestazione fatt: ')
        #dlg.span().simpleTextArea(value='^intfat', width='40em', height='50px')
        #dlg.hr()
        #btn_dlg=dlg.button('Visualizza Intestazione')
        #btn_dlg1=dlg.button('Email Intestazione fattura')
        #btn_dlg1.dataRpc('nome_temp', self.email_intfat,record='=#FORM.record',nome_vs='=#FORM.record.@vessel_details_id.@imbarcazione_id.nome')
        #fb_extra.button('Intestazione Fatture', action="genro.wdgById('mydialog').show()")
        #btn_dlg.dataRpc('intfat',self.intfat,record='=#FORM.record')

        fb_extra.div(height='30px').dock(id='mydock_fat')
        div_extra.paletteGrid(paletteCode='Intestazione Fatture',table='shipsteps.invoice_det',viewResource='ViewIntFat', dockTo='mydock_fat',condition='$id=:invid',condition_invid='^#FORM.record.invoice_det_id')
        
        dlgws = rg_extra.dialog(nodeId='dialog_ws',style='width:300px;height:100px;',title='Water supply',closable=True)
        dlgws.span('Quantity mt.: ')
        dlgws.span().textbox(value='^.acqua',table='shipsteps.arrival', columns='$acqua', width='10em')
        dlgws.hr()
        btn_dlgws=dlgws.button('Email request ws')
        btn_dlgws.dataRpc('nome_temp', self.email_ws,record='=#FORM.record',servizio=['ws'], email_template_id='email_water_supply',
                            nome_template = 'shipsteps.arrival:water_supply',nome_vs='=#FORM.record.@vessel_details_id.@imbarcazione_id.nome',
                            _onResult="""if(result=='no_int'){genro.publish('floating_message',{message:"manca l'intestazione: inseriscila",messageType:'error',duration_out:6});}
                                         if(result=='ws')genro.publish("floating_message",{message:"email ready to be sent", messageType:"message"});this.form.save();""")#this.form.save();
                                         #this.form.reload()""")
        fb_extra.button('Water supply', action="genro.wdgById('dialog_ws').show()")
        #avendo creato la formula column nella tasklist per verificare che i documenti del bunker sono stati inviati o no alla cp
        #inseriamo il div che è in riferimento alla valore della formula column 
        fb_extra.div('^.doc_bunker',_virtual_column='@arr_tasklist.doc_bunker', font_weight='bold', color='red')
       

    @public_method
    def intfat(self,record, **kwargs):  
        intfat_id = record['invoice_det_id']
        tbl_invoice = self.db.table('shipsteps.invoice_det')
        int_fat = tbl_invoice.readColumns(columns="""$rag_sociale || coalesce(' - '|| $address, '') || coalesce(' - '|| $cap,'') || coalesce(' - '|| $city,'') || coalesce(' - Vat: ' || $vat,'') || 
                                     coalesce(' - unique code: ' || $cod_univoco,'') || coalesce(' - pec: ' || $pec,'') AS rag_soc""", where='$id=:id_inv',id_inv=intfat_id)
        return int_fat

    def allegatiArrivo(self,pane):
        pane.attachmentGrid(viewResource='ViewFromArrivalAtc')
        #pane.attachmentGallery(viewResource='ViewFromArrivalAtc')
            
    @public_method
    def shorePassLazyMode(self,pane):
        pane.stackTableHandler(relation='@shorepass_arr',formResource='Form',view_store__onBuilt=True)

    @public_method
    def paxListLazyMode(self,pane):
        pane.stackTableHandler(relation='@paxlist_arr',formResource='Form',view_store__onBuilt=True)    

    #def shorepass(self, pane):
    #    pane.stackTableHandler(relation='@shorepass_arr',formResource='Form')

    def garbage(self, pane):
        pane.inlineTableHandler(relation='@garbage_arr',viewResource='ViewFromGarbage')

    def tributi(self, pane):
        pane.dialogTableHandler(relation='@tributi_arr')#,viewResource='ViewFromTributi')

    @public_method
    def rinfusaLazyMode(self,pane):
        pane.stackTableHandler(relation='@rinfusa_arr',formResource='FormFromRinfusa',view_store__onBuilt=True)

    #def rinfusa(self, pane):
    #    pane.stackTableHandler(relation='@rinfusa_arr',formResource='FormFromRinfusa')
    
    @public_method
    def bunkerLazyMode(self,pane):
        pane.stackTableHandler(relation='@bunker_arr',formResource='FormFromBunker',view_store__onBuilt=True)

    #def bunker(self, pane):   
    #    pane.stackTableHandler(relation='@bunker_arr',formResource='FormFromBunker')

    @public_method
    def certificateLazyMode(self,pane):
        pane.stackTableHandler(relation='@istanza_cert_arr',formResource='FormFromCertificates',view_store__onBuilt=True)

    @public_method
    def email_services(self,record,email_template_id=None,servizio=[],nome_temp=None,**kwargs):
    #def email_services(self, record,email_template_id=None,servizio=[],selPkeys_att=None,**kwargs):
        record_arr=record['id']
        arrival_id=record['id']
        flag=record['flag']
        #creiamo la variabile lista attcmt dove tramite il ciclo for andremo a sostituire la parola 'site' con '/home'
        attcmt=[]
        #verifichiamo se abbiamo inserito il numero di protocollo nsis prima di inviare email a sanimare
        if email_template_id == 'email_sanimare':
            if kwargs['nsisprot'] is None or kwargs['nsisprot'] == '':
                nome_temp ='no_nsisprot'
                return nome_temp
        #se il servizio è mod61_arr appendiamo agli attachments il fal1_arr e la nota_arrivo e il mod61
        if nome_temp == 'mod61_arr':
            attcmt.append(self.fal1_path)
            attcmt.append(self.notacp_path)
            file_path_mod61 = 'site:stampe_template/mod61_arr_'+self.vessel_name+'.pdf'
            fileSn_mod61 = self.site.storageNode(file_path_mod61)
            attcmt.append(fileSn_mod61.internal_path)
        #se il servizio è mod61_dep appendiamo agli attachments il fal1_dep e la dichiarazione integrativa di partenza e il mod61
        if nome_temp == 'mod61_dep':
            attcmt.append(self.fal1_path)
            attcmt.append(self.notacp_path)
            #se la bandiera è diversa da quella Italiana appendiamo negli attachments anche la disponibilità valuta
            if flag != 'IT - ITALY':
                attcmt.append(self.dispval_path)
            file_path_mod61 = 'site:stampe_template/mod61_dep_'+self.vessel_name+'.pdf'
            fileSn_mod61 = self.site.storageNode(file_path_mod61)
            attcmt.append(fileSn_mod61.internal_path)
        #se il servizio è form_immigration appendiamo agli attachments il form immigration
        if nome_temp == 'form_immigration':
            file_path_imm = 'site:stampe_template/form_immigration.pdf'
            fileSn_imm = self.site.storageNode(file_path_imm)
            attcmt.append(fileSn_imm.internal_path)    
        #trasformiamo la stringa pkeys allegati in una lista prelevandoli dai kwargs ricevuti tramite bottone
        #ma verifichiamo se nei kwargs gli allegati ci sono per non ritrovarci la variabile lista_all senza assegnazione
        #prendiamo anche il valore dei services che corrisponde al consignee in modo di avere l'intestazione corretta nell'email
        services = None
        std_att= None
        for chiavi in kwargs.keys():
            if chiavi=='allegati':
                if kwargs['allegati']:
                    lista_all=list(kwargs['allegati'].split(","))
                else:
                    lista_all=None
            if chiavi=='services':
                if kwargs['services']:
                    services=kwargs['services']
            if chiavi=='std_att':
                if kwargs['std_att']:
                    std_att=list(kwargs['std_att'].split(","))   
                else:
                    std_att=None   
                               
        #avendo preso il valore services nei kwargs ossia il consignee dell'email sevices andiamo a copiarlo nel record della tasklist nome_servizio 
        record_tasklist=record['@arr_tasklist.id'] 
        tbl_tasklist = self.db.table('shipsteps.tasklist')  
        tbl_tasklist.batchUpdate(dict(nome_servizio=services),
                                    where='$id=:id_task', id_task=record_tasklist)
        self.db.commit()
        
       # with tbl_tasklist.recordToUpdate('id'==record_tasklist) as rec_tasklist:
       #    rec_tasklist['nome_servizio'] = services
                 
       #if kwargs['allegati']:
       #    lista_all=list(kwargs['allegati'].split(","))
       #else:
       #    lista_all=None
        
        #lettura degli attachment
        if lista_all:
            len_allegati = len(lista_all) #verifichiamo la lunghezza della lista pkeys tabella allegati
            file_url=[]
            tbl_att =  self.db.table('shipsteps.arrival_atc') #definiamo la variabile della tabella allegati
            #ciclo for per la lettura dei dati sulla tabella allegati ritornando su ogni ciclo tramite la pkey dell'allegato la colonna $fileurl e alla fine
            #viene appesa alla variabile lista file_url
            for e in range(len_allegati):
                pkeys_att=lista_all[e]
                fileurl = tbl_att.readColumns(columns='$fileurl',
                      where='$id=:att_id',
                        att_id=pkeys_att)
                if fileurl is not None and fileurl !='':
                    file_url.append(fileurl)
        
            ln = len(file_url)
            for r in range(ln):
                fileurl = file_url[r]
                file_path = fileurl.replace('/home','site')
                fileSn = self.site.storageNode(file_path)
                attcmt.append(fileSn.internal_path)
        
        #lettura degli attachment standard in email_services_atc
        if std_att is not None:
            
            len_allegati = len(std_att) #verifichiamo la lunghezza della lista pkeys tabella allegati email_services_atc
            file_url=[]
            tbl_att =  self.db.table('shipsteps.email_services_atc') #definiamo la variabile della tabella allegati
            #ciclo for per la lettura dei dati sulla tabella allegati ritornando su ogni ciclo tramite la pkey dell'allegato la colonna $fileurl e alla fine
            #viene appesa alla variabile lista file_url
            for e in range(len_allegati):
                pkeys_att=std_att[e]
                fileurl = tbl_att.readColumns(columns='$fileurl',
                      where='$id=:att_id',
                        att_id=pkeys_att)
                if fileurl is not None and fileurl !='':
                    file_url.append(fileurl)
        
            ln = len(file_url)
            for r in range(ln):
                fileurl = file_url[r]
                file_path = fileurl.replace('/home','site')
                fileSn = self.site.storageNode(file_path)
                attcmt.append(fileSn.internal_path)
        
        #lettura degli attachment in email_service_atc
        #verifichiamo il numero di servizi
        ln_serv=len(servizio)
        #definiamo le tabelle su cui effettuare le ricerche
        tbl_emailservices = self.db.table('shipsteps.email_services')
        tbl_emailserv_atc = self.db.table('shipsteps.email_services_atc')
        #preleviamo il nome esatto del servizio service_for_email_id
        service_for_email_id = tbl_emailservices.readColumns(columns="$service_for_email_id", where='$service_for_email=:serv AND $consignee=:cons', serv=servizio[0], cons=services)
        #aggiorniamo a true la colonna default del servizio richiesto
        tbl_emailservices.batchUpdate(dict(default=True),
                                    where='$consignee=:cons', cons=services)                         
        #aggiorniamo a false la colonna default degli eventuali servizi aggiunti non richiesti
        tbl_emailservices.batchUpdate(dict(default=False),
                                    where='$consignee != :cons AND $service_for_email_id =:servizio',cons=services, servizio=service_for_email_id)  
                                                                  
        self.db.commit()
        #leggiamo prima gli id dei servizi su email_services così passiamo gli id alla tabella di attachment per la lettura dell'url
        #per poi trasformarlo nel giusto path che appendiamo agli attachment dell'email
        #for e in range(ln_serv):
        #    serv=servizio[e]
        #    service_id = tbl_emailservices.readColumns(columns="$id", where='$service_for_email=:serv AND $consignee=:cons', serv=serv, cons=services)
        #
        #
        #    att_services = tbl_emailserv_atc.query(columns="$filepath", where='$maintable_id=:m_id' ,
        #                                                            m_id=service_id).fetch()
        #    if att_services	!= []:
        #        file_url = att_services[e][0]
        #        file_path = file_url.replace('/home','site')
        #        fileSn = self.site.storageNode(file_path)
        #        attcmt.append(fileSn.internal_path)
         
        #vecchio codice con rilevamento attachments tramite casella checkbox
       #if not record:
       #    return
       #tbl_att =  self.db.table('shipsteps.arrival_atc')
       #fileurl = tbl_att.query(columns='$fileurl',
       #          where='$att_email=:a_att AND $maintable_id=:mt_id',
       #            a_att='true',mt_id=record).fetch()
      
       #ln = len(fileurl)
       #attcmt=[]
       #for r in range(ln):
       #    file_url = fileurl[r][0]
       #    file_path = file_url.replace('/home','site')
       #    fileSn = self.site.storageNode(file_path)
       #    attcmt.append(fileSn.internal_path)

        #Condizioniamo l'aggiunta dell'allegato se il servizio invio email è il garbage
        if servizio==['garbage']:
            file_path_gb = 'site:stampe_template/garbage_request.pdf'
            fileSn_gb = self.site.storageNode(file_path_gb)
            attcmt.append(fileSn_gb.internal_path)
        #Condizioniamo l'aggiunta dell'allegato se il servizio invio email è la deroga rifiuti
        if email_template_id == 'email_deroga_garbage':
            file_path_gb = 'site:stampe_template/deroga_rifiuti.pdf'
            fileSn_gb = self.site.storageNode(file_path_gb)
            attcmt.append(fileSn_gb.internal_path)    
           #if r < (ln-1):
           #    attcmt = attcmt + fileSn.internal_path + ','
           #else:
           #    attcmt = attcmt + fileSn.internal_path

        # Lettura degli account email predefiniti all'interno di Agency e Staff
        tbl_staff =  self.db.table('agz.staff')
        account_email = tbl_staff.readColumns(columns='$email_account_id',
                  where='$agency_id=:ag_id',
                    ag_id=self.db.currentEnv.get('current_agency_id'))
        tbl_agency =  self.db.table('agz.agency')
        account_emailpec = tbl_agency.readColumns(columns='$emailpec_account_id',
                  where='$id=:ag_id',
                    ag_id=self.db.currentEnv.get('current_agency_id'))

      
        #Lettura degli indirizzi email destinatari
        ln_serv=len(servizio)
        
        #email_d, email_cc_d, email_pec_d, email_pec_cc_d=[],[],[],[]
        email_d, email_cc_d,email_bcc_d, email_pec_d, email_pec_cc_d='','','','',''
        tbl_email_services=self.db.table('shipsteps.email_services')

        if services !='' and services is not None:
           
            for e in range(ln_serv):
                serv=servizio[e]

                email_dest, email_cc_dest,email_bcc_dest, email_pec_dest, email_pec_cc_dest = tbl_email_services.readColumns(columns="""$email,$email_cc,$email_bcc,$email_pec,$email_cc_pec""",
                                                        where='$service_for_email=:serv AND $agency_id=:ag_id AND $consignee=:cons', serv=serv, cons=services,
                                                        ag_id=self.db.currentEnv.get('current_agency_id'))

               #email_dest, email_cc_dest,email_bcc_dest, email_pec_dest, email_pec_cc_dest = tbl_email_services.readColumns(columns="""$email,$email_cc,$email_bcc,$email_pec,$email_cc_pec""",
               #                                        where='$service_for_email=:serv AND $agency_id=:ag_id AND $consignee=:cons', serv=serv,cons=services,
               #                                        ag_id=self.db.currentEnv.get('current_agency_id'))                                                    


                if e < (ln_serv-1):
                    if email_dest is not None:
                        email_d = email_d + email_dest + ','
                    if email_cc_dest is not None:
                        email_cc_d = email_cc_d + email_cc_dest + ','
                    if email_bcc_dest is not None:
                        email_bcc_d = email_bcc_d + email_bcc_dest + ','
                    if email_pec_dest is not None:
                        email_pec_d = email_pec_d + email_pec_dest + ','
                    if email_pec_cc_dest is not None:
                        email_pec_cc_d = email_pec_cc_d + email_pec_cc_dest + ','
                else:
                    if email_dest is not None:
                        email_d = email_d + email_dest
                    if email_cc_dest is not None:
                        email_cc_d = email_cc_d + email_cc_dest
                    if email_bcc_dest is not None:
                        email_bcc_d = email_bcc_d + email_bcc_dest
                    if email_pec_dest is not None:
                        email_pec_d = email_pec_d + email_pec_dest
                    if email_pec_cc_dest is not None:
                        email_pec_cc_d = email_pec_cc_d + email_pec_cc_dest
        else:
            
            for e in range(ln_serv):
                serv=servizio[e]

                email_dest, email_cc_dest,email_bcc_dest, email_pec_dest, email_pec_cc_dest = tbl_email_services.readColumns(columns="""$email,$email_cc,$email_bcc,$email_pec,$email_cc_pec""",
                                                        where='$service_for_email=:serv AND $agency_id=:ag_id', serv=serv,
                                                        ag_id=self.db.currentEnv.get('current_agency_id'))

               #email_dest, email_cc_dest,email_bcc_dest, email_pec_dest, email_pec_cc_dest = tbl_email_services.readColumns(columns="""$email,$email_cc,$email_bcc,$email_pec,$email_cc_pec""",
               #                                        where='$service_for_email=:serv AND $agency_id=:ag_id AND $consignee=:cons', serv=serv,cons=services,
               #                                        ag_id=self.db.currentEnv.get('current_agency_id'))                                                    


                if e < (ln_serv-1):
                    if email_dest is not None:
                        email_d = email_d + email_dest + ','
                    if email_cc_dest is not None:
                        email_cc_d = email_cc_d + email_cc_dest + ','
                    if email_bcc_dest is not None:
                        email_bcc_d = email_bcc_d + email_bcc_dest + ','
                    if email_pec_dest is not None:
                        email_pec_d = email_pec_d + email_pec_dest + ','
                    if email_pec_cc_dest is not None:
                        email_pec_cc_d = email_pec_cc_d + email_pec_cc_dest + ','
                else:
                    if email_dest is not None:
                        email_d = email_d + email_dest
                    if email_cc_dest is not None:
                        email_cc_d = email_cc_d + email_cc_dest
                    if email_bcc_dest is not None:
                        email_bcc_d = email_bcc_d + email_bcc_dest
                    if email_pec_dest is not None:
                        email_pec_d = email_pec_d + email_pec_dest
                    if email_pec_cc_dest is not None:
                        email_pec_cc_d = email_pec_cc_d + email_pec_cc_dest
        
        if (email_dest) is not None:
            self.db.table('email.message').newMessageFromUserTemplate(
                                                          record_id=record_arr,
                                                          table='shipsteps.arrival',
                                                          account_id = account_email,
                                                          to_address=email_d,
                                                          cc_address=email_cc_d,
                                                          bcc_address=email_bcc_d,
                                                          attachments=attcmt,
                                                          template_code=email_template_id,
                                                          arrival_id=arrival_id)
            
            self.db.commit()
    
        if (email_pec_dest) is not None:
            self.db.table('email.message').newMessageFromUserTemplate(
                                                          record_id=record_arr,
                                                          table='shipsteps.arrival',
                                                          account_id = account_emailpec,
                                                          to_address=email_pec_d,
                                                          cc_address=email_pec_cc_d,
                                                          attachments=attcmt,
                                                          template_code=email_template_id,
                                                          arrival_id=arrival_id)
            self.db.commit()
        
        
        if (email_dest or email_pec_dest) is not None:

            if email_template_id == 'email_dogana':
                nome_temp = 'val_dog'
            elif email_template_id == 'email_frontiera':
                nome_temp = 'val_imm'
            elif email_template_id == 'email_sanimare':
                nome_temp = 'val_usma'
            elif email_template_id == 'email_pfso':
                nome_temp = 'val_pfso'
            elif email_template_id == 'email_pilot_moor':
                nome_temp = 'val_pil_moor'
            elif email_template_id == 'email_tug':
                nome_temp = 'val_tug'
            elif email_template_id == 'email_tug_dep':
                nome_temp = 'val_tug_dep'      
            elif email_template_id == 'email_chemist':
                nome_temp = 'val_chemist'
            elif email_template_id == 'email_gpg':
                nome_temp = 'val_gpg'
            elif email_template_id == 'email_ens':
                nome_temp = 'val_ens'
            elif email_template_id == 'email_ric_lps':
                nome_temp = 'val_lps'
            elif email_template_id == 'email_lps_cp':
                nome_temp = 'val_lps_cp'    
            elif email_template_id == 'email_garbage_cp':
                nome_temp = 'val_garb_cp'
            elif email_template_id == 'email_integrazione_alim':
                nome_temp = 'val_integr'
            elif email_template_id == 'email_pmou':
                nome_temp = 'val_pmou'
            elif email_template_id == 'email_chimico_cp':
                nome_temp = 'val_chemist_cp'
            elif email_template_id == 'not_rifiuti':
                nome_temp = 'val_adsp'
            elif email_template_id == 'email_arrivo_cp':
                nome_temp = 'val_mod61arr'
            elif email_template_id == 'email_partenza_cp':
                nome_temp = 'val_mod61dep'
            elif email_template_id == 'email_water_supply':
                nome_temp = 'val_ws'
            elif email_template_id == 'email_deroga_garbage':
                nome_temp = 'val_deroga_gb'  
            elif email_template_id == 'email_ricevutarifiuti_cp':
                nome_temp = 'val_ricrifiuti'       
            elif email_template_id == 'email_tributi_cp':
                nome_temp = 'val_tributi'  
            return nome_temp
    
    @public_method
    def email_serv_upd(self, record,email_template_id=None,selPkeys_att=None, **kwargs):
        arrival_id=record['id']
        if not record:
            return
        #lettura del record_id della tabella arrival
        record_id=record['id']
        #lettura dati su tabella arrival
        tbl_arrival = self.db.table('shipsteps.arrival')
        vessel_type,vessel_name,eta_arr,info_moor = tbl_arrival.readColumns(columns='@vessel_details_id.@imbarcazione_id.tipo,@vessel_details_id.@imbarcazione_id.nome,$eta,$info_moor',
                  where='$agency_id=:ag_id AND $id=:rec_id',
                    ag_id=self.db.currentEnv.get('current_agency_id'),rec_id=record_id)
        eta = eta_arr.strftime("%d/%m/%Y, %H:%M")    
        
        #creiamo la variabile lista attcmt dove tramite il ciclo for andremo a sostituire la parola 'site' con '/home'
        attcmt=[]
        
        #verifichiamo che nei kwargs['allegati'] non abbiamo il valore nullo e trasformiamo la stringa pkeys allegati in una lista prelevandoli dai kwargs ricevuti tramite bottone
        if kwargs['allegati'] is not None:
            lista_all = list(kwargs['allegati'].split(","))
        else:
            lista_all = None
        
        #lettura degli attachment
        if lista_all is not None:
            len_allegati = len(lista_all) #verifichiamo la lunghezza della lista pkeys tabella allegati
            file_url=[]
            tbl_att =  self.db.table('shipsteps.arrival_atc') #definiamo la variabile della tabella allegati
            #ciclo for per la lettura dei dati sulla tabella allegati ritornando su ogni ciclo tramite la pkey dell'allegato la colonna $fileurl e alla fine
            #viene appesa alla variabile lista file_url
            for e in range(len_allegati):
                pkeys_att=lista_all[e]
                fileurl = tbl_att.readColumns(columns='$fileurl',
                      where='$id=:att_id',
                        att_id=pkeys_att)
                if fileurl is not None and fileurl !='':
                    file_url.append(fileurl)
        
            ln = len(file_url)
            for r in range(ln):
                fileurl = file_url[r]
                file_path = fileurl.replace('/home','site')
                fileSn = self.site.storageNode(file_path)
                attcmt.append(fileSn.internal_path)

        # Lettura degli account email predefiniti all'interno di Agency e Staff
        tbl_staff =  self.db.table('agz.staff')
        account_email,email_mittente,user_fullname = tbl_staff.readColumns(columns='$email_account_id,@email_account_id.address,$fullname',
                  where='$agency_id=:ag_id',
                    ag_id=self.db.currentEnv.get('current_agency_id'))
        tbl_agency =  self.db.table('agz.agency')
        agency_name,ag_fullstyle,account_emailpec,emailpec_mitt = tbl_agency.readColumns(columns='$agency_name,$fullstyle,$emailpec_account_id, @emailpec_account_id.address',
                  where='$id=:ag_id',
                    ag_id=self.db.currentEnv.get('current_agency_id'))
        #preleviamo dai kwargs i servizi per gli aggiornamenti
        services=kwargs['services']
        #trasformiamo la stringa services in una lista
        servizio=list(services.split(","))
        #troviamo la lunghezza della variabile servizio
        #print(x)
        ln_serv=len(servizio)
        #assegnamo le varibili liste per inserire successivamente i risultati della ricerca sulla tabella email_services
        destinatario,destinatario_pec,email_d, email_cc_d,email_bcc_d, email_pec_d, email_pec_cc_d=[],[],[],[],[],[],[]
        #definiamo la variabile contentente la tabella email_services
        tbl_email_services=self.db.table('shipsteps.email_services')
        #con il ciclo for ad ogni passaggio otteniamo il nome del servizio che passeremo alla query e i risultati saranno appesi alle liste
        for e in range(ln_serv):
            serv=servizio[e]

            dest,email_dest, email_cc_dest,email_bcc_dest = tbl_email_services.readColumns(columns="""$consignee,$email,$email_cc,$email_bcc""",
                                                    where="$service_for_email_id=:serv AND $agency_id=:ag_id", serv=serv,
                                                    ag_id=self.db.currentEnv.get('current_agency_id'))
            dest_pec,email_pec_dest, email_pec_cc_dest = tbl_email_services.readColumns(columns="""$consignee,$email_pec,$email_cc_pec""",
                                                    where='$service_for_email_id=:serv AND $agency_id=:ag_id AND $email_pec IS NOT NULL', serv=serv,
                                                    ag_id=self.db.currentEnv.get('current_agency_id'))
            if dest is not None and dest !='':
                destinatario.append('a: ' + dest)
            if dest_pec is not None and dest_pec !='':
                destinatario_pec.append('a: ' + dest_pec)
            #verifichiamo se il servizio è solo uno sarà inserito all'email destinatario to:    
            if email_dest is not None and email_dest !='' and ln_serv==1:
                email_d.append(email_dest)
            #verifichiamo se il servizio è più di uno sarà inserito all'email bcc:     
            if email_dest is not None and email_dest !='' and ln_serv>1:
                email_bcc_d.append(email_dest)
            #verifichiamo se il servizio è solo uno sarà inserito all'email destinatario to:    
            if email_cc_dest is not None and email_cc_dest != '' and ln_serv==1:
                email_cc_d.append(email_cc_dest)
             #verifichiamo se il servizio è più di uno sarà inserito all'email bcc:    
            if email_cc_dest is not None and email_cc_dest != '' and ln_serv>1:
                email_bcc_d.append(email_cc_dest)    
            if email_bcc_dest is not None and email_bcc_dest != '':    
                email_bcc_d.append(email_bcc_dest)
            if email_pec_dest is not None and email_pec_dest != '':
                email_pec_d.append(email_pec_dest)
            if email_pec_cc_dest is not None and email_pec_cc_dest !='':
                email_pec_cc_d.append(email_pec_cc_dest)
        #trasformiamo le liste in stringhe assegnandole alle relative variabili
        consignee='<br>'.join([str(item) for item in destinatario])
        consignee_pec='<br>'.join([str(item) for item in destinatario_pec])
        if ln_serv == 1:
            email_to = ','.join([str(item) for item in email_d])
        else:
            email_to = email_mittente    
        email_cc = ','.join([str(item) for item in email_cc_d])
        email_bcc = ','.join([str(item) for item in email_bcc_d])
        email_pec = ','.join([str(item) for item in email_pec_d])
        email_pec_cc = ','.join([str(item) for item in email_pec_cc_d])
        
        now = datetime.now()
        cur_time = now.strftime("%H:%M:%S")    
        if cur_time < '13:00:00':
            sal='Buongiorno,'  
        elif cur_time < '17:00:00':
            sal='Buon pomeriggio'
        elif cur_time < '24:00:00':
            sal = 'Buonasera,' 
        elif cur_time < '04:00:00':
            sal = 'Buona notte,'      

        subject='aggiornamento '+vessel_type + ' ' + vessel_name + ' ref:' + record['reference_num']
        body_header="""<span style="font-family:courier new,courier,monospace;">""" + 'da: '+ agency_name + '<br>' + consignee + '<br><br>'
        body_header_pec="""<span style="font-family:courier new,courier,monospace;">""" + 'da: '+ agency_name + '<br>' + consignee_pec + '<br><br>'
        body_footer= '<br>Cordiali saluti<br><br>' + user_fullname + '<br><br>' + ag_fullstyle + """</span></div>"""
        if info_moor is None:
            info_moor=''
        body_msg=(sal + '<br>' + "con la presente si comunica che il nuovo ETA della " +vessel_type + ' ' + vessel_name + " è il " + eta + '<br><br>' + info_moor)
        body_html=(body_header + body_msg + body_footer )
        
        if email_to:
            self.db.table('email.message').newMessage(account_id=account_email,
                           to_address=email_to,
                           from_address=email_mittente,
                           subject=subject, body=body_html, 
                           cc_address=email_cc, 
                           bcc_address=email_bcc, attachments=attcmt,arrival_id=arrival_id,
                           html=True)
            self.db.commit()
        body_html_pec=(body_header_pec + body_msg + body_footer )
        if email_pec is not None and email_pec!='':
            
            self.db.table('email.message').newMessage(account_id=account_emailpec,
                           to_address=email_pec,
                           from_address=emailpec_mitt,
                           subject=subject, body=body_html_pec, 
                           cc_address=email_pec_cc, 
                           attachments=attcmt,arrival_id=arrival_id,
                           html=True)
            self.db.commit()
        if (email_dest or email_pec_dest) is not None:
            nome_temp='val_upd'
            return nome_temp
       # print(x)

    @public_method
    def email_intfat(self, record, **kwargs):
        arrival_id=record['id']
        if not record:
            return
        #lettura del record_id della tabella arrival
        record_id=record['id']
        vessel_type = record['@vessel_details_id.@imbarcazione_id.tipo']
        vessel_name = record['@vessel_details_id.@imbarcazione_id.nome']
        intfat_id = record['invoice_det_id']
        tbl_invoice = self.db.table('shipsteps.invoice_det')
        int_fat = tbl_invoice.readColumns(columns="""$rag_sociale || coalesce(' - '|| $address, '') || coalesce(' - '|| $cap,'') || coalesce(' - '|| $city,'') || coalesce(' - Vat: ' || $vat,'') || 
                                     coalesce(' - unique code: ' || $cod_univoco,'') || coalesce(' - pec: ' || $pec,'') AS rag_soc""", where='$id=:id_inv',id_inv=intfat_id)
        #rag_soc,indirizzo,cap,citta,vat,cf,cod_un,pec = tbl_invoice.readColumns(columns='$rag_sociale,$address,$cap,$city,$vat,$cf,$cod_univoco,$pec', where='$id=:id_inv',id_inv=intfat_id)
        if int_fat == [None, None, None, None, None, None, None]:
            nome_temp='no_int'
            return nome_temp
        
        # Lettura degli account email predefiniti all'interno di Agency e Staff
        tbl_staff =  self.db.table('agz.staff')
        account_email,email_mittente,user_fullname = tbl_staff.readColumns(columns='$email_account_id,@email_account_id.address,$fullname',
                  where='$agency_id=:ag_id',
                    ag_id=self.db.currentEnv.get('current_agency_id'))
        tbl_agency =  self.db.table('agz.agency')
        agency_name,ag_fullstyle,account_emailpec,emailpec_mitt = tbl_agency.readColumns(columns='$agency_name,$fullstyle,$emailpec_account_id, @emailpec_account_id.address',
                  where='$id=:ag_id',
                    ag_id=self.db.currentEnv.get('current_agency_id'))
        
        
        now = datetime.now()
        cur_time = now.strftime("%H:%M:%S")    
        if cur_time < '13:00:00':
            sal='Buongiorno,'  
        elif cur_time < '17:00:00':
            sal='Buon pomeriggio'
        elif cur_time < '24:00:00':
            sal = 'Buonasera,' 
        elif cur_time < '04:00:00':
            sal = 'Buona notte,'      
       
        subject='Intestazione fatture '+vessel_type + ' ' + vessel_name + ' ref:' + record['reference_num']
        #body_header="""<span style="font-family:courier new,courier,monospace;">""" + 'da: '+ agency_name + '<br>' + consignee + '<br><br>'
        body_footer= 'Cordiali saluti<br><br>' + user_fullname + '<br><br>' + ag_fullstyle + """</span></div>"""
        
        body_msg=("""<span style="font-family:courier new,courier,monospace;">""" + sal + '<br>' + "con la presente siamo a girarVi intestazione fatture per " +vessel_type + ' ' + vessel_name + " :" + '<br><br>' +
                        int_fat + '<br><br>')
        body_html=(body_msg + body_footer )
                
        self.db.table('email.message').newMessage(account_id=account_email,
                           to_address='',
                           from_address=email_mittente,
                           subject=subject, body=body_html, 
                           cc_address='',arrival_id=arrival_id,
                           html=True)
        self.db.commit()
        
        nome_temp='intfat'
        return nome_temp

    @public_method
    def email_ws(self, record, **kwargs):
        arrival_id=record['id']
        result = Bag()
        if not record:
            return
        #lettura del record_id della tabella arrival
        record_id=record['id']
        vessel_type = record['@vessel_details_id.@imbarcazione_id.tipo']
        vessel_name = record['@vessel_details_id.@imbarcazione_id.nome']
        intfat_id = record['invoice_det_id']
        qt_ws = record['@arr_tasklist.acqua']
        tbl_invoice = self.db.table('shipsteps.invoice_det')
        int_fat = tbl_invoice.readColumns(columns="""$rag_sociale ||' '|| coalesce($address, '') ||' '|| coalesce($cap,'') ||' '|| coalesce($city,'') || coalesce(' Vat: ' || $vat,'') || 
                                     coalesce(' unique code: ' || $cod_univoco,'') || coalesce(' pec: ' || $pec,'') AS rag_soc""", where='$id=:id_inv',id_inv=intfat_id)
        #rag_soc,indirizzo,cap,citta,vat,cf,cod_un,pec = tbl_invoice.readColumns(columns='$rag_sociale,$address,$cap,$city,$vat,$cf,$cod_univoco,$pec', where='$id=:id_inv',id_inv=intfat_id)
        if int_fat == [None, None, None, None, None, None, None]:
            nome_temp='no_int'
            result['int'] = 'no'
           
            return nome_temp
        
        # Lettura degli account email predefiniti all'interno di Agency e Staff
        tbl_staff =  self.db.table('agz.staff')
        account_email,email_mittente,user_fullname = tbl_staff.readColumns(columns='$email_account_id,@email_account_id.address,$fullname',
                  where='$agency_id=:ag_id',
                    ag_id=self.db.currentEnv.get('current_agency_id'))
        tbl_agency =  self.db.table('agz.agency')
        agency_name,ag_fullstyle,account_emailpec,emailpec_mitt = tbl_agency.readColumns(columns='$agency_name,$fullstyle,$emailpec_account_id, @emailpec_account_id.address',
                  where='$id=:ag_id',
                    ag_id=self.db.currentEnv.get('current_agency_id'))
        #preleviamo dai kwargs i servizi per gli aggiornamenti
        servizio=kwargs['servizio']
        #trasformiamo la stringa services in una lista
        #servizio=list(services.split(","))
        #troviamo la lunghezza della variabile servizio
        #print(x)
        ln_serv=len(servizio)
        #assegnamo le varibili liste per inserire successivamente i risultati della ricerca sulla tabella email_services
        destinatario,destinatario_pec,email_d, email_cc_d,email_bcc_d, email_pec_d, email_pec_cc_d=[],[],[],[],[],[],[]
        #definiamo la variabile contentente la tabella email_services
        tbl_email_services=self.db.table('shipsteps.email_services')
        #con il ciclo for ad ogni passaggio otteniamo il nome del servizio che passeremo alla query e i risultati saranno appesi alle liste
        for e in range(ln_serv):
            serv=servizio[e]

            dest,email_dest, email_cc_dest,email_bcc_dest = tbl_email_services.readColumns(columns="""$consignee,$email,$email_cc,$email_bcc""",
                                                    where="$service_for_email_id=:serv AND $agency_id=:ag_id", serv=serv,
                                                    ag_id=self.db.currentEnv.get('current_agency_id'))
            dest_pec,email_pec_dest, email_pec_cc_dest = tbl_email_services.readColumns(columns="""$consignee,$email_pec,$email_cc_pec""",
                                                    where='$service_for_email_id=:serv AND $agency_id=:ag_id AND $email_pec IS NOT NULL', serv=serv,
                                                    ag_id=self.db.currentEnv.get('current_agency_id'))
            if dest is not None and dest !='':
                destinatario.append('a: ' + dest)
            if dest_pec is not None and dest_pec !='':
                destinatario_pec.append('a: ' + dest_pec)
            if email_dest is not None and email_dest !='':
                email_d.append(email_dest)
            if email_cc_dest is not None and email_cc_dest != '':
                email_cc_d.append(email_cc_dest)
            if email_bcc_dest is not None and email_bcc_dest != '':    
                email_bcc_d.append(email_bcc_dest)
            if email_pec_dest is not None and email_pec_dest != '':
                email_pec_d.append(email_pec_dest)
            if email_pec_cc_dest is not None and email_pec_cc_dest !='':
                email_pec_cc_d.append(email_pec_cc_dest)
        #trasformiamo le liste in stringhe assegnandole alle relative variabili
        consignee='<br>'.join([str(item) for item in destinatario])
        consignee_pec='<br>'.join([str(item) for item in destinatario_pec])
        email_to = ','.join([str(item) for item in email_d])
        email_cc = ','.join([str(item) for item in email_cc_d])
        email_bcc = ','.join([str(item) for item in email_bcc_d])
        email_pec = ','.join([str(item) for item in email_pec_d])
        email_pec_cc = ','.join([str(item) for item in email_pec_cc_d])
        
        now = datetime.now()
        cur_time = now.strftime("%H:%M:%S")    
        if cur_time < '13:00:00':
            sal='Buongiorno,'  
        elif cur_time < '17:00:00':
            sal='Buon pomeriggio'
        elif cur_time < '24:00:00':
            sal = 'Buonasera,' 
        elif cur_time < '04:00:00':
            sal = 'Buona notte,'      
        
        subject='Richiesta fornitura acqua '+vessel_type + ' ' + vessel_name + ' ref:' + record['reference_num']
        #body_header="""<span style="font-family:courier new,courier,monospace;">""" + 'da: '+ agency_name + '<br>' + consignee + '<br><br>'
        body_footer= 'Cordiali saluti<br><br>' + user_fullname + '<br><br>' + ag_fullstyle + """</span></div>"""
        
        body_msg=("""<span style="font-family:courier new,courier,monospace;">""" + sal + '<br>' + "con la presente siamo a richiederVi fornitura di tonn."+qt_ws +" di acqua poatbile per " +vessel_type + ' ' + vessel_name + " :" + '<br><br>' +
                       'Potete fatturare a:<br>' + int_fat + '<br><br> e inviare a:<br>' + agency_name + '<br><br>')
        body_html=(body_msg + body_footer )
        #print(x)
        
        self.db.table('email.message').newMessage(account_id=account_email,
                           to_address=email_to,
                           from_address=email_mittente,
                           subject=subject, body=body_html, 
                           cc_address=email_cc,arrival_id=arrival_id,
                           html=True)
        self.db.commit()
        
        nome_temp='ws'
        return nome_temp

    @public_method
    def email_arrival_sof(self, record,email_template_id=None,servizio=[],selPkeys_att=None, **kwargs):
        tbl_arrival = self.db.table('shipsteps.arrival')
        #verifichiamo che ci sia il record
        if not record:
            return
        record_arr=record['id']
        arrival_id=record['id']
        #creiamo la variabile lista attcmt dove tramite il ciclo for andremo a sostituire la parola 'site' con '/home'
        attcmt=[]
        #verifichiamo che nei kwargs['allegati'] non abbiamo il valore nullo e trasformiamo la stringa pkeys allegati in una lista prelevandoli dai kwargs ricevuti tramite bottone
        if kwargs['allegati'] is not None:
            lista_all = list(kwargs['allegati'].split(","))
        else:
            lista_all = None
        
        #lettura degli attachment
        if lista_all is not None:
            len_allegati = len(lista_all) #verifichiamo la lunghezza della lista pkeys tabella allegati
            file_url=[]
            tbl_att =  self.db.table('shipsteps.arrival_atc') #definiamo la variabile della tabella allegati
            #ciclo for per la lettura dei dati sulla tabella allegati ritornando su ogni ciclo tramite la pkey dell'allegato la colonna $fileurl e alla fine
            #viene appesa alla variabile lista file_url
            for e in range(len_allegati):
                pkeys_att=lista_all[e]
                fileurl = tbl_att.readColumns(columns='$fileurl',
                      where='$id=:att_id',
                        att_id=pkeys_att)
                if fileurl is not None and fileurl !='':
                    file_url.append(fileurl)
        
            ln = len(file_url)
            for r in range(ln):
                fileurl = file_url[r]
                file_path = fileurl.replace('/home','site')
                fileSn = self.site.storageNode(file_path)
                attcmt.append(fileSn.internal_path) 

        # Lettura degli account email predefiniti all'interno di Agency e Staff
        tbl_staff =  self.db.table('agz.staff')
        account_email,email_mittente = tbl_staff.readColumns(columns='$email_account_id,@email_account_id.address',
                  where='$agency_id=:ag_id',
                    ag_id=self.db.currentEnv.get('current_agency_id'))
        tbl_agency =  self.db.table('agz.agency')
        account_emailpec = tbl_agency.readColumns(columns='$emailpec_account_id',
                  where='$id=:ag_id',
                    ag_id=self.db.currentEnv.get('current_agency_id'))

        #verifichiamo e assegnamo alla variabile sof_id arrivatoci dal bottone di invio email
        if kwargs:
            sof_id=kwargs['sof_id']
            if sof_id is None:
                nome_temp='no_sof'
                return nome_temp
        else:
            return
        #inizializziamo le variabili per le email
        email_arr_to,email_arr_cc,email_arr_bcc='','',''
        email_a_to,email_a_cc,email_a_bcc=[],[],[]
        #definiamo le tabelle dove prelevare l'email
        tbl_email_sof=self.db.table('shipsteps.email_sof')
        tbl_email_arr=self.db.table('shipsteps.email_arr')
        #verifichiamo la lunghezza del servizio arrivatoci tramite il bottone di invio email e con i cicli for preleviamo i
        #dati email dalle relative tabelle- tramite l'uso delle liste con gli append aggiungiamo l'email 
        ln_serv=len(servizio)
        for e in range(ln_serv):
            serv=servizio[e]
            if serv=='arr':
                email_to = tbl_email_arr.query(columns="$email",
                                                    where='$arrival_id=:a_id and $email_type=:type', a_id=record_arr,
                                                    type='to').fetch()
                for e in range(len(email_to)):
                    email_a_to.append(email_to[e][0])

                email_cc = tbl_email_arr.query(columns="$email",
                                                    where='$arrival_id=:a_id and $email_type=:type', a_id=record_arr,
                                                    type='cc').fetch()  
                for e in range(len(email_cc)):
                    email_a_cc.append(email_cc[e][0])

                email_bcc = tbl_email_arr.query(columns="$email",
                                                    where='$arrival_id=:a_id and $email_type=:type', a_id=record_arr,
                                                    type='ccn').fetch()
                for e in range(len(email_bcc)):
                    email_a_bcc.append(email_bcc[e][0])
                
            elif serv=='sof':
                email_to = tbl_email_sof.query(columns="$email",
                                                    where='$sof_id=:s_id and $email_type=:type', s_id=sof_id,
                                                    type='to').fetch()  
                for e in range(len(email_to)):
                    email_a_to.append(email_to[e][0])

                email_cc = tbl_email_sof.query(columns="$email",
                                                    where='$sof_id=:s_id and $email_type=:type', s_id=sof_id,
                                                    type='cc').fetch()
                for e in range(len(email_cc)):
                    email_a_cc.append(email_cc[e][0]) 

                email_bcc = tbl_email_sof.query(columns="$email",
                                                    where='$sof_id=:s_id and $email_type=:type', s_id=sof_id,
                                                    type='ccn').fetch()  
                for e in range(len(email_bcc)):
                    email_a_bcc.append(email_bcc[e][0])
        #estraiamo le stringhe email dalle liste
        email_arr_to=','.join([str(item) for item in email_a_to])
        email_arr_cc=','.join([str(item) for item in email_a_cc])                    
        email_arr_bcc=','.join([str(item) for item in email_a_bcc])
        #verifichiamo che non mancano email destinatari TO e CCN altrimenti ritorniamo con la variabile nome_temp che innesca il messaggio
        #verifichiamo se non presente l'email to allora inseriamo l'email del mittente
        if email_arr_to == email_arr_bcc == '':  
            nome_temp = 'no_email'
            return nome_temp
        elif email_arr_to =='':    
            email_arr_to=email_mittente 
        #creiamo il nuovo messaggio e con il db.commit lo salviamo nella tabella di uscita email pronto per l'invio                                        
        self.db.table('email.message').newMessageFromUserTemplate(
                                                          record_id=sof_id,
                                                          table='shipsteps.sof',
                                                          account_id = account_email,
                                                          to_address=email_arr_to,
                                                          cc_address=email_arr_cc,
                                                          bcc_address=email_arr_bcc,
                                                          attachments=attcmt,
                                                          template_code=email_template_id,
                                                          arrival_id=arrival_id)
        self.db.commit()
        #ritorniamo con la variabile nome_temp per l'innesco del messaggio e il settaggio della checklist invio email a vero
        if email_template_id == 'email_updating_shiprec':
            nome_temp = 'ship_rec_upd'
        else:
            nome_temp = 'ship_rec'
        return nome_temp

    @public_method
    def email_arrdep(self, record,email_template_id=None,servizio=[],selPkeys_att=None, **kwargs):
        record_arr=record['id']
        arrival_id=record['id']
        #verifichiamo che ci sia il record
   
        if not record:
            return
        #creiamo la variabile lista attcmt dove tramite il ciclo for andremo a sostituire la parola 'site' con '/home'
        attcmt=[]

        #verifichiamo che nei kwargs['allegati'] non abbiamo il valore nullo e trasformiamo la stringa pkeys allegati in una lista prelevandoli dai kwargs ricevuti tramite bottone
        if kwargs['allegati'] is not None:
            lista_all = list(kwargs['allegati'].split(","))
        else:
            lista_all = None
        
        #verifichiamo nei kwargs['template'] il valore assegnato dalla nostra scelta al lancio dell'email ossia l'email_template_id
        if 'template' in kwargs.keys():
            email_template_id=kwargs['template']

        #lettura degli attachment
        if lista_all is not None:
            len_allegati = len(lista_all) #verifichiamo la lunghezza della lista pkeys tabella allegati
            file_url=[]
            tbl_att =  self.db.table('shipsteps.arrival_atc') #definiamo la variabile della tabella allegati
            #ciclo for per la lettura dei dati sulla tabella allegati ritornando su ogni ciclo tramite la pkey dell'allegato la colonna $fileurl e alla fine
            #viene appesa alla variabile lista file_url
            for e in range(len_allegati):
                pkeys_att=lista_all[e]
                fileurl = tbl_att.readColumns(columns='$fileurl',
                      where='$id=:att_id',
                        att_id=pkeys_att)
                if fileurl is not None and fileurl !='':
                    file_url.append(fileurl)
        
            ln = len(file_url)
            for r in range(ln):
                fileurl = file_url[r]
                file_path = fileurl.replace('/home','site')
                fileSn = self.site.storageNode(file_path)
                attcmt.append(fileSn.internal_path)

        # Lettura degli account email predefiniti all'interno di Agency e Staff
        tbl_staff =  self.db.table('agz.staff')
        account_email,email_mittente = tbl_staff.readColumns(columns='$email_account_id,@email_account_id.address',
                  where='$agency_id=:ag_id',
                    ag_id=self.db.currentEnv.get('current_agency_id'))
        tbl_agency =  self.db.table('agz.agency')
        account_emailpec = tbl_agency.readColumns(columns='$emailpec_account_id',
                  where='$id=:ag_id',
                    ag_id=self.db.currentEnv.get('current_agency_id'))

        #inizializziamo le variabili per le email
        email_arr_to,email_arr_cc,email_arr_bcc='','',''
        email_a_to,email_a_cc,email_a_bcc=[],[],[]
        #definiamo le tabelle dove prelevare l'email
        tbl_email_sof=self.db.table('shipsteps.email_sof')
        tbl_email_arr=self.db.table('shipsteps.email_arr')
        #verifichiamo la lunghezza del servizio arrivatoci tramite il bottone di invio email e con i cicli for preleviamo i
        #dati email dalle relative tabelle- tramite l'uso delle liste con gli append aggiungiamo l'email 
        ln_serv=len(servizio)
        for e in range(ln_serv):
            serv=servizio[e]
            if serv=='arr':
                email_to = tbl_email_arr.query(columns="$email",
                                                    where='$arrival_id=:a_id and $email_type=:type', a_id=record_arr,
                                                    type='to').fetch()
                for e in range(len(email_to)):
                    email_a_to.append(email_to[e][0])

                email_cc = tbl_email_arr.query(columns="$email",
                                                    where='$arrival_id=:a_id and $email_type=:type', a_id=record_arr,
                                                    type='cc').fetch()  
                for e in range(len(email_cc)):
                    email_a_cc.append(email_cc[e][0])

                email_bcc = tbl_email_arr.query(columns="$email",
                                                    where='$arrival_id=:a_id and $email_type=:type', a_id=record_arr,
                                                    type='ccn').fetch()
                for e in range(len(email_bcc)):
                    email_a_bcc.append(email_bcc[e][0])
                
        #estraiamo le stringhe email dalle liste
        email_arr_to=','.join([str(item) for item in email_a_to])
        email_arr_cc=','.join([str(item) for item in email_a_cc])                    
        email_arr_bcc=','.join([str(item) for item in email_a_bcc])
        #verifichiamo che non mancano email destinatari TO e CCN altrimenti ritorniamo con la variabile nome_temp che innesca il messaggio
        #verifichiamo se non presente l'email to allora inseriamo l'email del mittente
        if email_arr_to == email_arr_bcc == '':  
            nome_temp = 'no_email'
            return nome_temp
        elif email_arr_to =='':    
            email_arr_to=email_mittente 
        #creiamo il nuovo messaggio e con il db.commit lo salviamo nella tabella di uscita email pronto per l'invio    
        #print(x)                                    
        self.db.table('email.message').newMessageFromUserTemplate(
                                                          record_id=record_arr,
                                                          table='shipsteps.arrival',
                                                          account_id = account_email,
                                                          to_address=email_arr_to,
                                                          cc_address=email_arr_cc,
                                                          bcc_address=email_arr_bcc,
                                                          attachments=attcmt,
                                                          template_code=email_template_id,
                                                          arrival_id=arrival_id)
        
        self.db.commit()
        #ritorniamo con la variabile nome_temp per l'innesco del messaggio e il settaggio della checklist invio email a vero
        #if email_template_id == 'email_ormeggio' or email_template_id == 'email_operations' or email_template_id == 'email_operations_mov' or email_template_id == 'email_partenza':
        if email_template_id != '' or email_template_id != None:
            nome_temp = 'val_upd'
            return nome_temp

    @public_method
    def print_template(self, record, resultAttr=None, nome_template=None, email_template_id=None,servizio=[],  nome_vs=None, format_page=None, **kwargs):
        record_arr=record['id']
        flag=record['flag']
        #verifichiamo che stiamo stampando la checklist e che tipo di movimentazione è stata assegnato all'arrivo
        #al fine di assegnare il nome template della check list 
        if nome_template=='shipsteps.arrival:check_list' and record['@tip_mov.code'] == 'alim':
            nome_template='shipsteps.arrival:check_list_alim'
        if nome_template=='shipsteps.arrival:check_list' and record['@tip_mov.code'] == 'pass':
            nome_template='shipsteps.arrival:check_list_pass'
        # Crea stampa
        self.vessel_name = nome_vs
        tbl_arrival = self.db.table('shipsteps.arrival')
        builder = TableTemplateToHtml(table=tbl_arrival)
        #nome_template = nome_template #'shipsteps.arrival:check_list'

        nome_temp = nome_template.replace('shipsteps.arrival:','')
        if nome_template == 'shipsteps.arrival:mod61_arr' or nome_template == 'shipsteps.arrival:mod61_dep':
            nome_file = '{cl_id}.pdf'.format(
                    cl_id=nome_temp +'_' + nome_vs)
        else:
            nome_file = '{cl_id}.pdf'.format(
                    cl_id=nome_temp)
        #nome_file_st = 'laboratorio_piazza_sta_{cl_id}.pdf'.format(
        #    cl_id=self.avatar.user_id)
        template = self.loadTemplate(nome_template)  # nome del template
        pdfpath = self.site.storageNode('home:stampe_template', nome_file)
        
       #tbl_template=self.db.table('adm.htmltemplate')
       #letterhead = tbl_template.readColumns(columns='$id',
       #          where='$name=:tp_name', tp_name='A3_orizz')
        # (pdfpath.internal_path)
        #print(x)
       #if kwargs:
       #    letterhead=kwargs['letterhead_id']
       #else:
       #    letterhead=''
        
        #verifichiamo se nelle keys di kwargs troviamo la chiave pratique_date e lo assegnamo alla variabile data_pratica 
        data_pratica = None
        for chiavi in kwargs.keys():
            
            if chiavi=='pratique_date':
                if kwargs['pratique_date']:
                    data_pratica=kwargs['pratique_date']
        #avendo preso il valore data_pratica nei kwargs andiamo a copiarlo nel record della tasklist p_date che ci servirà nella variabile 
        #che utilizzeremo nei template per avere la data della form                     
        record_tasklist=record['@arr_tasklist.id'] 
        tbl_tasklist = self.db.table('shipsteps.tasklist')  
        tbl_tasklist.batchUpdate(dict(p_date=data_pratica),
                                    where='$id=:id_task', id_task=record_tasklist)
        self.db.commit()

        #Verifichiamo nel caso stampa sia del mod61_arr se ci sono i file da allegare in cartella altrimenti ritorniamo con il msg di errore
        if nome_temp == 'mod61_arr':
            nome_fal1arr = 'Fal1_arr_' + nome_vs
            nome_notarr = 'Nota_arrivo_' + nome_vs
            self.fal1_path = self.site.site_path+'/stampe_template/'+nome_fal1arr+'.pdf'
            self.notacp_path = self.site.site_path+'/stampe_template/'+nome_notarr+'.pdf'
            if not os.path.isfile(self.fal1_path):
                fal1_arrival = 'no'
            else:
                fal1_arrival = 'yes'    
            if not os.path.isfile(self.notacp_path):
                nota_arrivo = 'no'
            else:
                nota_arrivo = 'yes'
           
            if fal1_arrival == 'no' and nota_arrivo == 'no':
                nome_temp = 'fal1arr_notarr'  
                return nome_temp      
            elif fal1_arrival=='no':
                nome_temp = 'fal1_arr_no'
                return nome_temp
            elif nota_arrivo == 'no':
                nome_temp='nota_arr_no'
                return nome_temp
        
        if nome_temp == 'mod61_dep':
            sailed= self.db.table('shipsteps.arrival').readColumns(columns="$ets", where='$id=:a_id',a_id=record_arr)
            if sailed is None:
                nome_temp = 'no_sailed'
                return nome_temp
            nome_fal1dep = 'Fal1_dep_' + nome_vs
            nome_notdep = 'Dich_integrativa_partenza_' + nome_vs
            self.fal1_path = self.site.site_path+'/stampe_template/'+nome_fal1dep+'.pdf'
            self.notacp_path = self.site.site_path+'/stampe_template/'+nome_notdep+'.pdf'
            #verifica del file disponibilità valuta se la bandiera è diversa da IT
            nome_dispval=None
            if flag != 'IT - ITALY':
                nome_dispval = 'Disp_valuta_' + nome_vs
                self.dispval_path = self.site.site_path+'/stampe_template/'+nome_dispval+'.pdf'
                if not os.path.isfile(self.dispval_path):
                    dispval_departure = 'no'
                else:
                    dispval_departure = 'yes'  
            #verifica del file fal1              
            if not os.path.isfile(self.fal1_path):
                fal1_departure = 'no'
            else:
                fal1_departure = 'yes'
            #verifica del file notacp        
            if not os.path.isfile(self.notacp_path):
                nota_partenza = 'no'
            else:
                nota_partenza = 'yes'
            
            #verifica condizione dei file fal1 e notacp e disp_valuta se la bandiera è diversa da IT 
            if flag != 'IT - ITALY':
                if fal1_departure == 'no' and nota_partenza == 'no' and dispval_departure == 'no':
                    nome_temp = 'fal1dep_notapart_dispval'  
                    return nome_temp      
                elif fal1_departure=='no':
                    nome_temp = 'fal1_dep_no'
                    return nome_temp
                elif nota_partenza == 'no':
                    nome_temp='nota_part_no'
                    return nome_temp
                elif dispval_departure == 'no':
                    nome_temp='dispval_no'
                    return nome_temp
            #verifica condizione dei file fal1 e notacp           
            if fal1_departure == 'no' and nota_partenza == 'no':
                nome_temp = 'fal1dep_notapart'  
                return nome_temp      
            elif fal1_departure=='no':
                nome_temp = 'fal1_dep_no'
                return nome_temp
            elif nota_partenza == 'no':
                nome_temp='nota_part_no'
                return nome_temp
            
                
        tbl_htmltemplate = self.db.table('adm.htmltemplate')
        templates= tbl_htmltemplate.query(columns='$id,$name', where='').fetch()
        letterhead=''       
        for r in range(len(templates)):
            if templates[r][1] == 'A4_vert':
                letterhead = templates[r][0]    
            if format_page=='A3':
                if templates[r][1] == 'A3_orizz':
                    letterhead = templates[r][0]
         
        builder(record=record_arr, template=template,letterhead_id=letterhead)
        
        #if format_page=='A3':
        #    builder.page_format='A3'
        #    builder.page_width=427
        #    builder.page_height=290
#
        result = builder.writePdf(pdfpath=pdfpath)
        
        self.setInClientData(path='gnr.clientprint',
                             value=result.url(timestamp=datetime.now()), fired=True)
                               
       #inseriamo nella tabella di attachment il mod61_arr 
       #if nome_temp == 'mod61_arr':
       #    tbl_arrival_atc = self.db.table('shipsteps.arrival_atc')
       #    if not tbl_arrival_atc.checkDuplicate(maintable_id=record,description=nome_file):
       #        
       #        tbl_arrival_atc.addAttachment(maintable_id=record,
       #                                     origin_filepath=pdfpath,
       #                                     description=nome_file,
       #                                     copyFile=True)

        #inviamo l'email se si tratta di mod61_arr e rispetta le condizioni dei file da allegare
        if nome_temp == 'mod61_arr':
            if  fal1_arrival == 'yes' and nota_arrivo == 'yes':
                self.email_services(record,email_template_id,servizio, nome_temp, **kwargs)
                return nome_temp
        #inviamo l'email se si tratta di mod61_dep e rispetta le condizioni dei file da allegare
        if nome_temp == 'mod61_dep':
            if flag != 'IT - ITALY':
                if  fal1_departure == 'yes' and nota_partenza == 'yes' and dispval_departure == 'yes':
                    self.email_services(record,email_template_id,servizio, nome_temp, **kwargs)
                    return nome_temp       
            elif  fal1_departure == 'yes' and nota_partenza == 'yes':
                self.email_services(record,email_template_id,servizio, nome_temp, **kwargs)
                return nome_temp
        #inviamo l'email se si tratta di immigration form e rispetta le condizioni dei file da allegare
        if kwargs:
            for chiave in kwargs.keys():
                if chiave == 'only_print':
                    if kwargs['only_print'] == 'yes':
                        nome_temp = 'form_immigration_print'
                        return nome_temp
        if nome_temp == 'form_immigration' :
            self.email_services(record,email_template_id,servizio, nome_temp, **kwargs)
            return nome_temp        
        return nome_temp

    @public_method
    def print_template_garbage(self, record, resultAttr=None,selId=None, nome_template=None, email_template_id=None,servizio=[] , format_page=None, **kwargs):
        #msg_special=None
        selPkeys_att=kwargs['selPkeys_att']
        if selId is None:
            nome_temp = 'yes'
            return nome_temp

        tbl_garbage = self.db.table('shipsteps.garbage')
        builder = TableTemplateToHtml(table=tbl_garbage)

        nome_temp = nome_template.replace('shipsteps.garbage:','')
        nome_file = '{cl_id}.pdf'.format(
                    cl_id=nome_temp)

        template = self.loadTemplate(nome_template)  # nome del template
        pdfpath = self.site.storageNode('home:stampe_template', nome_file)

        tbl_htmltemplate = self.db.table('adm.htmltemplate')
        templates= tbl_htmltemplate.query(columns='$id,$name', where='').fetch()
        letterhead=''       
        for r in range(len(templates)):
            if templates[r][1] == 'A4_vert':
                letterhead = templates[r][0]    
            if format_page=='A3':
                if templates[r][1] == 'A3_orizz':
                    letterhead = templates[r][0]
          
        builder(record=selId, template=template,letterhead_id=letterhead)
        #builder(record=selId, template=template)
        #if format_page=='A3':
        #    builder.page_format='A3'
        #    builder.page_width=427
        #    builder.page_height=290

        result = builder.writePdf(pdfpath=pdfpath)

        self.setInClientData(path='gnr.clientprint',
                              value=result.url(timestamp=datetime.now()), fired=True)
        self.email_services(record,email_template_id,servizio, **kwargs)
        #se ritorna il valore di nome_temp dalla funzione sopra lanciata self.email_services
        # facciamo ritornare il valore di nome_temp alla chiamata iniziale del bottone di stampa per far scattare
        # il msg con il dataController
        nome_temp='val_garbage'
        return nome_temp    
    
    @public_method
    def print_template_gate(self, record, resultAttr=None,selId=None, nome_template=None, email_template_id=None,servizio=[] , format_page=None, **kwargs):
        #msg_special=None
        ag_id=record['agency_id']
        tbl_gate = self.db.table('shipsteps.opening_gate')
        record_gate = tbl_gate.query(columns='$id',where='$agency_id=:a_id',
                                                                    a_id=ag_id).fetch()            

        builder = TableTemplateToHtml(table=tbl_gate)
        storagePath=[]
        for r in range(len(record_gate)):
            nome_temp = nome_template.replace('shipsteps.opening_gate:','')+str(r)
            nome_file = '{cl_id}.pdf'.format(
                        cl_id=nome_temp)

            template = self.loadTemplate(nome_template)  # nome del template
            pdfpath = self.site.storageNode('home:stampe_template', nome_file)
            
            storagePath.append(pdfpath.fullpath)
            record=record_gate[r][0]

            tbl_htmltemplate = self.db.table('adm.htmltemplate')
            templates= tbl_htmltemplate.query(columns='$id,$name', where='').fetch()
            letterhead=''       
            for r in range(len(templates)):
                if templates[r][1] == 'A4_vert':
                    letterhead = templates[r][0]    
                if format_page=='A3':
                    if templates[r][1] == 'A3_orizz':
                        letterhead = templates[r][0]

            builder(record=record, template=template,letterhead_id=letterhead)
        #builder(record=selId, template=template)
        #if format_page=='A3':
        #    builder.page_format='A3'
        #    builder.page_width=427
        #    builder.page_height=290

            result = builder.writePdf(pdfpath=pdfpath)
        builder.pdf_handler.joinPdf(storagePath,'home:stampe_template/master_info.pdf')
        master_info = self.site.storageNode('home:stampe_template', 'master_info.pdf' )
        self.setInClientData(path='gnr.clientprint',
                              value=master_info.url(timestamp=datetime.now()), fired=True)
        
        #self.email_services(record,email_template_id,servizio, **kwargs)
        #se ritorna il valore di nome_temp dalla funzione sopra lanciata self.email_services
        # facciamo ritornare il valore di nome_temp alla chiamata iniziale del bottone di stampa per far scattare
        # il msg con il dataController
        nome_temp='master_info'
        return nome_temp    
    
    @public_method
    def print_template_derogagb(self, record, imbarcazione_id=None,resultAttr=None,selId=None,moored=None,nextport=None,nome_template=None, email_template_id=None,servizio=[] , format_page=None, **kwargs):
        #msg_special=None
        #facciamo arrivare alla variabile moored la datetime dell'ormeggio e se non presente torna indietro il valore no_moored per far scattare il dataController   
        if moored is None or moored == '':
            nome_temp = 'no_moored'
            return nome_temp
        #verifichiamo la variabile nextport e se non presente o con valore 'ORDER' o 'N/A' o 'MARE' torna indietro il valore no_next_port per far scattare il dataController
        if nextport is None or nextport == 'ORDER' or nextport == 'N/A' or nextport == 'MARE':
            nome_temp = 'no_next_port'
            return nome_temp
        
        tbl_arrival = self.db.table('shipsteps.arrival')
        builder = TableTemplateToHtml(table=tbl_arrival)

        nome_temp = nome_template.replace('shipsteps.arrival:','')
        nome_file = '{cl_id}.pdf'.format(
                    cl_id=nome_temp)

        template = self.loadTemplate(nome_template)  # nome del template
        pdfpath = self.site.storageNode('home:stampe_template', nome_file)

        tbl_htmltemplate = self.db.table('adm.htmltemplate')
        templates= tbl_htmltemplate.query(columns='$id,$name', where='').fetch()
        letterhead=''       
        for r in range(len(templates)):
            if templates[r][1] == 'A4_vert':
                letterhead = templates[r][0]    
            if format_page=='A3':
                if templates[r][1] == 'A3_orizz':
                    letterhead = templates[r][0]
          
        builder(record=record, template=template,letterhead_id=letterhead)

        #builder(record=record, template=template)
        #if format_page=='A3':
        #    builder.page_format='A3'
        #    builder.page_width=427
        #    builder.page_height=290

        result = builder.writePdf(pdfpath=pdfpath)

        self.setInClientData(path='gnr.clientprint',
                              value=result.url(timestamp=datetime.now()), fired=True)
        tbl_bolli = self.db.table('shipsteps.bolli')
        agency_id = record['agency_id']
        
        if not tbl_bolli.checkDuplicate(istanza='Deroga Rifiuti',ref_number=record['reference_num']):
            nuovo_record = dict(date=datetime.now(),imbarcazione_id=imbarcazione_id,istanza='Deroga Rifiuti',
                                ref_number=record['reference_num'],bolli_tr14=1,bolli_tr22=1,agency_id=agency_id)
            tbl_bolli.insert(nuovo_record)
            self.db.commit()    

        self.email_services(record,email_template_id,servizio, **kwargs)
        #se ritorna il valore di nome_temp dalla funzione sopra lanciata self.email_services
        # facciamo ritornare il valore di self.ms_special alla chiamata iniziale del bottone di stampa per far scattare
        # il msg con il dataController
        nome_temp='val_deroga_gb'
        return nome_temp    
    
    @public_method
    def apridoc(self,record,nome_form=None, **kwargs):
        workport=record['workport']
        eta=record['eta'].strftime("%d/%m/%Y")
        etb=record['etb'].strftime("%d/%m/%Y")
        lastport=record['lastport']
        nextport=record['nextport']
        vesselname=record['vesselname']
        flag=record['flag']
        imo=record['imo']
        gt=record['tsl']
        master_name=record['master_name']
        crew_n=str(record['n_crew'])
        pax_n=str(record['n_passengers'])
        vessel_details_id=record['vessel_details_id']
        workdate = self.db.workdate.strftime("%d/%m/%Y")
        if nome_form=='DichSanimare':
            #cerchiamo nella tabella certificati nave la sanitation
            tbl_shipsdoc = self.db.table('shipsteps.ship_doc')
            sanitation = tbl_shipsdoc.query(columns="$issued,to_char($date_cert,:df)", where='$cert=:cert and $vessel_details=:vess_det', 
                                                            cert='06_sanitation',vess_det=vessel_details_id, df='DD/MM/YYYY').fetch() 
            if sanitation == []:
                nome_temp ='no_sanitation'
                return nome_temp
            san_place_id = sanitation[0][0]
            san_date = sanitation[0][1]
            #san_place_id,san_date=tbl_shipsdoc.readColumns(columns="$issued,to_char($date_cert,:df)", where='$cert=:cert and $vessel_details=:vess_det', 
            #                                                cert='06_sanitation',vess_det=vessel_details_id, df='DD/MM/YYYY')
            #tramite l'id del luogo di rilascio del certificato andiamo a cercare nella tabella degli unlocode il place
            tbl_place = self.db.table('unlocode.place')
            san_place=tbl_place.readColumns(columns="$descrizione || ' - ' || @nazione_code.nome", where='$id=:place_id', place_id=san_place_id)
        
            nome_file = 'DichSanimare.docx'
            nome_file_out = 'DichSanimare_filled.docx'
            file_sn_out = self.site.storageNode('home:form_standard', 'DichSanimare_filled.docx')
           
            variables = {
            "${porto}": workport,
            "${date_arr}": eta,
            "${etb}": etb,
            "${nome_imb}": vesselname,
            "${imo}": imo,
            "${last_port}": lastport,
            "${next_port}": nextport,
            "${flag}": flag,
            "${master_name}": master_name,
            "${gt}": gt,
            "${sanitation_place}": san_place,
            "${date_sanitation}": str(san_date),
            "${crew_n}": crew_n,
            "${pax_n}": pax_n,
            "${current_date}": workdate,
            "${medico}": "/////",
            }
        if nome_form == 'InterferenzeFiore':
            nome_file = 'InterferenzeFiore.docx'
            nome_file_out = 'InterferenzeFiore_filled.docx'
            file_sn_out = self.site.storageNode('home:form_standard', 'InterferenzeFiore_filled.docx')    
            variables = {
            "${etb}": etb,
            "${nome_imb}": vesselname,
            }
        file_sn = self.site.storageNode('home:form_standard', nome_file)
        template_file_path = file_sn.internal_path
        #template_file_path = '/home/tommaso/Documenti/Linux/Python/ModificaDocx/test.docx'
        
        output_file_path = file_sn_out.internal_path
        #output_file_path = '/home/tommaso/Documenti/Linux/Python/ModificaDocx/result.docx'
        
        

        template_document = Document(template_file_path)

        for variable_key, variable_value in variables.items():
            for paragraph in template_document.paragraphs:
                self.replace_text_in_paragraph(paragraph, variable_key, variable_value)

            for table in template_document.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            self.replace_text_in_paragraph(paragraph, variable_key, variable_value)
        
        template_document.save(output_file_path)
        #apriamo direttamente il file salvato con il programma standard di sistema
        filename=output_file_path
       
       # subprocess.call(('xdg-open', filename))
        
      # path=self.site.site_path + str('/form_standard')
      # doc_path=filename
      # subprocess.call(['libreoffice',
      #          # '--headless',
      #          '--convert-to',
      #          'pdf',
      #          '--outdir',
      #          path,
      #          doc_path])
      # nome_file = nome_form + str('.pdf')  
              
        path_pdf = self.site.storageNode('home:form_standard', nome_file_out)
        #path_pdf=path + str('/') + nome_form + str('.pdf')
        result=self.site.storageNode(path_pdf)
        self.setInClientData(path='gnr.clientprint',
                              value=result.url(), fired=True)
        #print(x)
        return nome_form

    def replace_text_in_paragraph(self,paragraph, key, value):
        if key in paragraph.text:
            inline = paragraph.runs
            for item in inline:
                if key in item.text:
                    item.text = item.text.replace(key, value)
        


    
    def th_options(self):
        return dict(dialog_height='400px', dialog_width='600px', duplicate=True)
   